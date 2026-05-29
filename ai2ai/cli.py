"""
AI Debate Pipeline v5 â€” Expert Council
=======================================
StrukturĂˇlt tĂ¶bbfĂˇzisĂş vita tĂ¶bb AI modell kĂ¶zĂ¶tt.

SCENARIĂ“K:
  quick          â€” 3 modell, 1 fĂĽggetlen kĂ¶r + szintĂ©zis (DEFAULT)
  expert-council â€” 5 modell, 5 fĂˇzis: evidence â†’ fĂĽggetlen â†’ issue matrix â†’ rebuttal â†’ revision â†’ judge
  red-team       â€” 1 strategist + 1 engineer + 2 skeptic + judge

FĂZISOK (expert-council):
  0. Evidence Pack    â€” tĂ¶mĂ¶rĂ­tett forrĂˇscsomagot kap mindenki, nem nyers fĂˇjlokat
  1. FĂĽggetlen vĂ©lemĂ©nyek â€” modellek NEM lĂˇtjĂˇk egymĂˇst (opcionĂˇlisan pĂˇrhuzamosan)
  2. ModerĂˇtor Issue Matrix â€” konszenzus/vita tĂ©rkĂ©p, 3-5 kĂ©rdĂ©s
  3. CĂ©lzott Rebuttal â€” csak az issue matrixra reagĂˇlnak
  4. Position Revision â€” JSON: changed_mind, confidence, remaining_disagreement
  5. Final Judge       â€” kĂĽlĂ¶n modell a vĂ©gsĹ‘ szintĂ©zishez (nem a moderĂˇtor)

KAPCSOLĂ“K:
  --scenario quick|expert-council|red-team
  --quality fast|balanced|best
  --parallel          pĂˇrhuzamos Phase 1 hĂ­vĂˇsok
  --roles key=model   szerepek felĂĽlĂ­rĂˇsa (pl. judge=gpt,skeptic=deepseek)
  --estimate          csak kĂ¶ltsĂ©gbecslĂ©s
  --resume LOG.json   megszakadt futĂˇs folytatĂˇsa

HASZNĂLAT:
  python ai_debate.py project.zip --prompt "..." --scenario expert-council
  python ai_debate.py --folder ./docs --prompt "..." --quality best --parallel
  python ai_debate.py ... --roles judge=gpt,moderator=gemini --scenario red-team

API KULCSOK (.env):
  ANTHROPIC_API_KEY=sk-ant-...
  OPENAI_API_KEY=sk-...
  GOOGLE_API_KEY=...
  DEEPSEEK_API_KEY=sk-...
  XAI_API_KEY=xai-...

FĂśGGĹSĂ‰GEK:
  pip install anthropic openai google-genai python-docx openpyxl python-dotenv
  npm install -g docx
"""

from __future__ import annotations

import argparse, base64, concurrent.futures, json, mimetypes, os, re, tempfile
import shutil, subprocess, sys, tempfile, zipfile
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(errors="replace")
    except Exception:
        pass

# â”€â”€ .env â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _load_dotenv():
    try:
        from dotenv import load_dotenv
    except ImportError:
        return
    for folder in [Path(__file__).resolve().parent, *Path(__file__).resolve().parents]:
        if (folder / ".env").exists():
            load_dotenv(dotenv_path=folder / ".env", override=False)
            print(f"đź”‘ .env: {folder / '.env'}")
            return
    home = Path.home() / ".env"
    if home.exists():
        load_dotenv(dotenv_path=home, override=False)
_load_dotenv()

try:
    import anthropic
except ImportError:
    print("âťŚ pip install anthropic"); sys.exit(1)
try:
    from openai import OpenAI
except ImportError:
    print("âťŚ pip install openai"); sys.exit(1)
try:
    from docx import Document as DocxReader
except ImportError:
    print("âťŚ pip install python-docx"); sys.exit(1)


# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Modell katalĂłgus
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Modell katalĂłgus â€” ID-k .env-bĹ‘l tĂ¶ltĹ‘dnek, fallback az id mezĹ‘re
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _env(key, default):
    return os.environ.get(key, default)

def _build_catalog():
    return {
        # â”€â”€ Claude â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        "claude-opus": {
            "id":      _env("CLAUDE_BEST_MODEL",     "claude-opus-4-7"),
            "id_fast": _env("CLAUDE_FAST_MODEL",     "claude-sonnet-4-6"),
            "label": "Claude Opus", "emoji": "đźźŁ", "color": "\033[95m",
            "type": "anthropic", "pin": 15.0, "pout": 75.0,
        },
        "claude-sonnet": {
            "id":      _env("CLAUDE_BALANCED_MODEL", "claude-sonnet-4-6"),
            "id_fast": _env("CLAUDE_FAST_MODEL",     "claude-sonnet-4-6"),
            "label": "Claude Sonnet", "emoji": "đź”·", "color": "\033[96m",
            "type": "anthropic", "pin": 3.0, "pout": 15.0,
        },
        "claude-haiku": {
            "id":      _env("CLAUDE_FAST_MODEL",     "claude-haiku-4-5-20251001"),
            "id_fast": _env("CLAUDE_FAST_MODEL",     "claude-haiku-4-5-20251001"),
            "label": "Claude Haiku", "emoji": "đź”ą", "color": "\033[34m",
            "type": "anthropic", "pin": 0.8, "pout": 4.0,
        },
        # â”€â”€ OpenAI â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        # Responses API â€” reasoning_effort paramĂ©terrel
        "gpt-fast": {
            "id":      _env("OPENAI_FAST_MODEL",     "gpt-5.4-mini"),
            "id_fast": _env("OPENAI_FAST_MODEL",     "gpt-5.4-mini"),
            "label": "GPT Fast", "emoji": "đźź˘", "color": "\033[92m",
            "type": "openai_responses",
            "reasoning_effort": _env("OPENAI_REASONING_EFFORT_FAST", "low"),
            "pin": 0.40, "pout": 1.60,
        },
        "gpt": {
            "id":      _env("OPENAI_BALANCED_MODEL", "gpt-5.4"),
            "id_fast": _env("OPENAI_FAST_MODEL",     "gpt-5.4-mini"),
            "label": "GPT-5.4", "emoji": "đźź˘", "color": "\033[92m",
            "type": "openai_responses",
            "reasoning_effort": _env("OPENAI_REASONING_EFFORT_BALANCED", "medium"),
            "pin": 2.0, "pout": 8.0,
        },
        "gpt-best": {
            "id":      _env("OPENAI_BEST_MODEL",     "gpt-5.5"),
            "id_fast": _env("OPENAI_BALANCED_MODEL", "gpt-5.4"),
            "label": "GPT-5.5", "emoji": "đźź˘", "color": "\033[32m",
            "type": "openai_responses",
            "reasoning_effort": _env("OPENAI_REASONING_EFFORT_BEST", "high"),
            "pin": 5.0, "pout": 20.0,
        },
        # â”€â”€ Gemini â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        "gemini": {
            "id":      _env("GEMINI_BALANCED_MODEL", "gemini-3.1-pro-preview"),
            "id_fast": _env("GEMINI_FAST_MODEL",     "gemini-2.5-flash"),
            "label": "Gemini", "emoji": "đź”µ", "color": "\033[94m",
            "type": "gemini", "pin": 1.25, "pout": 5.0,
        },
        "gemini-fast": {
            "id":      _env("GEMINI_FAST_MODEL",     "gemini-2.5-flash"),
            "id_fast": _env("GEMINI_FAST_MODEL",     "gemini-2.5-flash"),
            "label": "Gemini Flash", "emoji": "đź”µ", "color": "\033[94m",
            "type": "gemini", "pin": 0.15, "pout": 0.60,
        },
        # â”€â”€ DeepSeek â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        # deepseek-chat/reasoner retire 2026-07-24 â€” v4 string-ek kellenek
        "deepseek": {
            "id":      _env("DEEPSEEK_BALANCED_MODEL", "deepseek-v4-flash"),
            "id_fast": _env("DEEPSEEK_FAST_MODEL",     "deepseek-v4-flash"),
            "label": "DeepSeek", "emoji": "đźźˇ", "color": "\033[93m",
            "type": "openai_compat",
            "base_url": "https://api.deepseek.com/v1", "env_key": "DEEPSEEK_API_KEY",
            "pin": 0.14, "pout": 0.28,
        },
        "deepseek-pro": {
            "id":      _env("DEEPSEEK_BEST_MODEL",   "deepseek-v4-pro"),
            "id_fast": _env("DEEPSEEK_BALANCED_MODEL","deepseek-v4-flash"),
            "label": "DeepSeek Pro", "emoji": "đźźˇ", "color": "\033[33m",
            "type": "openai_compat",
            "base_url": "https://api.deepseek.com/v1", "env_key": "DEEPSEEK_API_KEY",
            "pin": 0.435, "pout": 1.74,
        },
        # â”€â”€ Grok â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        # grok-4/grok-4-fast retired 2026-05-15; grok-4.3 az ajĂˇnlott default
        "grok": {
            "id":      _env("GROK_BALANCED_MODEL",   "grok-4.3"),
            "id_fast": _env("GROK_FAST_MODEL",       "grok-4.3"),
            "label": "Grok", "emoji": "đź”´", "color": "\033[91m",
            "type": "openai_compat",
            "base_url": "https://api.x.ai/v1", "env_key": "XAI_API_KEY",
            "pin": 1.25, "pout": 2.50,
        },
        "grok-best": {
            "id":      _env("GROK_BEST_MODEL",       "grok-4.20"),
            "id_fast": _env("GROK_BALANCED_MODEL",   "grok-4.3"),
            "label": "Grok Max", "emoji": "đź”´", "color": "\033[31m",
            "type": "openai_compat",
            "base_url": "https://api.x.ai/v1", "env_key": "XAI_API_KEY",
            "pin": 3.0, "pout": 15.0,
        },
    }

CATALOG = _build_catalog()

# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Scenario definĂ­ciĂłk
# Szerepek: moderator, strategist, engineer, skeptic, market, judge
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
SCENARIOS = {
    # â”€â”€ Quick: 3 modell, 1 fĂĽggetlen kĂ¶r + judge â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    "quick": {
        "desc": "3 modell, 1 fĂĽggetlen kĂ¶r + szintĂ©zis (default)",
        "phases": ["evidence", "independent", "judge"],
        "quality_map": {
            "fast": {
                "moderator":  "claude-sonnet",
                "debater1":   "gpt-fast",
                "debater2":   "gemini-fast",
                "judge":      "gpt-fast",
            },
            "balanced": {
                "moderator":  "claude-sonnet",
                "debater1":   "gpt",
                "debater2":   "gemini",
                "judge":      "gpt",
            },
            "best": {
                "moderator":  "claude-sonnet",
                "debater1":   "gpt-best",
                "debater2":   "gemini",
                "judge":      "claude-opus",
            },
        },
    },
    # â”€â”€ Expert Council: 5 fĂˇzis, 6 szereplĹ‘ â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    "expert-council": {
        "desc": "5 fĂˇzis: evidence â†’ fĂĽggetlen â†’ issue matrix â†’ rebuttal â†’ revision â†’ judge",
        "phases": ["evidence", "independent", "issue_matrix", "rebuttal", "revision", "judge"],
        "quality_map": {
            "fast": {
                "moderator":       "claude-sonnet",
                "strategist":      "gpt-fast",
                "engineer":        "claude-sonnet",
                "market_analyst":  "gemini-fast",
                "skeptic":         "grok",
                "cost_reasoner":   "deepseek",
                "judge":           "gpt-fast",
            },
            "balanced": {
                "moderator":       "claude-sonnet",
                "strategist":      "gpt",
                "engineer":        "claude-sonnet",
                "market_analyst":  "gemini",
                "skeptic":         "grok",
                "cost_reasoner":   "deepseek",
                "judge":           "gpt",
            },
            "best": {
                "moderator":       "claude-sonnet",
                "strategist":      "gpt-best",
                "engineer":        "claude-opus",
                "market_analyst":  "gemini",
                "skeptic":         "grok-best",
                "cost_reasoner":   "deepseek-pro",
                "judge":           "gpt-best",
            },
        },
    },
    # â”€â”€ Red Team: kritikus dĂ¶ntĂ©s, 2 szkeptikus â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    "red-team": {
        "desc": "1 strategist + 1 engineer + 2 skeptic + judge â€” biztonsĂˇgi/jogi/befektetĹ‘i dĂ¶ntĂ©sekhez",
        "phases": ["evidence", "independent", "issue_matrix", "rebuttal", "revision", "judge"],
        "quality_map": {
            "fast": {
                "moderator":  "claude-sonnet",
                "strategist": "gpt-fast",
                "engineer":   "claude-sonnet",
                "skeptic":    "grok",
                "skeptic2":   "deepseek",
                "judge":      "gpt-fast",
            },
            "balanced": {
                "moderator":  "claude-sonnet",
                "strategist": "gpt",
                "engineer":   "claude-sonnet",
                "skeptic":    "grok",
                "skeptic2":   "deepseek",
                "judge":      "gpt",
            },
            "best": {
                "moderator":  "claude-sonnet",
                "strategist": "gpt-best",
                "engineer":   "claude-opus",
                "skeptic":    "grok-best",
                "skeptic2":   "deepseek-pro",
                "judge":      "claude-opus",
            },
        },
    },
    # â”€â”€ Build Plan: implementĂˇciĂłs roadmap, sprint terv â”€â”€â”€â”€â”€â”€â”€
    "build-plan": {
        "desc": "Product + Senior Architect + Security + QA + PM szintĂ©zis â€” roadmap Ă©s sprint promptokhoz",
        "phases": ["evidence", "independent", "issue_matrix", "rebuttal", "revision", "judge"],
        "quality_map": {
            "fast": {
                "moderator":   "claude-sonnet",
                "product_mgr": "gpt-fast",
                "architect":   "claude-sonnet",
                "security":    "deepseek",
                "qa_reviewer": "gemini-fast",
                "judge":       "gpt-fast",
            },
            "balanced": {
                "moderator":   "claude-sonnet",
                "product_mgr": "gpt",
                "architect":   "claude-sonnet",
                "security":    "deepseek",
                "qa_reviewer": "gemini",
                "judge":       "gpt",
            },
            "best": {
                "moderator":   "claude-sonnet",
                "product_mgr": "gpt-best",
                "architect":   "claude-opus",
                "security":    "grok-best",
                "qa_reviewer": "gemini",
                "judge":       "claude-opus",
            },
        },
    },
}

RESET = "\033[0m"; BOLD = "\033[1m"; DIM = "\033[2m"; YELLOW = "\033[93m"

def _c(t, code): return f"{code}{t}{RESET}"
def _bold(t):   return _c(t, BOLD)
def _dim(t):    return _c(t, DIM)
def _yellow(t): return _c(t, YELLOW)
def _col(t, cfg): return _c(t, cfg.get("color",""))

def _enable_win_ansi():
    if sys.platform == "win32":
        try:
            import ctypes
            ctypes.windll.kernel32.SetConsoleMode(
                ctypes.windll.kernel32.GetStdHandle(-11), 7)
        except Exception: pass
_enable_win_ansi()

def _env_int(name: str, default: int) -> int:
    raw = os.environ.get(name)
    if raw is None or str(raw).strip() == "":
        return default
    try:
        return int(raw)
    except ValueError:
        print(_yellow(f"   Invalid {name}={raw!r}; using {default}"))
        return default

MAX_TOKENS     = _env_int("MAX_OUTPUT_TOKENS", 2000)
SYNTHESIS_MAX_OUTPUT_TOKENS = _env_int("SYNTHESIS_MAX_OUTPUT_TOKENS", 4000)
EVIDENCE_TOKENS = 3000
DEFAULT_MAX_CHARS = _env_int("MAX_SOURCE_CHARS", 10000)
MIN_VALID_PARTICIPANTS = _env_int("MIN_VALID_PARTICIPANTS", 2)
INFRASTRUCTURE_ROLES = {"moderator", "judge"}
ROLE_FALLBACKS = {
    "gemini": ["claude-sonnet", "gpt", "deepseek", "grok"],
    "openai_responses": ["claude-sonnet", "gemini", "deepseek", "grok"],
    "openai": ["claude-sonnet", "gemini", "deepseek", "grok"],
    "openai_compat": ["gpt", "claude-sonnet", "gemini", "grok"],
    "anthropic": ["gpt", "gemini", "deepseek", "grok"],
    "judge": ["gpt-best", "gpt", "claude-opus", "claude-sonnet", "gemini", "deepseek-pro", "grok-best"],
    "moderator": ["claude-sonnet", "claude-opus", "gpt", "gemini"],
}
RUN_METADATA = {}


TASK_TYPES = {
    "business_strategy",
    "technical_audit",
    "implementation_plan",
    "architecture_decision",
    "product_roadmap",
    "red_team_review",
    "document_review",
    "general_analysis",
}

TASK_TYPE_PERSPECTIVES = {
    "business_strategy": ["strategic", "market", "financial", "technical", "critical"],
    "technical_audit": ["architecture", "security", "operations", "maintainability", "critical"],
    "implementation_plan": ["product", "architecture", "engineering", "delivery", "critical"],
    "architecture_decision": ["architecture", "engineering", "security", "operations", "critical"],
    "product_roadmap": ["product", "market", "user_experience", "engineering", "critical"],
    "red_team_review": ["security", "adversarial", "operational", "legal", "critical"],
    "document_review": ["editorial", "domain_expert", "clarity", "risk", "critical"],
    "general_analysis": ["strategic", "technical", "critical"],
}

TASK_TYPE_DELIVERABLES = {
    "business_strategy": ["final_synthesis", "business_strategy", "implementation_plan", "ai_context_block"],
    "technical_audit": ["final_synthesis", "technical_findings", "risk_register", "remediation_plan"],
    "implementation_plan": ["final_synthesis", "implementation_plan", "task_roadmap", "ai_context_block"],
    "architecture_decision": ["final_synthesis", "architecture_decision_record", "tradeoff_analysis", "implementation_plan"],
    "product_roadmap": ["final_synthesis", "product_roadmap", "prioritized_backlog", "risks"],
    "red_team_review": ["final_synthesis", "threat_findings", "risk_register", "mitigation_plan"],
    "document_review": ["final_synthesis", "review_notes", "recommended_edits", "open_questions"],
    "general_analysis": ["final_synthesis"],
}


@dataclass
class TaskProfile:
    task_type: str
    user_goal: str
    target_audience: str
    language: str
    recommended_scenario: str
    required_perspectives: list[str] = field(default_factory=list)
    expected_deliverables: list[str] = field(default_factory=list)
    constraints: list[str] = field(default_factory=list)
    assumptions: list[str] = field(default_factory=list)
    profile_source: str = "deterministic"


def _infer_language(user_prompt: str, cli_lang: str = "hu") -> str:
    text = (user_prompt or "").lower()
    hu_markers = ("Ăˇ", "Ă©", "Ă­", "Ăł", "Ă¶", "Ĺ‘", "Ăş", "ĂĽ", "Ĺ±", " Ă©s ", " hogy ", " terv", "feladat")
    en_markers = (" the ", " and ", " architecture", " implementation", " review", " roadmap")
    if any(m in text for m in hu_markers):
        return "hu"
    if any(m in text for m in en_markers):
        return "en"
    return cli_lang or "hu"


def _infer_task_type(user_prompt: str, items: list = None) -> str:
    text = (user_prompt or "").lower()
    names = " ".join(getattr(i, "name", "") for i in (items or [])).lower()
    blob = f"{text} {names}"
    if any(k in blob for k in ("red team", "red-team", "threat", "fenyeget", "attack", "tĂˇmadĂˇs", "security review")):
        return "red_team_review"
    if any(k in blob for k in ("architecture decision", "technical architecture", "system architecture", "adr", "architektĂşra dĂ¶ntĂ©s", "architectural decision", "tradeoff")):
        return "architecture_decision"
    if any(k in blob for k in ("technical audit", "code audit", "audit", "refactor", "tech debt", "technikai audit", "kĂłdaudit")):
        return "technical_audit"
    if any(k in blob for k in ("business plan", "business strategy", "ĂĽzleti terv", "ĂĽzleti stratĂ©gia", "gtm", "go-to-market", "monetiz")):
        return "business_strategy"
    if any(k in blob for k in ("roadmap", "product roadmap", "backlog", "termĂ©k roadmap", "product plan")):
        return "product_roadmap"
    if any(k in blob for k in ("implementation plan", "megvalĂłsĂ­tĂˇsi terv", "sprint", "task roadmap", "build plan")):
        return "implementation_plan"
    if any(k in blob for k in ("document review", "review this document", "szerkeszd", "dokumentum", ".docx", ".md")):
        return "document_review"
    return "general_analysis"


def _infer_target_audience(task_type: str, user_prompt: str) -> str:
    text = (user_prompt or "").lower()
    if any(k in text for k in ("investor", "befektet", "pitch", "board", "vezetĹ‘sĂ©g")):
        return "investors and leadership"
    if any(k in text for k in ("developer", "engineer", "architect", "fejleszt", "mĂ©rnĂ¶k")):
        return "engineering team"
    if task_type in ("technical_audit", "architecture_decision", "implementation_plan"):
        return "engineering and product stakeholders"
    if task_type in ("business_strategy", "product_roadmap"):
        return "founders, product, and go-to-market stakeholders"
    if task_type == "document_review":
        return "document owner and reviewers"
    return "decision makers"


def _infer_recommended_scenario(task_type: str, current_scenario: str) -> str:
    if task_type == "red_team_review":
        return "red-team"
    if task_type in ("technical_audit", "architecture_decision", "product_roadmap"):
        return "expert-council"
    if current_scenario in SCENARIOS:
        return current_scenario
    return "quick"


def _infer_constraints(user_prompt: str, items: list) -> list[str]:
    constraints = []
    text = (user_prompt or "").lower()
    if any(k in text for k in ("do not", "don't", "ne ", "tilos", "without", "nĂ©lkĂĽl")):
        constraints.append("Respect explicit user constraints in the prompt.")
    if any(k in text for k in ("sprint", "deadline", "hatĂˇridĹ‘", "today", "mvp")):
        constraints.append("Keep recommendations phased and delivery-oriented.")
    if items:
        constraints.append("Ground claims in the provided source material.")
    else:
        constraints.append("No source files were provided; rely on the user prompt and mark assumptions.")
    return constraints


def _infer_assumptions(task_type: str, items: list) -> list[str]:
    assumptions = [
        "The profile is a lightweight routing aid, not a strict output contract.",
        "The existing CLI scenario remains authoritative unless the user overrides it.",
    ]
    if task_type in ("business_strategy", "implementation_plan"):
        assumptions.append("Preserve the current business plan, implementation plan, and AI context behavior.")
    if not items:
        assumptions.append("Source evidence may be incomplete because no files were loaded.")
    return assumptions


def build_task_profile(user_prompt: str, items: list, lang: str, scenario_key: str) -> TaskProfile:
    try:
        task_type = _infer_task_type(user_prompt, items)
        if task_type not in TASK_TYPES:
            task_type = "general_analysis"
        language = _infer_language(user_prompt, lang)
        return TaskProfile(
            task_type=task_type,
            user_goal=(user_prompt or "").strip() or "Analyze the provided source material.",
            target_audience=_infer_target_audience(task_type, user_prompt),
            language=language,
            recommended_scenario=_infer_recommended_scenario(task_type, scenario_key),
            required_perspectives=list(TASK_TYPE_PERSPECTIVES.get(task_type, TASK_TYPE_PERSPECTIVES["general_analysis"])),
            expected_deliverables=list(TASK_TYPE_DELIVERABLES.get(task_type, ["final_synthesis"])),
            constraints=_infer_constraints(user_prompt, items),
            assumptions=_infer_assumptions(task_type, items),
            profile_source="deterministic",
        )
    except Exception:
        return TaskProfile(
            task_type="general_analysis",
            user_goal=(user_prompt or "").strip() or "Analyze the provided source material.",
            target_audience="decision makers",
            language=_infer_language(user_prompt, lang or "hu"),
            recommended_scenario=scenario_key if scenario_key in SCENARIOS else "quick",
            required_perspectives=["strategic", "technical", "critical"],
            expected_deliverables=["final_synthesis"],
            constraints=["Task profiling failed; proceed conservatively."],
            assumptions=["Deterministic fallback profile was used."],
            profile_source="deterministic_fallback",
        )


def task_profile_to_dict(profile: TaskProfile | dict | None) -> dict:
    if profile is None:
        return {}
    if isinstance(profile, TaskProfile):
        return asdict(profile)
    return dict(profile)


def format_task_profile(profile: TaskProfile | dict | None) -> str:
    data = task_profile_to_dict(profile)
    if not data:
        return "(no task profile)"
    return json.dumps(data, ensure_ascii=False, indent=2)


def task_profile_prompt_block(profile: TaskProfile | dict | None) -> str:
    data = task_profile_to_dict(profile)
    if not data:
        return "TASK PROFILE: not available"
    return f"""TASK PROFILE:
- task_type: {data.get('task_type')}
- user_goal: {data.get('user_goal')}
- target_audience: {data.get('target_audience')}
- language: {data.get('language')}
- recommended_scenario: {data.get('recommended_scenario')}
- required_perspectives: {', '.join(data.get('required_perspectives') or [])}
- expected_deliverables: {', '.join(data.get('expected_deliverables') or [])}
- constraints: {'; '.join(data.get('constraints') or [])}
- assumptions: {'; '.join(data.get('assumptions') or [])}"""


def print_task_profile(profile: TaskProfile | dict):
    print("\n" + "â•"*64)
    print(_bold("TaskProfile"))
    print("â•"*64)
    print(format_task_profile(profile))


def save_task_profile_only(profile: TaskProfile | dict, out_dir: str, user_prompt: str,
                           output_contract: OutputContract | dict | None = None):
    os.makedirs(out_dir, exist_ok=True)
    profile_dict = task_profile_to_dict(profile)
    Path(os.path.join(out_dir, "task_profile.json")).write_text(
        json.dumps(profile_dict, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    payload = {
        "user_prompt": user_prompt,
        "debate": [],
        "metadata": {
            "task_profile": profile_dict,
            "output_contract": output_contract_summary(output_contract),
        },
    }
    Path(os.path.join(out_dir, "debate_log.json")).write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"âś… TaskProfile JSON: {os.path.join(out_dir, 'task_profile.json')}")
    print(f"âś… Log: {os.path.join(out_dir, 'debate_log.json')}")


DEFAULT_OUTPUT_CONTRACT = {
    "contract_id": "default_business_master_plan",
    "title": "SzintĂ©zis",
    "language": "hu",
    "audience": "founders, product, engineering, and AI implementation stakeholders",
    "output_format": "markdown",
    "deliverables": [
        {
            "id": "final_verdict",
            "title": "VĂ©gsĹ‘ Ă­tĂ©let",
            "required": True,
            "source": "verdict",
            "sections": [],
        },
        {
            "id": "business_plan",
            "title": "VĂ©gsĹ‘ ĂĽzleti terv",
            "required": True,
            "source": "main",
            "sections": [
                {"id": "product_value", "title": "TermĂ©k Ă©s Ă©rtĂ©kajĂˇnlat", "required": False},
                {"id": "market_gtm", "title": "Piac Ă©s go-to-market", "required": False},
                {"id": "current_state", "title": "Jelenlegi Ăˇllapot Ă©s korlĂˇtok", "required": False},
                {"id": "business_model", "title": "BevĂ©teli / mĹ±kĂ¶dĂ©si modell", "required": False},
                {"id": "competitive_position", "title": "VersenypozĂ­ciĂł", "required": False},
                {"id": "business_risks", "title": "Ăśzleti kockĂˇzatok", "required": False},
                {"id": "business_recommendations", "title": "Ăśzleti ajĂˇnlĂˇsok", "required": False},
            ],
        },
        {
            "id": "implementation_plan",
            "title": "VĂ©gsĹ‘ megvalĂłsĂ­tĂˇsi terv",
            "required": True,
            "source": "implementation",
            "sections": [
                {"id": "architecture_direction", "title": "ArchitektĂşra irĂˇny", "required": False},
                {"id": "roadmap", "title": "Roadmap", "required": False},
                {"id": "technical_risks", "title": "Technikai kockĂˇzatok", "required": False},
            ],
        },
        {
            "id": "ai_context_block",
            "title": "AI kontextus blokk",
            "required": True,
            "source": "ai_context",
            "sections": [],
        },
    ],
    "required_metadata": [
        "consensus_points",
        "risk_register",
        "decision_log",
        "recommendations",
    ],
}


@dataclass
class OutputContract:
    contract_id: str
    title: str
    language: str
    audience: str
    deliverables: list[dict] = field(default_factory=list)
    required_metadata: list[str] = field(default_factory=list)
    output_format: str = "markdown"
    contract_source: str = "builtin_default"


def _normalize_contract_section(section) -> dict:
    if isinstance(section, str):
        return {"id": _slug(section), "title": section, "required": True}
    if not isinstance(section, dict):
        _abort("Invalid OutputContract: every section must be a string or object.")
    title = str(section.get("title") or section.get("id") or "").strip()
    if not title:
        _abort("Invalid OutputContract: section is missing title/id.")
    return {
        "id": str(section.get("id") or _slug(title)),
        "title": title,
        "required": bool(section.get("required", True)),
    }


def _normalize_contract_deliverable(deliverable) -> dict:
    if not isinstance(deliverable, dict):
        _abort("Invalid OutputContract: every deliverable must be an object.")
    title = str(deliverable.get("title") or deliverable.get("id") or "").strip()
    if not title:
        _abort("Invalid OutputContract: deliverable is missing title/id.")
    did = str(deliverable.get("id") or _slug(title))
    sections = [_normalize_contract_section(s) for s in deliverable.get("sections", [])]
    return {
        "id": did,
        "title": title,
        "required": bool(deliverable.get("required", True)),
        "source": str(deliverable.get("source") or _infer_contract_source(did, title)),
        "sections": sections,
    }


def _slug(text: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", str(text).lower()).strip("_")
    return slug or "section"


def _infer_contract_source(deliverable_id: str, title: str) -> str:
    blob = f"{deliverable_id} {title}".lower()
    if any(k in blob for k in ("verdict", "Ă­tĂ©let", "itelet", "decision")):
        return "verdict"
    if any(k in blob for k in ("context", "kontextus", "ai_context")):
        return "ai_context"
    if any(k in blob for k in ("implementation", "megvalĂłsĂ­tĂˇs", "megvalositas", "roadmap", "remediation", "mitigation")):
        return "implementation"
    return "main"


def output_contract_from_dict(data: dict, source: str = "provided") -> OutputContract:
    if not isinstance(data, dict):
        _abort("Invalid OutputContract: root must be a JSON object.")
    contract_id = str(data.get("contract_id") or "").strip()
    title = str(data.get("title") or "").strip()
    if not contract_id:
        _abort("Invalid OutputContract: missing contract_id.")
    if not title:
        _abort("Invalid OutputContract: missing title.")
    deliverables = data.get("deliverables")
    if not isinstance(deliverables, list) or not deliverables:
        _abort("Invalid OutputContract: deliverables must be a non-empty list.")
    required_metadata = data.get("required_metadata", [])
    if not isinstance(required_metadata, list):
        _abort("Invalid OutputContract: required_metadata must be a list.")
    output_format = str(data.get("output_format") or "markdown").strip()
    if output_format not in ("markdown", "json", "docx-compatible"):
        _abort("Invalid OutputContract: output_format must be markdown/json/docx-compatible.")
    return OutputContract(
        contract_id=contract_id,
        title=title,
        language=str(data.get("language") or "hu"),
        audience=str(data.get("audience") or "decision makers"),
        deliverables=[_normalize_contract_deliverable(d) for d in deliverables],
        required_metadata=[str(m) for m in required_metadata],
        output_format=output_format,
        contract_source=source,
    )


def default_output_contract(language: str = "hu") -> OutputContract:
    data = json.loads(json.dumps(DEFAULT_OUTPUT_CONTRACT, ensure_ascii=False))
    data["language"] = language or data.get("language", "hu")
    return output_contract_from_dict(data, source="builtin_default")


def load_output_contract(path: str, language: str = "hu") -> OutputContract:
    if not path:
        return default_output_contract(language)
    p = Path(path)
    if not p.exists():
        _abort(f"Invalid OutputContract: file not found: {path}")
    try:
        raw = p.read_text(encoding="utf-8-sig")
        if p.suffix.lower() in (".yaml", ".yml"):
            try:
                import yaml
            except ImportError:
                _abort("Invalid OutputContract: YAML contract requires PyYAML; use JSON or install pyyaml.")
            data = yaml.safe_load(raw)
        else:
            data = json.loads(raw)
    except SystemExit:
        raise
    except Exception as e:
        _abort(f"Invalid OutputContract: could not parse {path}: {type(e).__name__}: {e}")
    contract = output_contract_from_dict(data, source=str(p))
    if not contract.language:
        contract.language = language
    return contract


def output_contract_to_dict(contract: OutputContract | dict | None) -> dict:
    if contract is None:
        return {}
    if isinstance(contract, OutputContract):
        return asdict(contract)
    return dict(contract)


def output_contract_summary(contract: OutputContract | dict | None) -> dict:
    data = output_contract_to_dict(contract)
    if not data:
        return {}
    return {
        "contract_id": data.get("contract_id"),
        "title": data.get("title"),
        "language": data.get("language"),
        "audience": data.get("audience"),
        "output_format": data.get("output_format"),
        "deliverables": [
            {
                "id": d.get("id"),
                "title": d.get("title"),
                "sections": [s.get("title") for s in d.get("sections", [])],
            }
            for d in data.get("deliverables", [])
        ],
        "required_metadata": data.get("required_metadata", []),
        "contract_source": data.get("contract_source", ""),
    }


def output_contract_prompt_block(contract: OutputContract | dict | None) -> str:
    summary = output_contract_summary(contract)
    if not summary:
        return "OUTPUT CONTRACT: not available"
    return "OUTPUT CONTRACT:\n" + json.dumps(summary, ensure_ascii=False, indent=2)


# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# FĂˇjlkezelĂ©s (vĂˇltozatlan v4-bĹ‘l)
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
TEXT_EXT  = {".txt",".md",".markdown",".rst",".json",".yaml",".yml",".xml",
             ".html",".htm",".css",".js",".ts",".tsx",".jsx",".py",".java",
             ".cs",".cpp",".c",".h",".go",".rs",".php",".rb",".sh",".bat",
             ".ps1",".sql",".toml",".ini",".cfg",".env",".gitignore",".kt",".swift"}
WORD_EXT  = {".docx",".doc"}
EXCEL_EXT = {".xlsx",".xls",".xlsm",".csv",".tsv"}
PDF_EXT   = {".pdf"}
IMG_EXT   = {".png",".jpg",".jpeg",".gif",".webp",".bmp",".tiff"}
ZIP_EXT   = {".zip",".tar",".gz",".tgz",".7z"}
EXCL_DIRS = {"node_modules",".git","dist","build",".next","__pycache__","vendor","venv",".venv"}
EXCL_FILES= {"package-lock.json","yarn.lock","bun.lockb","pnpm-lock.yaml"}
EXCL_PATS = ["components/ui/","components\\ui\\",".min.js",".min.css"]

def _excluded(path: Path, base: Path) -> bool:
    for p in path.parts:
        if p in EXCL_DIRS or p.startswith("~") or p.startswith("."): return True
    if path.name in EXCL_FILES: return True
    rel = str(path.relative_to(base)).replace("\\","/")
    return any(x.replace("\\","/") in rel for x in EXCL_PATS)

def _trunc(text: str, n: int) -> str:
    if len(text)<=n: return text
    h=n//2; return text[:h]+f"\n[...{len(text)-n} kihagyva...]\n"+text[-h:]

def _is_failed_response(text: str) -> bool:
    if not text or not str(text).strip():
        return True
    t = str(text).strip().lower()
    if t.startswith("[error"):
        return True
    failure_bits = (
        " hiba:", "hiba:", "api kulcs", "missing api key",
        "not available", "unavailable", "nem el", "nem el",
        "invalid api key", "authentication", "permission denied",
        "rate limit", "quota", "timeout", "timed out",
    )
    return t.startswith("[") and any(bit in t for bit in failure_bits)

def _valid_response(text: str) -> bool:
    return not _is_failed_response(text)

NON_FATAL_FAILED_STEP_REPAIRS = {
    "final_decision_outline_markdown_fallback": "final_decision_outline",
}

def _normalize_human_failed_steps(failed_steps: list[str],
                                  repaired_steps: list[str] | None = None) -> list[str]:
    repaired = set(repaired_steps or [])
    normalized = []
    for step in failed_steps or []:
        repair_step = NON_FATAL_FAILED_STEP_REPAIRS.get(step)
        if step in repaired or (repair_step and repair_step in repaired):
            continue
        normalized.append(step)
    return normalized

def _effective_human_failed_steps(step_results: list[dict],
                                  repaired_steps: list[str] | None,
                                  metadata_step_names: set[str]) -> list[str]:
    raw_failed = [
        s["step"] for s in step_results
        if not s.get("ok") and s.get("step") not in metadata_step_names
    ]
    return _normalize_human_failed_steps(raw_failed, repaired_steps)

def _exception_details(e: Exception) -> dict:
    status = getattr(e, "status_code", None) or getattr(e, "status", None)
    body = getattr(e, "body", None) or getattr(e, "response", None)
    if body is not None and not isinstance(body, str):
        body = str(body)
    msg = str(e) or "(empty exception message)"
    detail = {
        "exception_class": type(e).__name__,
        "exception_message": msg,
    }
    if status is not None:
        detail["status"] = status
    if body:
        detail["raw_api_error"] = body[:1000]
    return detail

def _health_detail_text(info: dict) -> str:
    parts = [
        f"provider={info.get('provider','?')}",
        f"model_key={info.get('model_key','?')}",
        f"model_id={info.get('model_id','?')}",
    ]
    if info.get("exception_class"):
        parts.append(f"exception={info.get('exception_class')}")
    if info.get("status") is not None:
        parts.append(f"status={info.get('status')}")
    msg = info.get("exception_message") or info.get("detail") or ""
    if msg:
        parts.append(f"message={msg}")
    if info.get("raw_api_error"):
        parts.append(f"raw={info['raw_api_error']}")
    return "; ".join(parts)

def _abort(message: str):
    print(f"\nERROR: {message}")
    raise SystemExit(2)

def _participant_roles(engines: dict) -> list[str]:
    return [k for k in engines if k not in INFRASTRUCTURE_ROLES]

def _required_participant_count(scenario_key: str) -> int:
    if scenario_key == "quick":
        return 2
    return MIN_VALID_PARTICIPANTS

def _valid_participant_roles(engines: dict, health: Optional[dict] = None) -> list[str]:
    roles = _participant_roles(engines)
    if health is None:
        return roles
    return [r for r in roles if health.get(r, {}).get("ok")]

def _require_min_participants(roles: list[str], context: str, required: int = None):
    required = required or MIN_VALID_PARTICIPANTS
    if len(roles) < required:
        _abort(
            f"{context}: only {len(roles)} valid debate participant(s), "
            f"need at least {required}. No final document was generated."
        )

def _role_mapping(engines: dict) -> dict:
    return {
        role: {
            "model_key": eng.key,
            "label": eng.label,
            "provider": eng.provider,
            "model_id": eng.mid,
            "role_type": "infrastructure" if role in INFRASTRUCTURE_ROLES else "debate_participant",
        }
        for role, eng in engines.items()
    }

class DocItem:
    def __init__(self,name,path,text="",b64="",mime="",kind="text"):
        self.name=name; self.path=path; self.text=text
        self.b64=b64; self.mime=mime; self.kind=kind
    def is_img(self): return self.kind=="image"

def _read_txt(p):
    for enc in ("utf-8","utf-8-sig","cp1250","latin-1"):
        try: return Path(p).read_text(encoding=enc)
        except UnicodeDecodeError: continue
    return Path(p).read_bytes().decode("utf-8",errors="replace")

def _read_docx(p):
    try:
        doc=DocxReader(p)
        parts=[par.text for par in doc.paragraphs if par.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                r=" | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if r: parts.append(r)
        return "\n\n".join(parts)
    except Exception as e: return f"[DOCX hiba: {e}]"

def _read_pdf(p):
    try:
        import fitz
        return "\n\n".join(f"[{i+1}]\n{pg.get_text()}" for i,pg in enumerate(fitz.open(p)))
    except ImportError: pass
    try:
        import pypdf
        r=pypdf.PdfReader(p)
        return "\n\n".join(f"[{i+1}]\n{pg.extract_text()}" for i,pg in enumerate(r.pages))
    except ImportError: return "[PDF: pip install pymupdf]"

def _read_excel(p):
    try:
        import pandas as pd
        ext=Path(p).suffix.lower()
        if ext in (".csv",".tsv"): return pd.read_csv(p,encoding_errors="replace").to_string(index=False,max_rows=100)
        dfs=pd.read_excel(p,sheet_name=None)
        return "\n\n".join(f"--- {s} ---\n{df.to_string(index=False,max_rows=50)}" for s,df in dfs.items())
    except ImportError: pass
    try:
        import openpyxl
        wb=openpyxl.load_workbook(p,read_only=True,data_only=True)
        parts=[]
        for s in wb.sheetnames:
            rows=[" | ".join(str(c) if c else "" for c in row) for row in wb[s].iter_rows(max_row=100,values_only=True)]
            parts.append(f"--- {s} ---\n"+"\n".join(r for r in rows if r.strip()))
        return "\n\n".join(parts)
    except: return "[Excel: pip install openpyxl]"

def _proc_file(path,max_chars,name=None):
    p=Path(path); ext=p.suffix.lower(); n=name or p.name
    if ext in IMG_EXT:
        try:
            mime=mimetypes.guess_type(path)[0] or "image/png"
            b64=base64.standard_b64encode(p.read_bytes()).decode()
            return DocItem(n,path,b64=b64,mime=mime,kind="image")
        except Exception as e: return DocItem(n,path,text=f"[kĂ©p hiba: {e}]")
    if ext in WORD_EXT:  text=_read_docx(path)
    elif ext in PDF_EXT: text=_read_pdf(path)
    elif ext in EXCEL_EXT: text=_read_excel(path)
    elif ext in TEXT_EXT:  text=_read_txt(path)
    else:
        try:
            text=_read_txt(path)
            if "\x00" in text[:200]: return DocItem(n,path,text="[binĂˇris]",kind="skip")
        except: return DocItem(n,path,text="[nem olvashatĂł]",kind="skip")
    return DocItem(n,path,text=_trunc(text,max_chars))

def _extract_zip(zp,tmp):
    out=os.path.join(tmp,Path(zp).stem); os.makedirs(out,exist_ok=True)
    try:
        if zp.endswith(".zip"):
            with zipfile.ZipFile(zp) as z: z.extractall(out)
        else:
            import tarfile
            with tarfile.open(zp,"r:*") as t: t.extractall(out)
    except Exception as e: print(f"   âš ď¸Ź  ZIP: {e}")
    return out

def load_sources(sources, max_chars):
    all_files=[]; tmp_dirs=[]
    for src in sources:
        sp=Path(src)
        if not sp.exists(): print(f"   âš ď¸Ź  Nem talĂˇlom: {src}"); continue
        if sp.is_dir():
            for f in sorted(sp.rglob("*")):
                if f.is_file() and not _excluded(f,sp):
                    all_files.append((str(f),str(f.relative_to(sp))))
        elif sp.suffix.lower() in ZIP_EXT:
            print(f"   đź—śď¸Ź  ZIP: {sp.name}")
            tmp=tempfile.mkdtemp(prefix="debate_"); tmp_dirs.append(tmp)
            ep=Path(_extract_zip(str(sp),tmp))
            for f in sorted(ep.rglob("*")):
                if f.is_file() and not _excluded(f,ep):
                    all_files.append((str(f),f"{sp.stem}/{f.relative_to(ep)}"))
        else:
            all_files.append((str(sp),sp.name))
    print(f"\nđź“‚ {len(all_files)} fĂˇjl...")
    items=[]
    for fp,disp in all_files:
        item=_proc_file(fp,max_chars,disp)
        if item.kind=="skip": continue
        items.append(item)
        icon="đź–Ľď¸Ź " if item.is_img() else "âś…"
        size=f"({len(item.text):,} kar)" if not item.is_img() else ""
        print(f"   {icon} {Path(fp).suffix.upper() or '?':6s} {disp} {size}")
    for d in tmp_dirs: shutil.rmtree(d,ignore_errors=True)
    if not items: print("âťŚ Nincs feldolgozhatĂł fĂˇjl."); sys.exit(1)
    imgs=sum(1 for i in items if i.is_img())
    print(f"\n   đź“Š {len(items)-imgs} szĂ¶veges + {imgs} kĂ©p")
    return items


# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# AI Engine â€” egysĂ©ges hĂ­vĂł, streaming + fallback
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
class AIEngine:
    """EgysĂ©ges interfĂ©sz minden modellhez."""
    def __init__(self, model_key: str, quality: str):
        cfg = dict(CATALOG[model_key])
        cfg["_mid"] = cfg["id_fast"] if quality == "fast" else cfg["id"]
        self.key   = model_key
        self.cfg   = cfg
        self.client = None
        self._init_client()

    def _init_client(self):
        t = self.cfg["type"]
        if t == "anthropic":
            k = os.environ.get("ANTHROPIC_API_KEY")
            if k: self.client = anthropic.Anthropic(api_key=k)
        elif t in ("openai", "openai_responses"):
            k = os.environ.get("OPENAI_API_KEY")
            if k: self.client = OpenAI(api_key=k)
        elif t == "openai_compat":
            k = os.environ.get(self.cfg.get("env_key",""))
            if k: self.client = OpenAI(api_key=k, base_url=self.cfg["base_url"])
        elif t == "gemini":
            k = os.environ.get("GOOGLE_API_KEY")
            if k:
                try:
                    from google import genai
                    self.client = genai.Client(api_key=k)
                except ImportError: pass

    @property
    def available(self): return self.client is not None
    @property
    def label(self): return self.cfg["label"]
    @property
    def emoji(self): return self.cfg["emoji"]
    @property
    def color(self): return self.cfg["color"]
    @property
    def mid(self): return self.cfg["_mid"]

    @property
    def provider(self): return self.cfg["type"]

    @property
    def required_env_key(self):
        t = self.cfg["type"]
        if t == "anthropic": return "ANTHROPIC_API_KEY"
        if t in ("openai", "openai_responses"): return "OPENAI_API_KEY"
        if t == "openai_compat": return self.cfg.get("env_key", "")
        if t == "gemini": return "GOOGLE_API_KEY"
        return ""

    def call(self, system: str, messages: list, stream: bool = True,
             label: str = None, max_output_tokens: int = None) -> str:
        """HĂ­vja a modellt, opcionĂˇlisan streamelve. Visszaadja a teljes szĂ¶veget."""
        if not self.available:
            key = self.required_env_key or "API key"
            return f"[ERROR {self.label}: missing {key}; provider={self.provider}; model={self.mid}]"
        lbl = label or self.label
        t = self.cfg["type"]
        text = ""
        max_out = int(max_output_tokens or MAX_TOKENS)

        if stream:
            sys.stdout.write(_c(f"\n{self.emoji} {lbl}:\n", self.color))
            sys.stdout.flush()

        try:
            if t == "anthropic":
                if stream:
                    with self.client.messages.stream(
                        model=self.mid, max_tokens=max_out,
                        system=system, messages=messages
                    ) as s:
                        for chunk in s.text_stream:
                            sys.stdout.write(_c(chunk, self.color))
                            sys.stdout.flush()
                            text += chunk
                else:
                    r = self.client.messages.create(
                        model=self.mid, max_tokens=max_out,
                        system=system, messages=messages)
                    text = r.content[0].text

            elif t == "openai_responses":
                # OpenAI Responses API (GPT-5.x reasoning modellek)
                reasoning_effort = self.cfg.get("reasoning_effort", "medium")
                # Responses API input: system + messages Ă¶sszefĹ±zve
                input_msgs = [{"role": "system", "content": system}] + messages
                if stream:
                    with self.client.responses.stream(
                        model=self.mid,
                        input=input_msgs,
                        reasoning={"effort": reasoning_effort},
                        max_output_tokens=max_out,
                    ) as s:
                        for event in s:
                            if hasattr(event, "delta") and event.delta:
                                sys.stdout.write(_c(event.delta, self.color))
                                sys.stdout.flush()
                                text += event.delta
                else:
                    r = self.client.responses.create(
                        model=self.mid,
                        input=input_msgs,
                        reasoning={"effort": reasoning_effort},
                        max_output_tokens=max_out,
                    )
                    text = r.output_text or ""

            elif t in ("openai","openai_compat"):
                msgs = [{"role":"system","content":system}] + messages
                if stream:
                    for chunk in self.client.chat.completions.create(
                        model=self.mid, max_tokens=max_out,
                        stream=True, messages=msgs
                    ):
                        d = chunk.choices[0].delta.content or ""
                        if d:
                            sys.stdout.write(_c(d, self.color))
                            sys.stdout.flush()
                            text += d
                else:
                    r = self.client.chat.completions.create(
                        model=self.mid, max_tokens=max_out, messages=msgs)
                    text = r.choices[0].message.content

            elif t == "gemini":
                from google import genai
                from google.genai import types as gt
                history = []
                for m in messages[:-1]:
                    role = "user" if m["role"]=="user" else "model"
                    history.append(gt.Content(role=role,
                                              parts=[gt.Part(text=m["content"])]))
                last_msg = messages[-1]["content"] if messages else ""
                chat = self.client.chats.create(
                    model=self.mid,
                    config=gt.GenerateContentConfig(
                        system_instruction=system,
                        max_output_tokens=max_out),
                    history=history)
                if stream:
                    for chunk in chat.send_message_stream(last_msg):
                        d = chunk.text or ""
                        if d:
                            sys.stdout.write(_c(d, self.color))
                            sys.stdout.flush()
                            text += d
                else:
                    text = chat.send_message(last_msg).text or ""

        except Exception as e:
            detail = _exception_details(e)
            err = (
                f"[ERROR {self.label}: provider={self.provider}; model_key={self.key}; "
                f"model_id={self.mid}; exception={detail['exception_class']}; "
                f"message={detail['exception_message']}"
                f"{'; status=' + str(detail['status']) if detail.get('status') is not None else ''}"
                f"{'; raw=' + detail['raw_api_error'] if detail.get('raw_api_error') else ''}]"
            )
            if stream: sys.stdout.write(_c(err, self.color))
            text = err

        if stream:
            sys.stdout.write("\n"); sys.stdout.flush()
        return text

    def health_check(self) -> dict:
        base = {
            "provider": self.provider,
            "model_key": self.key,
            "model_id": self.mid,
            "label": self.label,
        }
        if not self.available:
            key = self.required_env_key or "API key"
            return {
                **base,
                "ok": False,
                "detail": f"missing {key}",
                "exception_class": "MissingApiKey",
                "exception_message": f"missing {key}",
            }
        try:
            resp = self._health_probe()
            if _is_failed_response(resp):
                return {
                    **base,
                    "ok": False,
                    "detail": resp,
                    "exception_class": "HealthCheckFailed",
                    "exception_message": resp or "(empty provider response)",
                }
            return {**base, "ok": True, "detail": "ok"}
        except Exception as e:
            return {**base, "ok": False, "detail": str(e) or "(empty exception message)", **_exception_details(e)}

    def _health_probe(self) -> str:
        system = "Health check. Reply with OK only."
        prompt = "OK?"
        t = self.cfg["type"]
        if t == "anthropic":
            r = self.client.messages.create(
                model=self.mid,
                max_tokens=16,
                system=system,
                messages=[{"role": "user", "content": prompt}],
            )
            return r.content[0].text if r.content else ""
        if t == "openai_responses":
            r = self.client.responses.create(
                model=self.mid,
                input=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt},
                ],
                reasoning={"effort": self.cfg.get("reasoning_effort", "medium")},
                max_output_tokens=16,
            )
            return r.output_text or ""
        if t in ("openai", "openai_compat"):
            r = self.client.chat.completions.create(
                model=self.mid,
                max_tokens=16,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt},
                ],
            )
            return r.choices[0].message.content or ""
        if t == "gemini":
            from google.genai import types as gt
            chat = self.client.chats.create(
                model=self.mid,
                config=gt.GenerateContentConfig(
                    system_instruction=system,
                    max_output_tokens=16,
                ),
            )
            return chat.send_message(prompt).text or ""
        return ""


# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Szerepek Ă¶sszeĂˇllĂ­tĂˇsa
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def build_roles(scenario_key: str, quality: str, role_overrides: dict) -> dict[str, AIEngine]:
    """
    Visszaad: {role_name: AIEngine}
    role_overrides: {"judge": "gpt", "skeptic": "deepseek", ...}
    """
    scenario = SCENARIOS[scenario_key]
    base = dict(scenario["quality_map"][quality])
    base.update(role_overrides)  # felĂĽlĂ­rĂˇsok Ă©rvĂ©nyesĂ­tĂ©se

    engines = {}
    print(f"\nđź”Ś Szerepek Ă©s modellek ({scenario_key} | {quality}):")
    for role, model_key in base.items():
        if role == "debaters":
            # quick scenario: debaters lista
            for i, mk in enumerate(model_key):
                mk_resolved = role_overrides.get(f"debater{i+1}", mk)
                if mk_resolved not in CATALOG:
                    print(f"   âš ď¸Ź  Ismeretlen modell: {mk_resolved}, visszatĂ©rĂ©s: {mk}")
                    mk_resolved = mk
                eng = AIEngine(mk_resolved, quality)
                role_name = f"debater{i+1}"
                engines[role_name] = eng
                status = "âś…" if eng.available else "âťŚ"
                print(f"   {status} {role_name:12s} â†’ {eng.label} ({eng.mid})")
        else:
            mk_resolved = role_overrides.get(role, model_key)
            if mk_resolved not in CATALOG:
                print(f"   âš ď¸Ź  Ismeretlen modell override: {mk_resolved}, visszatĂ©rĂ©s: {model_key}")
                mk_resolved = model_key
            eng = AIEngine(mk_resolved, quality)
            engines[role] = eng
            status = "âś…" if eng.available else "âťŚ"
            print(f"   {status} {role:12s} â†’ {eng.label} ({eng.mid})")

    return engines

def _check_engine_health(eng: AIEngine, skip_network: bool = False) -> dict:
    if skip_network:
        ok = eng.available
        detail = "client configured" if ok else f"missing {eng.required_env_key or 'API key'}"
        synthesis = {
            "ok": ok,
            "detail": detail,
            "provider": eng.provider,
            "model_key": eng.key,
            "model_id": eng.mid,
            "label": eng.label,
            "exception_class": None if ok else "MissingApiKey",
            "exception_message": "" if ok else detail,
        }
    return eng.health_check()

def run_health_checks(engines: dict, skip_network: bool = False) -> dict:
    print("\nProvider/model health check:")
    health = {}
    for role, eng in engines.items():
        info = _check_engine_health(eng, skip_network=skip_network)
        health[role] = info
        status = "OK" if info.get("ok") else "FAIL"
        print(f"   {status:4s} {role:14s} -> {eng.label} ({eng.mid}) - {_health_detail_text(info)}")
    return health

def _fallback_candidates_for(role: str, eng: AIEngine) -> list[str]:
    if role in ("judge", "moderator"):
        base = ROLE_FALLBACKS[role]
    else:
        base = ROLE_FALLBACKS.get(eng.provider, ["claude-sonnet", "gpt", "gemini", "deepseek", "grok"])
    return [mk for mk in base if mk in CATALOG and mk != eng.key]

def _existing_model_ids(engines: dict, exclude_role: str) -> set[str]:
    return {
        eng.mid for role, eng in engines.items()
        if role != exclude_role and role not in INFRASTRUCTURE_ROLES
    }

def _catalog_model_id(model_key: str, quality: str) -> str:
    cfg = CATALOG[model_key]
    return cfg["id_fast"] if quality == "fast" else cfg["id"]

def _ordered_candidates(role: str, eng: AIEngine, engines: dict, quality: str) -> list[str]:
    candidates = _fallback_candidates_for(role, eng)
    if role in INFRASTRUCTURE_ROLES:
        return candidates
    used = _existing_model_ids(engines, exclude_role=role)
    fresh = [mk for mk in candidates if _catalog_model_id(mk, quality) not in used]
    reused = [mk for mk in candidates if mk not in fresh]
    return fresh + reused

def print_resolved_role_mapping(engines: dict, original_mapping: dict):
    print("\nResolved role mapping:")
    for role, eng in engines.items():
        original = original_mapping.get(role, {})
        changed = original.get("model_id") and original.get("model_id") != eng.mid
        suffix = ""
        if changed:
            suffix = f" (was {original.get('label')} / {original.get('model_id')})"
        role_type = "infrastructure" if role in INFRASTRUCTURE_ROLES else "debate participant"
        print(f"   {role:14s} [{role_type:20s}] -> {eng.label} ({eng.mid}) [{eng.provider}]{suffix}")

def _abort_fallback_failure(prefix: str, engines: dict, health: dict,
                            resolution: dict, scenario_key: str):
    failed = []
    for role, info in health.items():
        if not info.get("ok"):
            failed.append(f"{role}: {_health_detail_text(info)}")

    tried = []
    for role, attempts in resolution.get("tried", {}).items():
        if not attempts:
            tried.append(f"{role}: no fallback candidates")
            continue
        labels = []
        for attempt in attempts:
            status = "ok" if attempt.get("ok") else "fail"
            labels.append(
                f"{attempt.get('model_key')}({status}: "
                f"{attempt.get('exception_message') or attempt.get('detail') or 'no detail'})"
            )
        tried.append(f"{role}: " + ", ".join(labels))

    key_issues = []
    for role, info in health.items():
        msg = (info.get("exception_message") or info.get("detail") or "").lower()
        if not info.get("ok") and ("missing" in msg or "api_key" in msg or "api key" in msg):
            key_issues.append(f"{role}: {info.get('exception_message') or info.get('detail')}")
        elif not info.get("ok") and ("auth" in msg or "invalid" in msg or "permission" in msg):
            key_issues.append(f"{role}: {info.get('exception_message') or info.get('detail')}")

    role_hint = "debater2=claude-sonnet,judge=gpt"
    if scenario_key != "quick":
        role_hint = "strategist=gpt,engineer=claude-sonnet,judge=gpt"
    message = [
        prefix,
        "Failed roles: " + ("; ".join(failed) if failed else "none after fallback"),
        "Fallbacks tried: " + ("; ".join(tried) if tried else "none"),
        "API key issues: " + ("; ".join(key_issues) if key_issues else "none detected; check provider status/model id/quota"),
        f"Manual override example: python ai_debate.py --scenario {scenario_key} --roles {role_hint} ...",
    ]
    _abort("\n".join(message))

def resolve_role_fallbacks(engines: dict, health: dict, scenario_key: str,
                           quality: str, skip_network: bool = False) -> dict:
    resolution = {
        "events": [],
        "failed_roles": {},
        "tried": {},
    }

    for role in list(engines.keys()):
        if health.get(role, {}).get("ok"):
            continue
        original = engines[role]
        candidates = _ordered_candidates(role, original, engines, quality)
        resolution["tried"][role] = candidates
        print(_yellow(
            f"   {role} {original.label} failed health check; trying fallbacks: "
            f"{', '.join(candidates) if candidates else '(none)'}"
        ))

        replacement = None
        replacement_health = None
        tried_details = []
        for model_key in candidates:
            cand = AIEngine(model_key, quality)
            info = _check_engine_health(cand, skip_network=skip_network)
            tried_details.append({
                "model_key": model_key,
                "label": cand.label,
                "model_id": cand.mid,
                "provider": cand.provider,
                "ok": info.get("ok"),
                "detail": info.get("detail"),
                "exception_class": info.get("exception_class"),
                "exception_message": info.get("exception_message"),
                "status": info.get("status"),
                "raw_api_error": info.get("raw_api_error"),
            })
            status = "OK" if info.get("ok") else "FAIL"
            print(f"      {status:4s} fallback {model_key:14s} -> {_health_detail_text(info)}")
            if info.get("ok"):
                replacement = cand
                replacement_health = info
                break

        resolution["tried"][role] = tried_details
        if replacement:
            print(_yellow(
                f"   {role} {original.label} failed health check, "
                f"replaced with {replacement.label}."
            ))
            resolution["events"].append({
                "role": role,
                "original": {
                    "model_key": original.key,
                    "label": original.label,
                    "provider": original.provider,
                    "model_id": original.mid,
                    "health": health.get(role, {}),
                },
                "resolved": {
                    "model_key": replacement.key,
                    "label": replacement.label,
                    "provider": replacement.provider,
                    "model_id": replacement.mid,
                    "health": replacement_health,
                },
            })
            engines[role] = replacement
            health[role] = replacement_health
        else:
            resolution["failed_roles"][role] = {
                "original": {
                    "model_key": original.key,
                    "label": original.label,
                    "provider": original.provider,
                    "model_id": original.mid,
                    "health": health.get(role, {}),
                },
                "tried": tried_details,
            }

    required = _required_participant_count(scenario_key)
    valid_participants = _valid_participant_roles(engines, health)
    if "moderator" not in engines or not health.get("moderator", {}).get("ok"):
        _abort_fallback_failure("Moderator infrastructure role is not healthy.", engines, health, resolution, scenario_key)
    if "judge" not in engines or not health.get("judge", {}).get("ok"):
        _abort_fallback_failure("Judge infrastructure role is not healthy.", engines, health, resolution, scenario_key)
    if len(valid_participants) < required:
        _abort_fallback_failure(
            f"Provider/model health check: only {len(valid_participants)} valid debate participant(s), need at least {required}.",
            engines, health, resolution, scenario_key,
        )
    return resolution

def drop_unhealthy_participants(engines: dict, health: dict):
    for role in list(_participant_roles(engines)):
        if not health.get(role, {}).get("ok"):
            eng = engines.pop(role)
            print(_yellow(
                f"   Excluding unhealthy participant before debate: "
                f"{role} -> {eng.label} ({health.get(role, {}).get('detail', 'failed')})"
            ))


# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# FĂZISOK
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def phase_evidence(items: list, user_prompt: str, moderator: AIEngine,
                   task_profile: TaskProfile | dict | None = None) -> str:
    """
    Phase 0: Evidence Pack
    A moderĂˇtor tĂ¶mĂ¶rĂ­ti a forrĂˇsanyagot egy strukturĂˇlt csomagba.
    Ez kerĂĽl a vitĂˇzĂłkhoz, nem a nyers fĂˇjlok.
    """
    print("\n" + "â•"*64)
    print(_bold("âšˇ PHASE 0 â€” Evidence Pack generĂˇlĂˇsa"))
    print("â•"*64)

    txt_items = [i for i in items if not i.is_img()]
    raw_text  = "\n\n".join(
        f"[{i.name}]\n{i.text}" for i in txt_items
    )
    source_list = "\n".join(f"- {i.name}" for i in items)

    system = """Te egy precĂ­z dokumentum-Ă¶sszefoglalĂł AI vagy.
Feladatod: egy strukturĂˇlt Evidence Package elkĂ©szĂ­tĂ©se
amelyet AI modellek kapnak vitaalapkĂ©nt.
KIZĂRĂ“LAG a megadott forrĂˇsokbĂłl dolgozz, ne talĂˇlj ki semmit."""

    prompt = f"""{task_profile_prompt_block(task_profile)}

FELHASZNĂLĂ“I CĂ‰L:
{user_prompt}

FORRĂSDOKUMENTUMOK:
{source_list}

TARTALOM:
{_trunc(raw_text, 15000)}

KĂ©szĂ­tsd el az Evidence Package-t ebben a formĂˇtumban:

## USER_GOAL
[1-2 mondatos Ă¶sszefoglalĂˇs]

## SOURCE_FILES
[fĂˇjlok listĂˇja tĂ­pussal]

## KEY_FACTS
[10-15 legfontosabb tĂ©ny, szĂˇmadat, ĂˇllĂ­tĂˇs a forrĂˇsokbĂłl]

## CONSTRAINTS
[technikai, ĂĽzleti, idĹ‘beli korlĂˇtok]

## OPEN_QUESTIONS
[megvĂˇlaszolatlan kĂ©rdĂ©sek, hiĂˇnyzĂł informĂˇciĂłk]

## EVIDENCE_SNIPPETS
[3-5 kulcsfontossĂˇgĂş idĂ©zet/rĂ©szlet a forrĂˇsokbĂłl, kontextussal]

## TASK_FOR_MODELS
[mit vĂˇrunk a vitĂˇzĂł modellek vĂˇlaszĂˇtĂłl]"""

    resp = moderator.call(system, [{"role":"user","content":prompt}],
                          stream=True, label="ModerĂˇtor â†’ Evidence Pack")
    return resp


def phase_independent(engines: dict, evidence: str, roles_info: dict,
                       parallel: bool, session_log: list,
                       task_profile: TaskProfile | dict | None = None) -> dict[str, str]:
    """
    Phase 1: FĂĽggetlen vĂ©lemĂ©nyek
    Minden modell KĂśLĂ–N kapja az evidence packot, nem lĂˇtja a tĂ¶bbieket.
    """
    print("\n" + "â•"*64)
    print(_bold("đź”Ť PHASE 1 â€” FĂĽggetlen ĂˇllĂˇspont (modellek NEM lĂˇtjĂˇk egymĂˇst)"))
    print("â•"*64)

    debaters = {k: v for k, v in engines.items()
                if k not in ("moderator", "judge")}

    def _call_one(role_key: str, engine: AIEngine) -> tuple[str, str]:
        role_desc = roles_info.get(role_key, role_key)
        system = f"""Te {engine.label} vagy, {role_desc} szerepben.
Az alĂˇbbi Evidence Package alapjĂˇn adj Ă¶nĂˇllĂł, fĂĽggetlen szakvĂ©lemĂ©nyt.
NE prĂłbĂˇld kitalĂˇlni mit mondanak a tĂ¶bbi AI modellek.
Adj strukturĂˇlt vĂˇlaszt: erĹ‘ssĂ©gek, gyengesĂ©gek, kockĂˇzatok, javaslatok.
Max 500 szĂł. TĂ¶mĂ¶r, szakmai hangon."""

        prompt = f"""{task_profile_prompt_block(task_profile)}

EVIDENCE PACKAGE:
{evidence}

Add meg a sajĂˇt fĂĽggetlen szakvĂ©lemĂ©nyedet {role_desc} szemszĂ¶gĂ©bĹ‘l!
LĂ©gy kritikus Ă©s konkrĂ©t. Emelj ki 2-3 legfontosabb pontot."""

        resp = engine.call(system, [{"role":"user","content":prompt}],
                           stream=not parallel,
                           label=f"{engine.label} [{role_desc}]")
        return role_key, resp

    opinions = {}
    if parallel and len(debaters) > 1:
        print(_yellow("   âšˇ PĂˇrhuzamos hĂ­vĂˇsok..."))
        with concurrent.futures.ThreadPoolExecutor(max_workers=len(debaters)) as ex:
            futures = {ex.submit(_call_one, k, v): k for k, v in debaters.items()}
            for fut in concurrent.futures.as_completed(futures):
                try:
                    rk, resp = fut.result()
                    opinions[rk] = resp
                    # Parallel esetĂ©n kiĂ­rjuk utĂłlag
                    eng = debaters[rk]
                    role_d = roles_info.get(rk, rk)
                    sys.stdout.write(_c(f"\n{eng.emoji} {eng.label} [{role_d}] (pĂˇrhuzamos):\n{resp}\n", eng.color))
                    sys.stdout.flush()
                except Exception as e:
                    rk = futures[fut]
                    opinions[rk] = f"[Hiba: {e}]"
    else:
        for rk, eng in debaters.items():
            _, resp = _call_one(rk, eng)
            opinions[rk] = resp

    valid_opinions = {}
    for rk, resp in opinions.items():
        phase = "independent" if _valid_response(resp) else "independent_error"
        if phase == "independent":
            valid_opinions[rk] = resp
        else:
            print(_yellow(
                f"   Excluding failed independent response from {rk} "
                f"({engines[rk].label})"
            ))
        session_log.append({
            "phase": phase, "role": rk,
            "speaker": engines[rk].label, "text": resp,
            "error": phase.endswith("_error")
        })
    return valid_opinions


def phase_issue_matrix(moderator: AIEngine, opinions: dict[str, str],
                        evidence: str, engines: dict,
                        session_log: list) -> str:
    """
    Phase 2: ModerĂˇtor Issue Matrix
    Nem Ă¶sszefoglalĂł, hanem vita-tĂ©rkĂ©p: konszenzus, konfliktus, kĂ©rdĂ©sek.
    """
    print("\n" + "â•"*64)
    print(_bold("đź“Š PHASE 2 â€” ModerĂˇtor Issue Matrix"))
    print("â•"*64)

    opinions_text = "\n\n".join(
        f"[{engines[rk].label} â€” {rk}]\n{text}"
        for rk, text in opinions.items()
    )

    system = """Te egy precĂ­z vita-moderĂˇtor AI vagy.
Feladatod NEM Ă¶sszefoglalĂˇs, hanem vitatĂ©rkĂ©p kĂ©szĂ­tĂ©se.
StrukturĂˇlt, tĂˇblĂˇzatos formĂˇtumban dolgozz."""

    prompt = f"""EVIDENCE PACKAGE (emlĂ©keztetĹ‘):
{_trunc(evidence, 2000)}

A RĂ‰SZTVEVĹK FĂśGGETLEN VĂ‰LEMĂ‰NYEI:
{opinions_text}

KĂ©szĂ­tsd el az Issue Matrixot ebben a formĂˇtumban:

## KONSZENZUS_PONTOK
[Miben Ă©rtett egyet mindenki? Min. 3-5 pont ++ jelĂ¶lĂ©ssel]

## VITAPONTOK
[Miben tĂ©rtek el? Ki mit ĂˇllĂ­tott? TĂˇblĂˇzatos formĂˇtum:]
| KĂ©rdĂ©s | Egyik nĂ©zet | MĂˇsik nĂ©zet | BizonyĂ­tĂ©k erĹ‘ssĂ©ge |

## GYENGE_ĂLLĂŤTĂSOK
[Hol hozott valaki nem bizonyĂ­tott vagy spekulatĂ­v Ă©rvet?]

## TOP_3_VITAKĂ‰RDĂ‰S
[A kĂ¶vetkezĹ‘ rebuttal kĂ¶rben ezt a 3 kĂ©rdĂ©st kell eldĂ¶nteni:]
1. [KĂ©rdĂ©s] â€” [MiĂ©rt fontos?]
2. [KĂ©rdĂ©s] â€” [MiĂ©rt fontos?]
3. [KĂ©rdĂ©s] â€” [MiĂ©rt fontos?]

## HIĂNYZĂ“_INFORMĂCIĂ“K
[Mit kellene mĂ©g tudni a jĂł dĂ¶ntĂ©shez?]"""

    matrix = moderator.call(system, [{"role":"user","content":prompt}],
                             stream=True, label="ModerĂˇtor â†’ Issue Matrix")
    if _is_failed_response(matrix):
        session_log.append({
            "phase": "issue_matrix_error", "role": "moderator",
            "speaker": moderator.label, "text": matrix, "error": True
        })
        _abort(f"Issue Matrix generation failed for moderator {moderator.label}: {matrix}")
    session_log.append({
        "phase": "issue_matrix", "role": "moderator",
        "speaker": moderator.label, "text": matrix
    })
    return matrix


def phase_rebuttal(engines: dict, issue_matrix: str, evidence: str,
                    roles_info: dict, session_log: list,
                    valid_roles: set[str] = None) -> dict[str, str]:
    """
    Phase 3: CĂ©lzott rebuttal
    Mindenki CSAK az issue matrixra reagĂˇl, nem egymĂˇs teljes szĂ¶vegĂ©re.
    """
    print("\n" + "â•"*64)
    print(_bold("âš”ď¸Ź  PHASE 3 â€” CĂ©lzott Rebuttal (csak az Issue Matrixra)"))
    print("â•"*64)

    debaters = {k: v for k, v in engines.items()
                if k not in ("moderator","judge")
                and (valid_roles is None or k in valid_roles)}
    rebuttals = {}

    for rk, eng in debaters.items():
        role_desc = roles_info.get(rk, rk)
        system = f"""Te {eng.label} vagy, {role_desc} szerepben.
A moderĂˇtor issue matrixa alapjĂˇn kell reagĂˇlnod.
SZABĂLYOK:
- Csak a TOP_3_VITAKĂ‰RDĂ‰SRE reagĂˇlj
- Minden kĂ©rdĂ©snĂ©l: ĂˇllĂ­tsd meg az ĂˇllĂˇspontodat, hozz bizonyĂ­tĂ©kot
- KĂ¶telezĹ‘: nevezd meg a legerĹ‘sebb ELLENTĂ‰TES Ă©rvet
- KĂ¶telezĹ‘: kĂ©rdĹ‘jelezz meg legalĂˇbb 1 feltĂ©telezĂ©st
Max 400 szĂł."""

        prompt = f"""ISSUE MATRIX (a moderĂˇtortĂłl):
{issue_matrix}

EVIDENCE PACKAGE (referencia):
{_trunc(evidence, 1500)}

ReagĂˇlj a TOP_3_VITAKĂ‰RDĂ‰SRE {role_desc} szemszĂ¶gĂ©bĹ‘l!
KĂ¶telezĹ‘: hozd fel a legerĹ‘sebb ellentĂ©tes Ă©rvet is."""

        resp = eng.call(system, [{"role":"user","content":prompt}],
                        stream=True, label=f"{eng.label} [Rebuttal]")
        phase = "rebuttal" if _valid_response(resp) else "rebuttal_error"
        if phase == "rebuttal":
            rebuttals[rk] = resp
        else:
            print(_yellow(
                f"   Excluding failed rebuttal from {rk} ({eng.label})"
            ))
        session_log.append({
            "phase": phase, "role": rk,
            "speaker": eng.label, "text": resp,
            "error": phase.endswith("_error")
        })

    return rebuttals


def phase_revision(engines: dict, issue_matrix: str, rebuttals: dict[str, str],
                    roles_info: dict, session_log: list,
                    valid_roles: set[str] = None) -> dict[str, dict]:
    """
    Phase 4: Position Revision
    Mindenki kitĂ¶lti a strukturĂˇlt JSON Ă¶nĂ©rtĂ©kelĂ©st.
    """
    print("\n" + "â•"*64)
    print(_bold("đź”„ PHASE 4 â€” ĂllĂˇspont-frissĂ­tĂ©s (JSON struktĂşra)"))
    print("â•"*64)

    debaters = {k: v for k, v in engines.items()
                if k not in ("moderator","judge")
                and (valid_roles is None or k in valid_roles)}
    revisions = {}

    all_rebuttals = "\n\n".join(
        f"[{engines[rk].label}]\n{text}"
        for rk, text in rebuttals.items()
    )

    for rk, eng in debaters.items():
        role_desc = roles_info.get(rk, rk)
        system = f"""Te {eng.label} vagy, {role_desc} szerepben.
A rebuttal kĂ¶r utĂˇn frissĂ­tsd ĂˇllĂˇspontodat.
KIZĂRĂ“LAG vĂˇlid JSON-t adj vissza, semmi mĂˇs szĂ¶veget."""

        prompt = f"""ISSUE MATRIX:
{_trunc(issue_matrix, 1000)}

REBUTTAL KĂ–R (minden rĂ©sztvevĹ‘):
{_trunc(all_rebuttals, 2000)}

Adj vissza CSAK JSON-t:
{{
  "role": "{role_desc}",
  "model": "{eng.label}",
  "changed_my_mind": true/false,
  "what_changed": "mi vĂˇltozott az ĂˇllĂˇspontodban (ha semmi: null)",
  "strongest_opposing_argument": "a legjobb ellenĂ©rv amit hallottĂˇl",
  "remaining_disagreement": "miben maradsz mĂˇs vĂ©lemĂ©nyen",
  "confidence": 0.0-1.0,
  "final_recommendation": "egy mondatos vĂ©gsĹ‘ javaslat"
}}"""

        raw = eng.call(system, [{"role":"user","content":prompt}],
                       stream=False)
        if _is_failed_response(raw):
            print(_yellow(
                f"   Excluding failed revision from {rk} ({eng.label})"
            ))
            session_log.append({
                "phase": "revision_error", "role": rk,
                "speaker": eng.label, "text": raw, "error": True
            })
            continue
        # JSON parse
        clean = raw.strip()
        if "```" in clean:
            for part in clean.split("```"):
                p = part.strip().lstrip("json").strip()
                if p.startswith("{"): clean = p; break
        try:
            parsed = json.loads(clean)
        except:
            parsed = {
                "role": role_desc, "model": eng.label,
                "changed_my_mind": False, "what_changed": None,
                "strongest_opposing_argument": "N/A",
                "remaining_disagreement": "N/A",
                "confidence": 0.5,
                "final_recommendation": raw[:200]
            }

        revisions[rk] = parsed
        # KiĂ­rĂˇs
        changed = "đź”„ VĂLTOZTATOTT" if parsed.get("changed_my_mind") else "âžˇď¸Ź  Tartja"
        conf = parsed.get("confidence", 0)
        sys.stdout.write(
            _c(f"\n{eng.emoji} {eng.label} [{role_desc}]: {changed} | "
               f"Bizalom: {conf:.0%} | "
               f"{parsed.get('final_recommendation','')[:80]}\n", eng.color)
        )
        sys.stdout.flush()

        session_log.append({
            "phase": "revision", "role": rk,
            "speaker": eng.label, "text": json.dumps(parsed, ensure_ascii=False)
        })

    return revisions


class SynthesisEngine:
    """Small multi-step final synthesis runner. Keeps Sprint 1 internal to the CLI flow."""
    def __init__(self, judge: AIEngine, evidence: str, issue_matrix: str,
                 rebuttals: dict, revisions: dict, engines: dict,
                 roles_info: dict, user_prompt: str, output_types: list,
                 lang: str, session_log: list,
                 task_profile: TaskProfile | dict | None = None,
                 output_contract: OutputContract | dict | None = None):
        self.judge = judge
        self.evidence = evidence
        self.issue_matrix = issue_matrix
        self.rebuttals = {rk: text for rk, text in rebuttals.items() if _valid_response(text)}
        self.revisions = {
            rk: rv for rk, rv in revisions.items()
            if rk in self.rebuttals and isinstance(rv, dict)
        }
        self.engines = engines
        self.roles_info = roles_info
        self.user_prompt = user_prompt
        self.output_types = output_types
        self.lang = lang
        self.session_log = session_log
        self.task_profile = task_profile_to_dict(task_profile)
        self.output_contract = output_contract_to_dict(output_contract or default_output_contract(lang))
        self.step_results = []
        self.repaired_steps = []
        self.repaired_sections = []
        self.incomplete_sections = []
        self.participant_roles = sorted(set(self.rebuttals.keys()) | set(self.revisions.keys()))

    def run(self) -> dict:
        _require_min_participants(self.participant_roles, "Final judge input")
        if not self._uses_default_contract():
            return self._run_contract()

        print(_bold("\nSynthesis step 1/6: final decision outline"))
        outline = self._outline_with_fallback()

        print(_bold("\nSynthesis step 2/6: main deliverable draft"))
        main_md = self._markdown_step(
            "main_deliverable_draft",
            self._main_deliverable_prompt(outline),
            SYNTHESIS_MAX_OUTPUT_TOKENS,
        )

        print(_bold("\nSynthesis step 3/6: implementation plan draft"))
        impl_md = self._markdown_step(
            "implementation_plan_draft",
            self._implementation_prompt(outline, main_md),
            SYNTHESIS_MAX_OUTPUT_TOKENS,
        )

        print(_bold("\nSynthesis step 4/6: AI context block"))
        ai_ctx = self._markdown_step(
            "ai_context_block",
            self._ai_context_prompt(outline, main_md, impl_md),
            min(2400, SYNTHESIS_MAX_OUTPUT_TOKENS),
        )

        outline = self._fill_structured_metadata(outline, main_md, impl_md, ai_ctx)

        print(_bold("\nSynthesis step 5/6: consistency review"))
        review = self._json_step(
            "consistency_review",
            self._review_prompt(outline, main_md, impl_md, ai_ctx),
            required_keys=["passed", "issues", "fixes", "final_warnings"],
        )

        print(_bold("\nSynthesis step 6/6: final artifact assembly"))
        synthesis = self._assemble(outline, main_md, impl_md, ai_ctx, review)
        self._log_step("final_artifact_assembly", json.dumps({
            "title": synthesis.get("title"),
            "sections": list(synthesis.get("synthesis_sections", {}).keys()),
            "failed_steps": synthesis.get("synthesis_process", {}).get("failed_steps", []),
        }, ensure_ascii=False), "json", ok=True)
        if "synthesis" in RUN_METADATA:
            RUN_METADATA["synthesis"]["synthesis_steps"] = self.step_results
        self.session_log.append({
            "phase": "judge", "role": "judge", "speaker": self.judge.label,
            "text": json.dumps({
                "verdict": synthesis.get("verdict", ""),
                "synthesis_process": synthesis.get("synthesis_process", {}),
            }, ensure_ascii=False)
        })
        return synthesis

    def _run_contract(self) -> dict:
        print(_bold("\nContract synthesis step 1: decision outline"))
        outline = self._contract_outline()

        deliverable_results = {}
        deliverables = self.output_contract.get("deliverables", [])
        total = len(deliverables)
        for idx, deliverable in enumerate(deliverables, 1):
            did = deliverable.get("id", f"deliverable_{idx}")
            print(_bold(f"\nContract synthesis step {idx + 1}/{total + 3}: {did}"))
            step_name = f"contract_deliverable:{did}"
            draft = self._markdown_step(
                step_name,
                self._contract_deliverable_prompt(outline, deliverable, deliverable_results),
                self._contract_deliverable_budget(deliverable),
            )
            deliverable_results[did] = {
                "id": did,
                "title": deliverable.get("title", did),
                "source": deliverable.get("source", "main"),
                "required": deliverable.get("required", True),
                "sections": deliverable.get("sections", []),
                "markdown": self._clean_contract_deliverable_markdown(draft, deliverable),
            }

        self._repair_contract_deliverables(outline, deliverable_results)

        main_md = self._join_contract_source(deliverable_results, "main")
        impl_md = self._join_contract_source(deliverable_results, "implementation")
        ai_ctx = self._join_contract_source(deliverable_results, "ai_context")
        if not main_md:
            main_md = self._join_all_contract_deliverables(deliverable_results)

        outline = self._fill_structured_metadata(outline, main_md, impl_md, ai_ctx)

        print(_bold(f"\nContract synthesis step {total + 2}/{total + 3}: consistency review"))
        review = self._json_step(
            "consistency_review",
            self._contract_review_prompt(outline, deliverable_results),
            required_keys=["passed", "issues", "fixes", "final_warnings"],
        )

        print(_bold(f"\nContract synthesis step {total + 3}/{total + 3}: final artifact assembly"))
        synthesis = self._assemble_contract(outline, deliverable_results, review)
        self._log_step("final_artifact_assembly", json.dumps({
            "title": synthesis.get("title"),
            "contract_id": self.output_contract.get("contract_id"),
            "deliverables": list(deliverable_results.keys()),
            "failed_steps": synthesis.get("synthesis_process", {}).get("failed_steps", []),
        }, ensure_ascii=False), "json", ok=True)
        if "synthesis" in RUN_METADATA:
            RUN_METADATA["synthesis"]["synthesis_steps"] = self.step_results
        self.session_log.append({
            "phase": "judge", "role": "judge", "speaker": self.judge.label,
            "text": json.dumps({
                "verdict": synthesis.get("verdict", ""),
                "synthesis_process": synthesis.get("synthesis_process", {}),
            }, ensure_ascii=False)
        })
        return synthesis

    def _base_context(self) -> str:
        rev_summary = "\n".join(
            f"[{r.get('model','?')} | {r.get('role','?')}] "
            f"VĂˇltoztatott: {r.get('changed_my_mind')} | "
            f"Bizalom: {r.get('confidence',0):.0%} | "
            f"Javaslat: {r.get('final_recommendation','')[:160]}"
            for r in self.revisions.values()
        ) or "(quick scenario: no structured revision round)"
        rebuttals_text = "\n\n".join(
            f"[{self.engines[rk].label} | {rk}]\n{text}"
            for rk, text in self.rebuttals.items()
        )
        outputs = "\n".join(f"- {o}" for o in self.output_types)
        lang_i = "Magyarul Ă­rj." if self.lang == "hu" else "Write in English."
        contract_block = "" if self._uses_default_contract() else f"\n{output_contract_prompt_block(self.output_contract)}\n"
        return f"""LANGUAGE: {lang_i}

USER GOAL:
{self.user_prompt or "(not provided)"}

{task_profile_prompt_block(self.task_profile)}

{contract_block}

REQUESTED OUTPUTS:
{outputs}

EVIDENCE PACKAGE:
{_trunc(self.evidence, 2600)}

ISSUE MATRIX:
{_trunc(self.issue_matrix, 1800) if self.issue_matrix else "(quick scenario: no issue matrix)"}

VALID PARTICIPANT INPUTS:
{_trunc(rebuttals_text, 3200)}

POSITION REVISIONS:
{rev_summary}
"""

    def _uses_default_contract(self) -> bool:
        return self.output_contract.get("contract_id") == DEFAULT_OUTPUT_CONTRACT["contract_id"]

    def _contract_deliverables_for_source(self, source: str) -> list[dict]:
        return [
            d for d in self.output_contract.get("deliverables", [])
            if d.get("source") == source
        ]

    def _contract_headings_for_source(self, source: str) -> str:
        lines = []
        for deliverable in self._contract_deliverables_for_source(source):
            if deliverable.get("sections"):
                lines.append(f"## {deliverable.get('title')}")
                for section in deliverable.get("sections", []):
                    lines.append(f"### {section.get('title')}")
            else:
                lines.append(f"## {deliverable.get('title')}")
        return "\n".join(lines) or "(No dedicated sections; create a concise complete section.)"

    def _contract_outline(self) -> dict:
        today = datetime.now().strftime("%Y-%m-%d")
        participants = [self.engines[k].label for k in self.participant_roles if k in self.engines]
        required = ["title", "goal_summary", "verdict", "consensus_points",
                    "open_issues", "risk_register", "decision_log", "recommendations"]
        outline = self._json_step(
            "final_decision_outline",
            self._outline_prompt(),
            required_keys=required,
            log_failure=False,
        )
        if not outline.get("_error"):
            return outline
        fallback = {
            "title": self.output_contract.get("title", "SzintĂ©zis"),
            "date": today,
            "goal_summary": self.user_prompt[:500] if self.user_prompt else "",
            "participants": participants,
            "judge": self.judge.label,
            "verdict": (
                "A dĂ¶ntĂ©si vĂˇzlat strukturĂˇlt JSON formĂˇban nem kĂ©szĂĽlt el, "
                "ezĂ©rt a vĂ©gsĹ‘ dokumentum a szerzĹ‘dĂ©sben megadott szakaszok alapjĂˇn, "
                "a vitĂˇban szereplĹ‘ Ă©rvĂ©nyes rĂ©sztvevĹ‘i inputokra tĂˇmaszkodva kĂ©szĂĽlt."
            ),
            "consensus_points": [],
            "open_issues": [],
            "risk_register": [],
            "decision_log": [],
            "recommendations": [],
            "_metadata_fallback": True,
        }
        self.repaired_steps.append("final_decision_outline")
        self._log_step(
            "final_decision_outline",
            json.dumps({"metadata_fallback": True, "verdict": fallback["verdict"]}, ensure_ascii=False),
            "json",
            ok=True,
            note="deterministic_contract_fallback",
        )
        return fallback

    def _contract_deliverable_prompt(self, outline: dict, deliverable: dict,
                                     prior_results: dict) -> str:
        sections = deliverable.get("sections") or []
        if sections:
            section_lines = "\n".join(
                f"- ### {s.get('title')} ({'required' if s.get('required', True) else 'optional'})"
                for s in sections
            )
            section_instruction = (
                "Use exactly these subsection headings and no other major sections:\n"
                f"{section_lines}"
            )
        else:
            section_instruction = (
                "This deliverable has no requested subsections. Write body content only; "
                "do not repeat the deliverable title as a heading."
            )
        prior_excerpt = "\n\n".join(
            f"[{v.get('title')}]\n{_trunc(v.get('markdown', ''), 900)}"
            for v in prior_results.values()
        ) or "(none yet)"
        return f"""{self._base_context()}

DECISION OUTLINE:
{json.dumps(outline, ensure_ascii=False, indent=2)}

ACTIVE OUTPUT CONTRACT:
{json.dumps(output_contract_summary(self.output_contract), ensure_ascii=False, indent=2)}

DELIVERABLE TO GENERATE:
{json.dumps({
    "id": deliverable.get("id"),
    "title": deliverable.get("title"),
    "required": deliverable.get("required", True),
    "source": deliverable.get("source"),
    "sections": deliverable.get("sections", []),
}, ensure_ascii=False, indent=2)}

PREVIOUS CONTRACT DELIVERABLES:
{prior_excerpt}

Write only this deliverable in polished markdown.
Do not use JSON. Do not include a top-level # heading.
Do not add default business-plan, implementation-plan, or AI-context sections unless this deliverable explicitly requests them.
Do not include any other contract deliverable.
Language: {self.output_contract.get('language') or self.lang}
Audience: {self.output_contract.get('audience')}

{section_instruction}

Every required section must be non-empty, specific, and evidence-grounded."""

    def _contract_deliverable_budget(self, deliverable: dict) -> int:
        sections = deliverable.get("sections") or []
        if len(sections) >= 3:
            return SYNTHESIS_MAX_OUTPUT_TOKENS
        if deliverable.get("source") == "verdict":
            return min(1600, SYNTHESIS_MAX_OUTPUT_TOKENS)
        return min(max(2400, SYNTHESIS_MAX_OUTPUT_TOKENS // 2), SYNTHESIS_MAX_OUTPUT_TOKENS)

    def _clean_contract_deliverable_markdown(self, markdown: str, deliverable: dict) -> str:
        text = self._strip_top_heading(markdown or "").strip()
        title = str(deliverable.get("title") or "").strip()
        if title:
            text = self._strip_duplicate_deliverable_headings(text, title)
        return text

    def _strip_duplicate_deliverable_headings(self, markdown: str, title: str) -> str:
        """Remove generated headings that duplicate the assembler-owned deliverable heading."""
        target = " ".join((title or "").strip().lower().split())
        if not target:
            return (markdown or "").strip()
        kept = []
        for line in (markdown or "").splitlines():
            stripped = line.strip()
            match = re.match(r"^#{1,6}\s+(.+?)\s*$", stripped)
            if match and " ".join(match.group(1).strip().lower().split()) == target:
                continue
            kept.append(line)
        return "\n".join(kept).strip()

    def _join_contract_source(self, deliverable_results: dict, source: str) -> str:
        parts = []
        for result in deliverable_results.values():
            if result.get("source") == source and result.get("markdown"):
                parts.append(f"## {result.get('title')}\n\n{result.get('markdown')}")
        return "\n\n".join(parts).strip()

    def _join_all_contract_deliverables(self, deliverable_results: dict) -> str:
        return "\n\n".join(
            f"## {r.get('title')}\n\n{r.get('markdown', '')}"
            for r in deliverable_results.values()
            if r.get("markdown")
        ).strip()

    def _contract_review_prompt(self, outline: dict, deliverable_results: dict) -> str:
        return f"""Review the contract-based synthesis for consistency. Return JSON only.

OUTPUT CONTRACT:
{json.dumps(output_contract_summary(self.output_contract), ensure_ascii=False, indent=2)}

OUTLINE:
{json.dumps(outline, ensure_ascii=False, indent=2)}

CONTRACT DELIVERABLE DRAFTS:
{_trunc(self._join_all_contract_deliverables(deliverable_results), 5000)}

Check only the requested contract deliverables and sections. Do not require default business/implementation/AI context sections unless the contract asks for them.

JSON shape:
{{
  "passed": true,
  "issues": ["specific inconsistency or gap"],
  "fixes": ["specific correction or note"],
  "final_warnings": ["important limitation to disclose"]
}}"""

    def _repair_contract_deliverables(self, outline: dict, deliverable_results: dict):
        max_rounds = 2
        for _ in range(max_rounds):
            current_md = self._build_contract_markdown_from_results(outline, deliverable_results)
            validation = self._validate_contract_markdown(current_md)
            issues = validation.get("section_issues", [])
            repairable = [
                issue for issue in issues
                if issue.get("level") in ("missing_section", "weak_section", "incomplete_section")
                and issue.get("section_id")
            ]
            if not repairable:
                return
            repaired_any = False
            for issue in repairable:
                deliverable = self._contract_deliverable_by_id(issue.get("deliverable_id"))
                if not deliverable:
                    continue
                did = deliverable.get("id")
                result = deliverable_results.get(did)
                if not result:
                    continue
                section = self._contract_section_by_id(deliverable, issue.get("section_id"))
                if not section:
                    continue
                previous = self._contract_section_content(result.get("markdown", ""), section.get("title", ""))
                repair = self.judge.call(
                    self._system("markdown"),
                    [{"role": "user", "content": self._contract_section_repair_prompt(
                        outline, deliverable, section, previous
                    )}],
                    stream=True,
                    label=f"Final Judge -> repair {did}:{section.get('id')}",
                    max_output_tokens=min(1800, SYNTHESIS_MAX_OUTPUT_TOKENS),
                )
                if not _valid_response(repair):
                    self.incomplete_sections.append({
                        "deliverable": did,
                        "section": section.get("id"),
                        "warnings": ["repair_failed_response"],
                    })
                    self._log_step(
                        f"contract_section_repair:{did}:{section.get('id')}",
                        repair,
                        "markdown",
                        ok=False,
                        note="repair_failed_response",
                    )
                    continue
                repaired_section = self._clean_repaired_contract_section(repair, section.get("title", ""))
                section_status = self._contract_section_status(repaired_section, heading_present=True)
                if section_status["status"] != "valid":
                    self.incomplete_sections.append({
                        "deliverable": did,
                        "section": section.get("id"),
                        "warnings": section_status.get("reasons", []),
                    })
                    self._log_step(
                        f"contract_section_repair:{did}:{section.get('id')}",
                        repaired_section,
                        "markdown",
                        ok=False,
                        note="repair_still_weak",
                    )
                    continue
                result["markdown"] = self._replace_contract_section(
                    result.get("markdown", ""),
                    section.get("title", ""),
                    repaired_section,
                )
                self._mark_step_repaired_if_present(f"contract_deliverable:{did}")
                repaired_any = True
                self.repaired_sections.append({
                    "deliverable": did,
                    "section": section.get("id"),
                    "warnings_before": issue.get("reasons", []),
                    "warnings_after": [],
                })
                self.repaired_steps.append(f"contract_section_repair:{did}:{section.get('id')}")
                self._log_step(
                    f"contract_section_repair:{did}:{section.get('id')}",
                    repaired_section,
                    "markdown",
                    ok=True,
                    note="targeted_contract_section_repair",
                )
            if not repaired_any:
                return

    def _mark_step_repaired_if_present(self, step_name: str):
        for step in self.step_results:
            if step.get("step") == step_name and not step.get("ok"):
                step["ok"] = True
                step["note"] = (step.get("note") + "; " if step.get("note") else "") + "repaired_by_targeted_section_repair"
        for entry in self.session_log:
            if entry.get("phase") == "synthesis_step" and entry.get("step") == step_name and entry.get("error"):
                entry["ok"] = True
                entry["error"] = False
                entry["note"] = (entry.get("note") + "; " if entry.get("note") else "") + "repaired_by_targeted_section_repair"

    def _contract_deliverable_by_id(self, deliverable_id: str) -> Optional[dict]:
        for deliverable in self.output_contract.get("deliverables", []):
            if deliverable.get("id") == deliverable_id:
                return deliverable
        return None

    def _contract_section_by_id(self, deliverable: dict, section_id: str) -> Optional[dict]:
        for section in deliverable.get("sections", []) or []:
            if section.get("id") == section_id:
                return section
        return None

    def _contract_section_repair_prompt(self, outline: dict, deliverable: dict,
                                        section: dict, previous: str) -> str:
        return f"""{self._base_context()}

Complete only the section "{section.get('title')}" for contract "{self.output_contract.get('contract_id')}".
Do not repeat other sections. Do not include the deliverable title.

ACTIVE CONTRACT:
- contract_id: {self.output_contract.get('contract_id')}
- title: {self.output_contract.get('title')}
- language: {self.output_contract.get('language') or self.lang}
- audience: {self.output_contract.get('audience')}

DELIVERABLE:
- id: {deliverable.get('id')}
- title: {deliverable.get('title')}

SECTION:
- id: {section.get('id')}
- title: {section.get('title')}

DECISION OUTLINE:
{json.dumps(outline, ensure_ascii=False, indent=2)}

SOURCE EVIDENCE SUMMARY:
{_trunc(self.evidence, 2200)}

PREVIOUS INCOMPLETE SECTION TEXT:
{previous or "(missing or placeholder)"}

Expected structure:
### {section.get('title')}
- Provide at least 3 concrete, actionable bullets.
- Include evidence-grounded technical actions, sequencing, or validation criteria.
- Avoid placeholders, TBD, generic filler, and unfinished fragments.

For technical audit remediation sections:
- Critical fixes: immediate blockers and correctness/security fixes.
- Near-term hardening: reliability, tests, observability, and maintainability.
- Long-term improvements: architecture, scalability, automation, and governance."""

    def _clean_repaired_contract_section(self, markdown: str, section_title: str) -> str:
        text = self._strip_top_heading(markdown or "").strip()
        extracted = self._contract_section_content(text, section_title)
        if extracted:
            return extracted.strip()
        lines = text.splitlines()
        while lines and not lines[0].strip():
            lines = lines[1:]
        if lines and re.sub(r"^#+\s*", "", lines[0].strip()).strip().lower() == section_title.lower():
            lines = lines[1:]
        return "\n".join(lines).strip()

    def _replace_contract_section(self, markdown: str, section_title: str, replacement: str) -> str:
        text = markdown or ""
        heading_re = re.compile(rf"^(?P<hashes>#+)\s+{re.escape(section_title)}\s*$", re.IGNORECASE | re.MULTILINE)
        match = heading_re.search(text)
        if not match:
            suffix = "\n\n" if text.strip() else ""
            return f"{text.strip()}{suffix}### {section_title}\n\n{replacement.strip()}".strip()
        start = match.end()
        level = len(match.group("hashes"))
        next_re = re.compile(rf"^#{{1,{level}}}\s+", re.MULTILINE)
        next_match = next_re.search(text, start)
        end = next_match.start() if next_match else len(text)
        return (text[:start] + "\n\n" + replacement.strip() + "\n\n" + text[end:]).strip()

    def _outline_prompt(self) -> str:
        today = datetime.now().strftime("%Y-%m-%d")
        participants = [self.engines[k].label for k in self.participant_roles if k in self.engines]
        return f"""{self._base_context()}

Create a compact decision outline as JSON only.
Use evidence-grounded, concrete claims. Do not write the full document here.

Required JSON shape:
{{
  "title": "Project name - Final Documentation",
  "date": "{today}",
  "goal_summary": "2-4 sentence goal summary",
  "participants": {json.dumps(participants, ensure_ascii=False)},
  "judge": "{self.judge.label}",
  "verdict": "4-6 sentence final decision with concrete tradeoffs",
  "consensus_points": ["specific consensus point"],
  "open_issues": ["specific open question"],
  "risk_register": [{{"risk": "...", "severity": "high/medium/low", "mitigation": "..."}}],
  "decision_log": [{{"decision": "...", "rationale": "...", "confidence": 0.0}}],
  "recommendations": ["actionable recommendation"]
}}"""

    def _main_deliverable_prompt(self, outline: dict) -> str:
        if not self._uses_default_contract():
            return f"""{self._base_context()}

DECISION OUTLINE:
{json.dumps(outline, ensure_ascii=False, indent=2)}

Write the main human-readable deliverable content in polished markdown.
Do not wrap it in JSON. Do not include a top-level # heading.
Follow the OutputContract. Cover the deliverables and sections whose source is "main".
Use exactly these headings where applicable:
{self._contract_headings_for_source("main")}

Each substantive section should contain concrete details from the debate/evidence."""
        return f"""{self._base_context()}

DECISION OUTLINE:
{json.dumps(outline, ensure_ascii=False, indent=2)}

Write the main human-readable deliverable in polished markdown.
Do not wrap it in JSON. Do not include a top-level # heading.
This text will be inserted under "## VĂ©gsĹ‘ ĂĽzleti terv".
Make it specific, useful, and structured.
Include these headings:
### TermĂ©k Ă©s Ă©rtĂ©kajĂˇnlat
### Piac Ă©s go-to-market
### Jelenlegi Ăˇllapot Ă©s korlĂˇtok
### BevĂ©teli / mĹ±kĂ¶dĂ©si modell
### VersenypozĂ­ciĂł
### Ăśzleti kockĂˇzatok
### Ăśzleti ajĂˇnlĂˇsok

Each substantive section should contain concrete details from the debate/evidence."""

    def _implementation_prompt(self, outline: dict, main_md: str) -> str:
        if not self._uses_default_contract():
            return f"""{self._base_context()}

DECISION OUTLINE:
{json.dumps(outline, ensure_ascii=False, indent=2)}

MAIN DELIVERABLE EXCERPT:
{_trunc(main_md, 1800)}

Write the implementation, roadmap, remediation, or action-plan deliverable content requested by the OutputContract.
Do not use JSON. Do not include a top-level # heading.
Follow the deliverables and sections whose source is "implementation".
Use exactly these headings where applicable:
{self._contract_headings_for_source("implementation")}

For task-like content, include objective, priority, dependencies, concrete steps, and definition of done."""
        return f"""{self._base_context()}

DECISION OUTLINE:
{json.dumps(outline, ensure_ascii=False, indent=2)}

MAIN DELIVERABLE EXCERPT:
{_trunc(main_md, 1800)}

Write a markdown implementation plan. Do not use JSON.
Do not include a top-level # heading.
This text will be inserted under "## VĂ©gsĹ‘ megvalĂłsĂ­tĂˇsi terv".
Include:
### ArchitektĂşra irĂˇny
### Roadmap
For every roadmap task use this exact structure:
#### Task N: [short title]
**Objective:** ...
**Priority:** high/medium/low
**Recommended phase/sprint:** ...
**Dependencies:** ...
**Concrete steps:**
- ...
**Definition of Done:**
- ...
### Technikai kockĂˇzatok
Use concrete tasks, dependencies, and complete definition-of-done bullets."""

    def _ai_context_prompt(self, outline: dict, main_md: str, impl_md: str) -> str:
        if not self._uses_default_contract():
            return f"""Create the AI/context deliverable requested by the OutputContract.
Do not use JSON. Do not include a top-level # heading.
Use exactly these headings where applicable:
{self._contract_headings_for_source("ai_context")}

{self._base_context()}

OUTLINE:
{json.dumps(outline, ensure_ascii=False, indent=2)}

MAIN DELIVERABLE:
{_trunc(main_md, 1800)}

IMPLEMENTATION / ACTION PLAN:
{_trunc(impl_md, 1800)}"""
        return f"""Create an AI-ready markdown context block from the material below.
Do not use JSON. Keep it concise but complete enough for a future AI agent.

OUTLINE:
{json.dumps(outline, ensure_ascii=False, indent=2)}

MAIN DELIVERABLE:
{_trunc(main_md, 1800)}

IMPLEMENTATION PLAN:
{_trunc(impl_md, 1800)}

Required headings:
Do not include a top-level # heading.
### Mi ez a projekt
### Jelenlegi Ăˇllapot
### Ăśzleti cĂ©lok
### Technikai irĂˇny
### Kritikus kockĂˇzatok
### AjĂˇnlott kĂ¶vetkezĹ‘ prompt"""

    def _review_prompt(self, outline: dict, main_md: str, impl_md: str, ai_ctx: str) -> str:
        return f"""Review the synthesis outputs for consistency. Return JSON only.

OUTLINE:
{json.dumps(outline, ensure_ascii=False, indent=2)}

MAIN DELIVERABLE:
{_trunc(main_md, 2200)}

IMPLEMENTATION PLAN:
{_trunc(impl_md, 1800)}

AI CONTEXT:
{_trunc(ai_ctx, 1200)}

JSON shape:
{{
  "passed": true,
  "issues": ["specific inconsistency or gap"],
  "fixes": ["specific correction or note"],
  "final_warnings": ["important limitation to disclose"]
}}"""

    def _system(self, output_type: str) -> str:
        if output_type == "json":
            return "You are the Final Judge. Return strict valid JSON only. No markdown fences."
        return "You are the Final Judge. Write polished, evidence-grounded markdown. Do not return JSON."

    def _outline_with_fallback(self) -> dict:
        required = ["title", "goal_summary", "verdict", "consensus_points",
                    "open_issues", "risk_register", "decision_log", "recommendations"]
        outline = self._json_step(
            "final_decision_outline",
            self._outline_prompt(),
            required_keys=required,
            log_failure=False,
        )
        if not outline.get("_error"):
            return outline

        fallback_prompt = f"""{self._base_context()}

The compact JSON decision outline could not be produced.
Write a concise markdown decision summary instead. Do not use JSON.
Include:
### VĂ©gsĹ‘ Ă­tĂ©let
### FĹ‘ konszenzusok
### Nyitott kĂ©rdĂ©sek
### KockĂˇzatok
### DĂ¶ntĂ©sek
### KĂ¶vetkezĹ‘ lĂ©pĂ©sek"""
        decision_md = self._markdown_step(
            "final_decision_outline_markdown_fallback",
            fallback_prompt,
            1200,
        )
        verdict = (
            self._extract_markdown_section(decision_md, "VĂ©gsĹ‘ Ă­tĂ©let")
            or self._extract_markdown_section(decision_md, "Végső ítélet")
            or self._first_markdown_body_block(decision_md, 1200)
        )
        fallback = {
            "title": "SzintĂ©zis",
            "date": datetime.now().strftime("%Y-%m-%d"),
            "goal_summary": self.user_prompt[:500] if self.user_prompt else "",
            "participants": [self.engines[k].label for k in self.participant_roles if k in self.engines],
            "judge": self.judge.label,
            "verdict": verdict.strip(),
            "consensus_points": [],
            "open_issues": [],
            "risk_register": [],
            "decision_log": [],
            "recommendations": [],
            "decision_summary_markdown": decision_md,
            "_metadata_fallback": True,
        }
        self.repaired_steps.append("final_decision_outline")
        self._log_step(
            "final_decision_outline",
            json.dumps({"metadata_fallback": True, "verdict": fallback["verdict"]}, ensure_ascii=False),
            "json",
            ok=True,
            note="markdown_metadata_fallback",
        )
        return fallback

    def _fill_structured_metadata(self, outline: dict, main_md: str,
                                  impl_md: str, ai_ctx: str) -> dict:
        required_lists = [
            "consensus_points", "open_issues", "risk_register",
            "decision_log", "recommendations",
        ]
        if not outline.get("_error") and all(outline.get(k) for k in required_lists):
            outline["_structured_metadata_status"] = "ok"
            outline["_metadata_available"] = True
            return outline

        prompt = f"""Extract compact structured metadata from the synthesis drafts.
Return strict JSON only. Keep every string short, single-line, and without markdown.
Use at most:
- consensus_points: 3 items, max 120 chars each
- open_issues: 3 items, max 120 chars each
- risk_register: 3 objects, max 120 chars per field
- decision_log: 3 objects, max 120 chars per field
- recommendations: 5 items, max 120 chars each

MAIN DELIVERABLE:
{_trunc(main_md, 1600)}

IMPLEMENTATION PLAN:
{_trunc(impl_md, 1800)}

AI CONTEXT:
{_trunc(ai_ctx, 800)}

Required JSON:
{{
  "consensus_points": ["..."],
  "open_issues": ["..."],
  "risk_register": [{{"risk": "...", "severity": "high/medium/low", "mitigation": "..."}}],
  "decision_log": [{{"decision": "...", "rationale": "...", "confidence": 0.0}}],
  "recommendations": ["..."]
}}"""
        metadata = self._json_step(
            "structured_metadata_extraction",
            prompt,
            required_keys=required_lists,
        )
        metadata = self._sanitize_metadata_payload(metadata)
        if metadata.get("_error") or not self._valid_metadata_payload(metadata):
            self._apply_empty_metadata(
                outline,
                "structured metadata extraction failed",
                detail=metadata.get("message", "") if isinstance(metadata, dict) else "",
            )
            return outline

        filled = []
        for key in required_lists:
            outline[key] = metadata.get(key) or []
            if outline[key]:
                filled.append(key)
        outline["_structured_metadata_status"] = "ok"
        outline["_metadata_available"] = bool(filled)
        if not filled:
            self._apply_empty_metadata(outline, "structured metadata extraction returned empty values")
            return outline
        outline["_structured_metadata_filled_fields"] = filled
        if filled:
            self.repaired_steps.append("structured_metadata")
        return outline

    def _sanitize_metadata_payload(self, metadata: dict) -> dict:
        if not isinstance(metadata, dict) or metadata.get("_error"):
            return metadata

        def clean_text(value, max_len=120):
            text = " ".join(str(value or "").split())
            text = re.sub(r"^[#>*\-\s]+", "", text).strip()
            text = text.replace("**", "").replace("__", "").replace("`", "")
            if len(text) > max_len:
                text = text[:max_len].rsplit(" ", 1)[0].rstrip(".,;:") + "..."
            return text

        sanitized = {}
        for key in ("consensus_points", "open_issues", "recommendations"):
            values = metadata.get(key) if isinstance(metadata.get(key), list) else []
            sanitized[key] = [
                clean_text(v)
                for v in values[:5]
                if clean_text(v)
            ]

        risks = metadata.get("risk_register") if isinstance(metadata.get("risk_register"), list) else []
        sanitized["risk_register"] = []
        for item in risks[:3]:
            if not isinstance(item, dict):
                continue
            risk = clean_text(item.get("risk"))
            if not risk:
                continue
            severity = clean_text(item.get("severity", "medium"), 20).lower()
            if severity not in ("high", "medium", "low"):
                severity = "medium"
            sanitized["risk_register"].append({
                "risk": risk,
                "severity": severity,
                "mitigation": clean_text(item.get("mitigation"), 120),
            })

        decisions = metadata.get("decision_log") if isinstance(metadata.get("decision_log"), list) else []
        sanitized["decision_log"] = []
        for item in decisions[:3]:
            if not isinstance(item, dict):
                continue
            decision = clean_text(item.get("decision"))
            if not decision:
                continue
            confidence = item.get("confidence", 0.0)
            try:
                confidence = max(0.0, min(1.0, float(confidence)))
            except (TypeError, ValueError):
                confidence = 0.0
            sanitized["decision_log"].append({
                "decision": decision,
                "rationale": clean_text(item.get("rationale"), 120),
                "confidence": confidence,
            })
        return sanitized

    def _valid_metadata_payload(self, metadata: dict) -> bool:
        if not isinstance(metadata, dict):
            return False
        list_keys = ("consensus_points", "open_issues", "risk_register", "decision_log", "recommendations")
        if not any(metadata.get(k) for k in list_keys):
            return False
        for key in ("consensus_points", "open_issues", "recommendations"):
            if metadata.get(key) and not isinstance(metadata[key], list):
                return False
        for key in ("risk_register", "decision_log"):
            if metadata.get(key) and not isinstance(metadata[key], list):
                return False
        return True

    def _apply_empty_metadata(self, outline: dict, reason: str, detail: str = ""):
        outline["consensus_points"] = []
        outline["open_issues"] = []
        outline["risk_register"] = []
        outline["decision_log"] = []
        outline["recommendations"] = []
        outline["_structured_metadata_status"] = "failed"
        outline["_metadata_available"] = False
        outline["_structured_metadata_reason"] = reason
        outline["_structured_metadata_error_detail"] = detail
        outline["_structured_metadata_filled_fields"] = []

    def _markdown_step(self, name: str, prompt: str, max_tokens: int) -> str:
        raw = self.judge.call(
            self._system("markdown"),
            [{"role": "user", "content": prompt}],
            stream=True,
            label=f"Final Judge -> {name}",
            max_output_tokens=min(max_tokens, SYNTHESIS_MAX_OUTPUT_TOKENS),
        )
        ok = _valid_response(raw)
        if not ok:
            text = self._failed_section(name, raw)
            self._log_step(name, text, "markdown", ok=False)
            return text
        warnings = self._truncation_warnings(raw)
        if warnings:
            repaired = self.judge.call(
                self._system("markdown"),
                [{"role": "user", "content": self._completion_repair_prompt(name, prompt, raw, warnings)}],
                stream=True,
                label=f"Final Judge -> {name} repair",
                max_output_tokens=min(max_tokens, SYNTHESIS_MAX_OUTPUT_TOKENS),
            )
            if _valid_response(repaired) and not self._truncation_warnings(repaired):
                self.repaired_steps.append(name)
                self._log_step(name, repaired, "markdown", ok=True, note="repaired_truncation")
                return repaired.strip()
            text = repaired if _valid_response(repaired) else raw
            self._log_step(name, text, "markdown", ok=False, note="truncation_detected")
            return (text or raw).strip()
        self._log_step(name, raw, "markdown", ok=True)
        return raw.strip()

    def _json_step(self, name: str, prompt: str, required_keys: list[str],
                   log_failure: bool = True) -> dict:
        raw = self.judge.call(
            self._system("json"),
            [{"role": "user", "content": prompt}],
            stream=True,
            label=f"Final Judge -> {name}",
            max_output_tokens=min(1400, SYNTHESIS_MAX_OUTPUT_TOKENS),
        )
        parsed, err = self._parse_json(raw, required_keys)
        if parsed is not None:
            self._log_step(name, json.dumps(parsed, ensure_ascii=False), "json", ok=True)
            return parsed

        repair_prompt = f"""Repair this into strict JSON only.
Parse error: {err}
Required keys: {', '.join(required_keys)}
Use compact single-line strings only. Do not use markdown, comments, or multiline strings.

RAW OUTPUT:
{_trunc(raw, 3000)}"""
        repaired = self.judge.call(
            self._system("json"),
            [{"role": "user", "content": repair_prompt}],
            stream=True,
            label=f"Final Judge -> {name} repair",
            max_output_tokens=min(1200, SYNTHESIS_MAX_OUTPUT_TOKENS),
        )
        parsed, err2 = self._parse_json(repaired, required_keys)
        if parsed is not None:
            self._log_step(name, json.dumps(parsed, ensure_ascii=False), "json", ok=True,
                           note="repaired")
            self.repaired_steps.append(name)
            return parsed

        failure = {
            "_error": True,
            "step": name,
            "message": f"JSON step failed; initial parse: {err}; repair parse: {err2}",
            "raw_excerpt": _trunc(repaired or raw, 1200),
        }
        if log_failure:
            self._log_step(name, json.dumps(failure, ensure_ascii=False), "json", ok=False)
        return failure

    def _parse_json(self, raw: str, required_keys: list[str]) -> tuple[Optional[dict], str]:
        if _is_failed_response(raw):
            return None, raw or "empty response"
        clean = (raw or "").strip()
        if "```" in clean:
            for part in clean.split("```"):
                p = part.strip().lstrip("json").strip()
                if p.startswith("{"):
                    clean = p
                    break
        try:
            parsed = json.loads(clean)
        except Exception as e:
            return None, f"{type(e).__name__}: {e}"
        missing = [k for k in required_keys if k not in parsed]
        if missing:
            return None, "missing keys: " + ", ".join(missing)
        return parsed, ""

    def _completion_repair_prompt(self, name: str, original_prompt: str,
                                  raw: str, warnings: list[str]) -> str:
        return f"""The previous markdown output for {name} appears incomplete.
Warnings: {', '.join(warnings)}

Regenerate the complete section from scratch. Keep the required structure,
finish every bullet and definition of done, and do not end with a dangling heading,
colon, TBD, or fragment.

ORIGINAL TASK:
{_trunc(original_prompt, 2600)}

INCOMPLETE OUTPUT:
{_trunc(raw, 1800)}"""

    def _truncation_warnings(self, text: str) -> list[str]:
        if not text or _is_failed_response(text):
            return ["empty_or_failed"]
        stripped = text.rstrip()
        raw_lines = stripped.splitlines()
        lines = [ln.rstrip() for ln in raw_lines if ln.strip()]
        if not lines:
            return ["empty"]
        last = lines[-1].strip()
        warnings = []
        if not _ends_with_closed_fenced_code_block(stripped):
            low = last.lower()
            if low in ("- minden", "* minden", "minden", "- all", "* all"):
                warnings.append("suspicious_fragment_minden")
            if "tbd" in low:
                warnings.append("tbd_marker")
            if last.endswith(":"):
                warnings.append("ends_after_colon")
            if last.startswith("#"):
                warnings.append("ends_with_heading")
            if last in ("-", "*") or last.startswith(("- [", "* [")):
                warnings.append("unfinished_bullet")
            if len(last.split()) <= 2 and last.startswith(("- ", "* ")):
                warnings.append("too_short_final_bullet")
            if self._looks_mid_word_ending(last):
                warnings.append("possible_mid_word_ending")
        warnings.extend(self._abrupt_heading_transitions(_lines_outside_code_fences(raw_lines)))
        warnings.extend(_short_heading_content_warnings(text))
        warnings.extend(self._task_block_warnings(stripped))
        return warnings

    def _looks_mid_word_ending(self, line: str) -> bool:
        if not line:
            return False
        stripped = line.strip()
        if stripped.endswith(("...", "â€¦")):
            return True
        if stripped[-1] in ".!?)]}â€ť\"'`":
            return False
        if stripped.startswith(("- ", "* ", "#", "|")):
            return False
        if ":" in stripped[-3:]:
            return True
        words = stripped.split()
        if not words:
            return False
        last_word = words[-1].strip(".,;:!?)]}â€ť\"'")
        if len(words) <= 3 and len(last_word) >= 4:
            return True
        common_incomplete_suffixes = ("eszk", "funk", "valid", "integr", "konfig", "architekt", "implement")
        return any(last_word.lower().endswith(sfx) for sfx in common_incomplete_suffixes)

    def _abrupt_heading_transitions(self, lines: list[str]) -> list[str]:
        warnings = []
        prev = ""
        for line in lines:
            cur = line.strip()
            if cur.startswith("#") and prev:
                prev_clean = prev.strip()
                if (
                    len(prev_clean) > 24
                    and not prev_clean.startswith(("- ", "* ", "|", "**"))
                    and not re.match(r"^\d+[\.)]\s+", prev_clean)
                    and not _strip_trailing_markdown_emphasis(prev_clean).endswith((".", "!", "?", ":", ";", ")", "]", "â€ť", "\""))
                    and not prev_clean.endswith((".", "!", "?", ":", ";", ")", "]", "â€ť", "\""))
                ):
                    warnings.append("abrupt_transition_before_heading")
            if cur:
                prev = cur
        return warnings

    def _task_block_warnings(self, markdown: str) -> list[str]:
        warnings = []
        task_blocks = self._task_blocks(markdown)
        for idx, block in enumerate(task_blocks, 1):
            lower = block.lower()
            required = {
                "objective": "**objective:**",
                "priority": "**priority:**",
                "dependencies": "**dependencies:**",
                "concrete_steps": "**concrete steps:**",
                "definition_of_done": "**definition of done:**",
            }
            for label, marker in required.items():
                if marker not in lower:
                    warnings.append(f"task_{idx}_missing_{label}")
            dod = self._field_block(block, "**Definition of Done:**")
            dod_bullets = [ln for ln in dod.splitlines() if ln.strip().startswith(("- ", "* ")) and len(ln.strip()[2:].split()) >= 3]
            if len(dod_bullets) < 2:
                warnings.append(f"task_{idx}_incomplete_definition_of_done")
            steps = self._field_block(block, "**Concrete steps:**")
            step_bullets = [ln for ln in steps.splitlines() if ln.strip().startswith(("- ", "* ")) and len(ln.strip()[2:].split()) >= 3]
            if len(step_bullets) < 2:
                warnings.append(f"task_{idx}_incomplete_concrete_steps")
        return warnings

    def _task_blocks(self, markdown: str) -> list[str]:
        lines = _lines_outside_code_fences((markdown or "").splitlines())
        starts = [i for i, ln in enumerate(lines) if ln.strip().lower().startswith("#### task ")]
        blocks = []
        for pos, start in enumerate(starts):
            end = starts[pos + 1] if pos + 1 < len(starts) else len(lines)
            blocks.append("\n".join(lines[start:end]))
        return blocks

    def _field_block(self, block: str, marker: str) -> str:
        lines = block.splitlines()
        start = None
        for i, ln in enumerate(lines):
            if ln.strip().lower() == marker.lower():
                start = i + 1
                break
        if start is None:
            return ""
        end = len(lines)
        for j in range(start, len(lines)):
            txt = lines[j].strip()
            if txt.startswith("**") and txt.endswith(":**"):
                end = j
                break
            if txt.startswith("#### Task "):
                end = j
                break
        return "\n".join(lines[start:end])

    def _strip_top_heading(self, markdown: str) -> str:
        lines = (markdown or "").strip().splitlines()
        while lines and lines[0].strip().startswith("# "):
            lines = lines[1:]
            while lines and not lines[0].strip():
                lines = lines[1:]
        return "\n".join(lines).strip()

    def _strip_leading_section_heading(self, markdown: str) -> str:
        lines = (markdown or "").strip().splitlines()
        while lines and re.match(r"^#{1,3}\s+\S", lines[0].strip()):
            lines = lines[1:]
            while lines and not lines[0].strip():
                lines = lines[1:]
        return "\n".join(lines).strip()

    def _sanitize_verdict_markdown(self, markdown: str) -> str:
        cleaned = self._strip_top_heading(markdown or "")
        if not cleaned:
            return ""
        lines = []
        for line in cleaned.splitlines():
            # A verdict is inserted under the canonical final-verdict heading.
            # Nested document titles from the judge would create duplicate top-level
            # structures and confuse downstream validation.
            if line.startswith("# "):
                continue
            lines.append(line)
        cleaned = "\n".join(lines).strip()
        lower = cleaned.lower()
        for marker in ("\n## tartalom", "\n## table of contents"):
            idx = lower.find(marker)
            if idx > 0:
                cleaned = cleaned[:idx].strip()
                break
        return cleaned

    def _extract_markdown_section(self, markdown: str, heading_text: str) -> str:
        lines = (markdown or "").splitlines()
        start = None
        level = 0
        in_fence = False
        for i, line in enumerate(lines):
            if line.strip().startswith("```"):
                in_fence = not in_fence
                continue
            if not in_fence and line.lstrip("#").strip().lower() == heading_text.lower():
                start = i + 1
                level = len(line) - len(line.lstrip("#"))
                break
        if start is None:
            return ""
        end = len(lines)
        in_fence = False
        for j in range(start, len(lines)):
            if lines[j].strip().startswith("```"):
                in_fence = not in_fence
                continue
            if not in_fence and lines[j].startswith("#"):
                next_level = len(lines[j]) - len(lines[j].lstrip("#"))
                if next_level <= level:
                    end = j
                    break
        return "\n".join(lines[start:end]).strip()

    def _first_markdown_body_block(self, markdown: str, max_chars: int = 1200) -> str:
        lines = []
        for line in (markdown or "").splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                if lines:
                    break
                continue
            if stripped == "---" and lines:
                break
            lines.append(line)
        text = "\n".join(lines).strip()
        if len(text) <= max_chars:
            return text
        cut = text[:max_chars]
        last_break = max(cut.rfind("\n\n"), cut.rfind(". "), cut.rfind("! "), cut.rfind("? "))
        if last_break > 120:
            return cut[:last_break + 1].strip()
        return cut.rstrip()

    def _failed_section(self, name: str, raw: str) -> str:
        return (
            f"## Synthesis step failed: {name}\n\n"
            "This section was not generated because the model call failed. "
            "The pipeline preserved the failure instead of inventing content.\n\n"
            f"```text\n{_trunc(raw or '(empty response)', 1200)}\n```"
        )

    def _log_step(self, name: str, text: str, output_type: str, ok: bool, note: str = ""):
        entry = {
            "phase": "synthesis_step",
            "role": "judge",
            "speaker": self.judge.label,
            "step": name,
            "output_type": output_type,
            "ok": ok,
            "text": text,
        }
        if note:
            entry["note"] = note
        if not ok:
            entry["error"] = True
        self.session_log.append(entry)
        self.step_results.append({
            "step": name,
            "output_type": output_type,
            "ok": ok,
            "note": note,
            "chars": len(text or ""),
        })

    def _assemble(self, outline: dict, main_md: str, impl_md: str,
                  ai_ctx: str, review: dict) -> dict:
        today = datetime.now().strftime("%Y-%m-%d")
        metadata_step_names = {"structured_metadata_extraction"}
        failed_steps = _effective_human_failed_steps(
            self.step_results, self.repaired_steps, metadata_step_names
        )
        metadata_failed_steps = [
            s["step"] for s in self.step_results
            if not s.get("ok") and s.get("step") in metadata_step_names
        ]
        title = outline.get("title") if not outline.get("_error") else f"SzintĂ©zis - {today}"
        verdict = outline.get("verdict") if not outline.get("_error") else "A vĂ©gsĹ‘ Ă­tĂ©let strukturĂˇlt metaadatai nem kĂ©szĂĽltek el, de a szintĂ©zis rĂ©szletes szĂ¶veges Ă©rtĂ©kelĂ©st ad."
        metadata_status = outline.get("_structured_metadata_status", "failed")
        metadata_available = bool(outline.get("_metadata_available")) and metadata_status == "ok"
        recommendations = outline.get("recommendations", []) if metadata_available and isinstance(outline.get("recommendations"), list) else []
        verdict = self._sanitize_verdict_markdown(verdict)
        main_clean = self._strip_leading_section_heading(main_md)
        impl_clean = self._strip_leading_section_heading(impl_md)
        ai_clean = self._strip_leading_section_heading(ai_ctx)
        main_clean, impl_clean, ai_clean = self._repair_sections_before_assembly(
            main_clean, impl_clean, ai_clean
        )
        final_markdown = self._build_final_markdown(outline, verdict, main_clean, impl_clean, ai_clean, review)
        validation = self._validate_final_markdown(final_markdown)
        human_artifact_status = validation["human_artifact_status"]
        final_status = self._final_validation_status(human_artifact_status, metadata_status, failed_steps)
        validation["structured_metadata_status"] = metadata_status
        validation["final_validation_status"] = final_status
        if failed_steps:
            validation["status"] = "failed"
            validation["truncation_warnings"].append({
                "section": "synthesis_steps",
                "warning": "required_step_failed: " + ", ".join(failed_steps),
            })
            validation["human_artifact_status"] = "failed"
            validation["final_validation_status"] = "failed"
        if self.incomplete_sections and validation["human_artifact_status"] == "ok":
            validation["human_artifact_status"] = "ok_with_warnings"
            validation["final_validation_status"] = self._final_validation_status("ok_with_warnings", metadata_status, failed_steps)
        implementation_items = []
        sections = [
            {"heading": "VĂ©gsĹ‘ ĂĽzleti terv", "content": main_clean, "bullets": []},
            {"heading": "VĂ©gsĹ‘ megvalĂłsĂ­tĂˇsi terv", "content": impl_clean, "bullets": []},
        ]
        synthesis = {
            "title": title,
            "date": outline.get("date", today) if not outline.get("_error") else today,
            "goal_summary": outline.get("goal_summary", self.user_prompt[:300] if self.user_prompt else ""),
            "participants": [self.engines[k].label for k in self.participant_roles if k in self.engines],
            "judge": self.judge.label,
            "task_profile": self.task_profile,
            "output_contract": output_contract_summary(self.output_contract),
            "verdict": verdict,
            "final_markdown": final_markdown,
            "main_deliverable_markdown": main_clean,
            "implementation_plan_markdown": impl_clean,
            "ai_context_block": ai_clean,
            "md_context_block": ai_clean,
            "synthesis_sections": {
                "main_deliverable": main_clean,
                "implementation_plan": impl_clean,
                "ai_context_block": ai_clean,
            },
            "synthesis_process": {
                "steps": self.step_results,
                "failed_steps": failed_steps,
                "metadata_failed_steps": metadata_failed_steps,
                "repaired_steps": self.repaired_steps,
                "repaired_sections": self.repaired_sections,
                "incomplete_sections": self.incomplete_sections,
                "summary": self._process_summary(failed_steps, review),
                "consistency_review": review,
                "human_artifact_status": validation["human_artifact_status"],
                "structured_metadata_status": metadata_status,
                "final_validation_status": validation["final_validation_status"],
                "missing_sections": validation["missing_sections"],
                "truncation_warnings": validation["truncation_warnings"],
                "metadata_available": metadata_available,
                "structured_metadata_reason": outline.get("_structured_metadata_reason", ""),
                "structured_metadata_error_detail": outline.get("_structured_metadata_error_detail", ""),
                "output_contract": output_contract_summary(self.output_contract),
            },
            "final_validation": validation,
            "consensus_points": outline.get("consensus_points", []) if metadata_available and isinstance(outline.get("consensus_points"), list) else [],
            "open_issues": outline.get("open_issues", []) if metadata_available and isinstance(outline.get("open_issues"), list) else [],
            "risk_register": outline.get("risk_register", []) if metadata_available and isinstance(outline.get("risk_register"), list) else [],
            "decision_log": outline.get("decision_log", []) if metadata_available and isinstance(outline.get("decision_log"), list) else [],
            "recommendations": recommendations,
            "implementation_plan": implementation_items,
            "introduction": outline.get("goal_summary", ""),
            "sections": sections,
            "conclusion": self._conclusion_from_review(review, failed_steps),
        }
        RUN_METADATA["synthesis"] = {
            "synthesis_steps": self.step_results,
            "failed_steps": failed_steps,
            "metadata_failed_steps": metadata_failed_steps,
            "repaired_steps": self.repaired_steps,
            "repaired_sections": self.repaired_sections,
            "incomplete_sections": self.incomplete_sections,
            "human_artifact_status": validation["human_artifact_status"],
            "structured_metadata_status": metadata_status,
            "final_validation_status": validation["final_validation_status"],
            "missing_sections": validation["missing_sections"],
            "weak_sections": validation.get("weak_sections", []),
            "contract_section_issues": validation.get("section_issues", []),
            "truncation_warnings": validation["truncation_warnings"],
            "metadata_available": metadata_available,
            "structured_metadata_reason": outline.get("_structured_metadata_reason", ""),
            "structured_metadata_error_detail": outline.get("_structured_metadata_error_detail", ""),
            "output_contract": output_contract_summary(self.output_contract),
        }
        return synthesis

    def _assemble_contract(self, outline: dict, deliverable_results: dict,
                           review: dict) -> dict:
        today = datetime.now().strftime("%Y-%m-%d")
        metadata_step_names = {"structured_metadata_extraction"}
        failed_steps = _effective_human_failed_steps(
            self.step_results, self.repaired_steps, metadata_step_names
        )
        metadata_failed_steps = [
            s["step"] for s in self.step_results
            if not s.get("ok") and s.get("step") in metadata_step_names
        ]
        title = self.output_contract.get("title") or outline.get("title") or f"SzintĂ©zis - {today}"
        verdict = outline.get("verdict") if not outline.get("_error") else ""
        metadata_status = outline.get("_structured_metadata_status", "failed")
        metadata_available = bool(outline.get("_metadata_available")) and metadata_status == "ok"
        recommendations = outline.get("recommendations", []) if metadata_available and isinstance(outline.get("recommendations"), list) else []
        final_markdown = self._build_contract_markdown_from_results(outline, deliverable_results)
        validation = self._validate_contract_markdown(final_markdown)
        human_artifact_status = validation["human_artifact_status"]
        final_status = self._final_validation_status(human_artifact_status, metadata_status, failed_steps)
        validation["structured_metadata_status"] = metadata_status
        validation["final_validation_status"] = final_status
        if failed_steps:
            validation["status"] = "failed"
            validation["truncation_warnings"].append({
                "section": "synthesis_steps",
                "warning": "required_step_failed: " + ", ".join(failed_steps),
            })
            validation["human_artifact_status"] = "failed"
            validation["final_validation_status"] = "failed"
        contract_deliverables = [
            {
                "id": r.get("id"),
                "title": r.get("title"),
                "source": r.get("source"),
                "required": r.get("required", True),
                "sections": r.get("sections", []),
                "markdown": r.get("markdown", ""),
                "chars": len(r.get("markdown", "")),
            }
            for r in deliverable_results.values()
        ]
        synthesis_sections = {
            r.get("id"): r.get("markdown", "")
            for r in contract_deliverables
        }
        section_records = [
            {"heading": r.get("title"), "content": r.get("markdown", ""), "bullets": []}
            for r in contract_deliverables
        ]
        incomplete_sections = self.incomplete_sections + validation.get("incomplete_sections", [])
        process = {
            "steps": self.step_results,
            "failed_steps": failed_steps,
            "metadata_failed_steps": metadata_failed_steps,
            "repaired_steps": self.repaired_steps,
            "repaired_sections": self.repaired_sections,
            "incomplete_sections": incomplete_sections,
            "summary": self._process_summary(failed_steps, review),
            "consistency_review": review,
            "human_artifact_status": validation["human_artifact_status"],
            "structured_metadata_status": metadata_status,
            "final_validation_status": validation["final_validation_status"],
            "missing_sections": validation["missing_sections"],
            "weak_sections": validation.get("weak_sections", []),
            "contract_section_issues": validation.get("section_issues", []),
            "validation_reasons": validation.get("validation_reasons", []),
            "truncation_warnings": validation["truncation_warnings"],
            "metadata_available": metadata_available,
            "structured_metadata_reason": outline.get("_structured_metadata_reason", ""),
            "structured_metadata_error_detail": outline.get("_structured_metadata_error_detail", ""),
            "output_contract": output_contract_summary(self.output_contract),
            "contract_deliverables": contract_deliverables,
        }
        synthesis = {
            "title": title,
            "date": outline.get("date", today) if not outline.get("_error") else today,
            "goal_summary": outline.get("goal_summary", self.user_prompt[:300] if self.user_prompt else ""),
            "participants": [self.engines[k].label for k in self.participant_roles if k in self.engines],
            "judge": self.judge.label,
            "task_profile": self.task_profile,
            "output_contract": output_contract_summary(self.output_contract),
            "verdict": verdict,
            "final_markdown": final_markdown,
            "main_deliverable_markdown": self._join_contract_source(deliverable_results, "main"),
            "implementation_plan_markdown": self._join_contract_source(deliverable_results, "implementation"),
            "ai_context_block": self._join_contract_source(deliverable_results, "ai_context"),
            "md_context_block": self._join_contract_source(deliverable_results, "ai_context"),
            "synthesis_sections": synthesis_sections,
            "synthesis_process": process,
            "final_validation": validation,
            "consensus_points": outline.get("consensus_points", []) if metadata_available and isinstance(outline.get("consensus_points"), list) else [],
            "open_issues": outline.get("open_issues", []) if metadata_available and isinstance(outline.get("open_issues"), list) else [],
            "risk_register": outline.get("risk_register", []) if metadata_available and isinstance(outline.get("risk_register"), list) else [],
            "decision_log": outline.get("decision_log", []) if metadata_available and isinstance(outline.get("decision_log"), list) else [],
            "recommendations": recommendations,
            "implementation_plan": [],
            "introduction": outline.get("goal_summary", ""),
            "sections": section_records,
            "conclusion": self._conclusion_from_review(review, failed_steps),
        }
        RUN_METADATA["synthesis"] = {
            "synthesis_steps": self.step_results,
            "failed_steps": failed_steps,
            "metadata_failed_steps": metadata_failed_steps,
            "repaired_steps": self.repaired_steps,
            "repaired_sections": self.repaired_sections,
            "incomplete_sections": incomplete_sections,
            "human_artifact_status": validation["human_artifact_status"],
            "structured_metadata_status": metadata_status,
            "final_validation_status": validation["final_validation_status"],
            "missing_sections": validation["missing_sections"],
            "weak_sections": validation.get("weak_sections", []),
            "contract_section_issues": validation.get("section_issues", []),
            "validation_reasons": validation.get("validation_reasons", []),
            "truncation_warnings": validation["truncation_warnings"],
            "metadata_available": metadata_available,
            "structured_metadata_reason": outline.get("_structured_metadata_reason", ""),
            "structured_metadata_error_detail": outline.get("_structured_metadata_error_detail", ""),
            "output_contract": output_contract_summary(self.output_contract),
            "contract_deliverables": contract_deliverables,
        }
        return synthesis

    def _build_contract_markdown_from_results(self, outline: dict,
                                              deliverable_results: dict) -> str:
        lines = [
            f"# {self.output_contract.get('title') or 'SzintĂ©zis'}",
            "",
            f"> {outline.get('date', datetime.now().strftime('%Y-%m-%d'))} | Judge: {self.judge.label} | Contract: {self.output_contract.get('contract_id')}",
            "",
        ]
        for deliverable in self.output_contract.get("deliverables", []):
            did = deliverable.get("id")
            result = deliverable_results.get(did, {})
            title = deliverable.get("title") or result.get("title") or did
            content = self._clean_contract_deliverable_markdown(
                result.get("markdown", ""),
                deliverable,
            )
            content = self._filter_contract_deliverable_content(deliverable, content)
            lines += [f"## {title}", ""]
            lines += [content or "_Ez a szakasz nem kĂ©szĂĽlt el teljesen._", ""]
        return "\n".join(lines).strip() + "\n"

    def _filter_contract_deliverable_content(self, deliverable: dict, content: str) -> str:
        sections = deliverable.get("sections") or []
        if not sections:
            return content.strip()
        lines = []
        for section in sections:
            title = section.get("title") or section.get("id")
            section_content = self._contract_section_content(content, title)
            lines += [
                f"### {title}",
                "",
                section_content.strip() or "_Ez a szakasz nem kĂ©szĂĽlt el teljesen._",
                "",
            ]
        return "\n".join(lines).strip()

    def _build_final_markdown(self, outline: dict, verdict: str, main_md: str,
                              impl_md: str, ai_ctx: str, review: dict) -> str:
        metadata_status = outline.get("_structured_metadata_status", "failed")
        metadata_available = bool(outline.get("_metadata_available")) and metadata_status == "ok"
        if not self._uses_default_contract():
            return self._build_contract_markdown(
                outline, verdict, main_md, impl_md, ai_ctx, metadata_available
            )
        lines = [
            "# SzintĂ©zis",
            "",
            f"> {outline.get('date', datetime.now().strftime('%Y-%m-%d'))} | Judge: {self.judge.label} | RĂ©sztvevĹ‘k: {', '.join([self.engines[k].label for k in self.participant_roles if k in self.engines])}",
            "",
            "## VĂ©gsĹ‘ Ă­tĂ©let",
            "",
            verdict.strip() or "A vĂ©gsĹ‘ Ă­tĂ©let nem volt elĂ©rhetĹ‘.",
            "",
            "## VĂ©gsĹ‘ ĂĽzleti terv",
            "",
            main_md.strip() or "_Ez a szakasz nem kĂ©szĂĽlt el teljesen._",
            "",
            "## VĂ©gsĹ‘ megvalĂłsĂ­tĂˇsi terv",
            "",
            impl_md.strip() or "_Ez a szakasz nem kĂ©szĂĽlt el teljesen._",
            "",
            "## AI kontextus blokk",
            "",
            ai_ctx.strip() or "_Ez a szakasz nem kĂ©szĂĽlt el teljesen._",
        ]
        if not metadata_available:
            lines += [
                "",
                "> A strukturĂˇlt dĂ¶ntĂ©si metaadatok nem kĂ©szĂĽltek el; a relevĂˇns kockĂˇzatok Ă©s dĂ¶ntĂ©sek a fenti ĂĽzleti Ă©s megvalĂłsĂ­tĂˇsi tervben szerepelnek.",
            ]
        return "\n".join(lines).strip() + "\n"

    def _build_contract_markdown(self, outline: dict, verdict: str, main_md: str,
                                 impl_md: str, ai_ctx: str, metadata_available: bool) -> str:
        lines = [
            f"# {self.output_contract.get('title') or 'SzintĂ©zis'}",
            "",
            f"> {outline.get('date', datetime.now().strftime('%Y-%m-%d'))} | Judge: {self.judge.label} | Contract: {self.output_contract.get('contract_id')}",
            "",
        ]
        sources = {
            "verdict": verdict.strip() or "A vĂ©gsĹ‘ Ă­tĂ©let nem volt elĂ©rhetĹ‘.",
            "main": main_md.strip(),
            "implementation": impl_md.strip(),
            "ai_context": ai_ctx.strip(),
        }
        for deliverable in self.output_contract.get("deliverables", []):
            title = deliverable.get("title") or deliverable.get("id")
            source_key = deliverable.get("source") or "main"
            content = sources.get(source_key, main_md).strip()
            lines += [f"## {title}", ""]
            sections = deliverable.get("sections") or []
            if sections and source_key != "verdict":
                for idx, section in enumerate(sections):
                    section_title = section.get("title") or section.get("id")
                    section_content = self._contract_section_content(content, section_title)
                    if not section_content and idx == 0:
                        section_content = content
                    lines += [
                        f"### {section_title}",
                        "",
                        section_content.strip() or "_Ez a szakasz nem kĂ©szĂĽlt el teljesen._",
                        "",
                    ]
            else:
                lines += [content or "_Ez a szakasz nem kĂ©szĂĽlt el teljesen._", ""]
        if not metadata_available and self.output_contract.get("required_metadata"):
            lines += [
                "> A strukturĂˇlt dĂ¶ntĂ©si metaadatok nem kĂ©szĂĽltek el; a relevĂˇns kockĂˇzatok Ă©s dĂ¶ntĂ©sek a fenti szakaszokban szerepelnek.",
                "",
            ]
        return "\n".join(lines).strip() + "\n"

    def _contract_section_content(self, markdown: str, heading_text: str) -> str:
        content = self._extract_markdown_section(markdown, heading_text)
        if content:
            return content
        pattern = re.compile(rf"^#+\s+{re.escape(heading_text)}\s*$", re.IGNORECASE | re.MULTILINE)
        match = pattern.search(markdown or "")
        if not match:
            return ""
        tail = markdown[match.end():]
        next_heading = re.search(r"^#+\s+", tail, re.MULTILINE)
        return tail[:next_heading.start()].strip() if next_heading else tail.strip()

    def _has_markdown_heading(self, markdown: str, heading_text: str) -> bool:
        pattern = re.compile(rf"^#+\s+{re.escape(heading_text)}\s*$", re.IGNORECASE | re.MULTILINE)
        return bool(pattern.search(markdown or ""))

    def _is_placeholder_text(self, text: str) -> bool:
        cleaned = " ".join((text or "").strip().lower().split())
        if not cleaned:
            return True
        placeholders = (
            "ez a szakasz nem kĂ©szĂĽlt el teljesen",
            "this section was not completed",
            "this section is incomplete",
            "tbd",
            "todo",
        )
        return any(p in cleaned for p in placeholders)

    def _meaningful_blocks(self, text: str) -> list[str]:
        blocks = []
        current = []
        for raw in (text or "").splitlines():
            line = raw.strip()
            if not line:
                if current:
                    blocks.append(" ".join(current).strip())
                    current = []
                continue
            if line.startswith("#"):
                continue
            if line.startswith(("- ", "* ")) or re.match(r"^\d+\.\s+", line):
                item = re.sub(r"^([-*]|\d+\.)\s+", "", line).strip()
                if len(item.split()) >= 4 and not self._is_placeholder_text(item):
                    blocks.append(item)
                continue
            current.append(line)
        if current:
            blocks.append(" ".join(current).strip())
        return [
            block for block in blocks
            if len(block.split()) >= 6 and not self._is_placeholder_text(block)
        ]

    def _contract_section_status(self, content: str, heading_present: bool,
                                 whole_deliverable: bool = False) -> dict:
        if not heading_present:
            return {"status": "missing", "level": "missing_section", "reasons": ["missing_heading"]}
        reasons = []
        if self._is_placeholder_text(content):
            reasons.append("placeholder_or_empty")
        meaningful = self._meaningful_blocks(content)
        if len(meaningful) < 2 and not (
            whole_deliverable and self._whole_deliverable_is_meaningful(content, meaningful)
        ):
            reasons.append("too_few_meaningful_blocks")
        for warning in self._truncation_warnings(content):
            if warning not in ("empty", "empty_or_failed"):
                reasons.append(warning)
        if reasons:
            return {"status": "weak", "level": "weak_section", "reasons": sorted(set(reasons))}
        return {"status": "valid", "level": "valid", "reasons": []}

    def _whole_deliverable_is_meaningful(self, content: str, meaningful: list[str]) -> bool:
        if _ends_with_closed_fenced_code_block(content):
            inner = re.sub(r"^\s*```[^\n]*\n", "", content.strip())
            inner = re.sub(r"\n```\s*$", "", inner)
            if len(inner.split()) >= 10 or len(inner.strip()) >= 120:
                return True
        plain = re.sub(r"```.*?```", " ", content or "", flags=re.DOTALL)
        plain = re.sub(r"^#+\s+", "", plain, flags=re.MULTILINE)
        plain = re.sub(r"[*_`>\[\]#|-]", " ", plain)
        plain = " ".join(plain.split())
        if len(plain) >= 500:
            return True
        if len(meaningful) >= 2:
            return True
        lower = plain.lower()
        has_summary = any(term in lower for term in (
            "overall assessment", "summary", "verdict", "assessment",
            "decision", "recommendation", "vĂ©gsĹ‘ Ă­tĂ©let", "Ă­tĂ©let", "Ă¶sszegzĂ©s",
        ))
        bullet_count = sum(
            1 for line in (content or "").splitlines()
            if line.strip().startswith(("- ", "* ")) and len(line.strip()[2:].split()) >= 4
        )
        return has_summary and bullet_count >= 1

    def _validate_final_markdown(self, markdown: str) -> dict:
        if not self._uses_default_contract():
            return self._validate_contract_markdown(markdown)
        required = {
            "VĂ©gsĹ‘ Ă­tĂ©let": "verdict",
            "VĂ©gsĹ‘ ĂĽzleti terv": "business_plan",
            "VĂ©gsĹ‘ megvalĂłsĂ­tĂˇsi terv": "implementation_plan",
            "AI kontextus blokk": "ai_context",
        }
        missing = []
        trunc_warnings = []
        for heading, key in required.items():
            content = self._extract_markdown_section(markdown, heading)
            if len(content.strip()) < 20:
                missing.append(key)
            for warning in self._truncation_warnings(content):
                trunc_warnings.append({"section": key, "warning": warning})
        top_level_count = sum(1 for line in _lines_outside_code_fences(markdown.splitlines()) if line.startswith("# "))
        if top_level_count != 1:
            trunc_warnings.append({"section": "document", "warning": f"top_level_heading_count={top_level_count}"})
        human_status = "ok"
        if missing:
            human_status = "failed"
        elif trunc_warnings:
            human_status = "ok_with_warnings"
        return {
            "status": human_status,
            "human_artifact_status": human_status,
            "missing_sections": missing,
            "truncation_warnings": trunc_warnings,
        }

    def _validate_contract_markdown(self, markdown: str) -> dict:
        missing = []
        weak = []
        incomplete = []
        section_issues = []
        trunc_warnings = []
        validation_reasons = []
        for deliverable in self.output_contract.get("deliverables", []):
            title = deliverable.get("title", "")
            did = deliverable.get("id", title)
            content = self._extract_markdown_section(markdown, title)
            deliverable_status = self._contract_section_status(
                content,
                self._has_markdown_heading(markdown, title),
                whole_deliverable=not bool(deliverable.get("sections")),
            )
            if deliverable.get("required", True):
                if deliverable_status["status"] == "missing":
                    missing.append(did)
                    issue = {
                        "deliverable_id": did,
                        "deliverable_title": title,
                        "section_id": "",
                        "section_title": title,
                        "level": "missing_deliverable",
                        "reasons": deliverable_status["reasons"],
                    }
                    section_issues.append(issue)
                    validation_reasons.append(issue)
                elif deliverable_status["status"] == "weak" and not deliverable.get("sections"):
                    weak.append(did)
                    incomplete.append(did)
                    issue = {
                        "deliverable_id": did,
                        "deliverable_title": title,
                        "section_id": "",
                        "section_title": title,
                        "level": "weak_deliverable",
                        "reasons": deliverable_status["reasons"],
                    }
                    section_issues.append(issue)
                    validation_reasons.append(issue)
            for warning in self._truncation_warnings(content):
                if warning not in ("empty", "empty_or_failed"):
                    item = {"section": did, "warning": warning}
                    trunc_warnings.append(item)
                    validation_reasons.append(item)
            for section in deliverable.get("sections", []):
                if not section.get("required", True):
                    continue
                sid = section.get("id", section.get("title", ""))
                section_title = section.get("title", "")
                section_content = self._extract_markdown_section(markdown, section_title)
                status = self._contract_section_status(
                    section_content,
                    self._has_markdown_heading(markdown, section_title),
                )
                if status["status"] == "missing":
                    missing.append(sid)
                    issue = {
                        "deliverable_id": did,
                        "deliverable_title": title,
                        "section_id": sid,
                        "section_title": section_title,
                        "level": "missing_section",
                        "reasons": status["reasons"],
                    }
                    section_issues.append(issue)
                    validation_reasons.append(issue)
                elif status["status"] == "weak":
                    weak.append(sid)
                    incomplete.append(sid)
                    issue = {
                        "deliverable_id": did,
                        "deliverable_title": title,
                        "section_id": sid,
                        "section_title": section_title,
                        "level": "weak_section",
                        "reasons": status["reasons"],
                    }
                    section_issues.append(issue)
                    validation_reasons.append(issue)
                for warning in self._truncation_warnings(section_content):
                    if warning not in ("empty", "empty_or_failed"):
                        item = {"section": sid, "warning": warning}
                        trunc_warnings.append(item)
                        validation_reasons.append(item)
        top_level_count = sum(1 for line in _lines_outside_code_fences(markdown.splitlines()) if line.startswith("# "))
        if top_level_count != 1:
            item = {"section": "document", "warning": f"top_level_heading_count={top_level_count}"}
            trunc_warnings.append(item)
            validation_reasons.append(item)
        human_status = "ok"
        if missing or weak:
            human_status = "failed"
        elif trunc_warnings:
            human_status = "ok_with_warnings"
        return {
            "status": human_status,
            "human_artifact_status": human_status,
            "missing_sections": missing,
            "weak_sections": weak,
            "incomplete_sections": incomplete,
            "section_issues": section_issues,
            "validation_reasons": validation_reasons,
            "truncation_warnings": trunc_warnings,
        }

    def _final_validation_status(self, human_status: str, metadata_status: str,
                                 failed_steps: list[str]) -> str:
        if failed_steps or human_status == "failed":
            return "failed"
        if human_status == "ok_with_warnings" or metadata_status != "ok":
            return "ok_with_warnings"
        return "ok"

    def _repair_sections_before_assembly(self, main_md: str, impl_md: str,
                                         ai_ctx: str) -> tuple[str, str, str]:
        sections = [
            ("business_plan", main_md, "VĂ©gsĹ‘ ĂĽzleti terv", SYNTHESIS_MAX_OUTPUT_TOKENS),
            ("implementation_plan", impl_md, "VĂ©gsĹ‘ megvalĂłsĂ­tĂˇsi terv", SYNTHESIS_MAX_OUTPUT_TOKENS),
            ("ai_context", ai_ctx, "AI kontextus blokk", min(1800, SYNTHESIS_MAX_OUTPUT_TOKENS)),
        ]
        repaired = {}
        for key, content, title, budget in sections:
            warnings = self._truncation_warnings(content)
            if not warnings:
                repaired[key] = content
                continue
            continuation = self.judge.call(
                self._system("markdown"),
                [{"role": "user", "content": self._continuation_prompt(title, content, warnings)}],
                stream=True,
                label=f"Final Judge -> continue {key}",
                max_output_tokens=budget,
            )
            merged = self._merge_continuation(content, continuation)
            remaining = self._truncation_warnings(merged)
            if _valid_response(continuation) and len(remaining) < len(warnings):
                self.repaired_sections.append({
                    "section": key,
                    "warnings_before": warnings,
                    "warnings_after": remaining,
                })
                self.repaired_steps.append(f"{key}_continuation")
                self._log_step(f"{key}_continuation", continuation, "markdown", ok=not remaining,
                               note="section_continuation")
                if remaining:
                    self.incomplete_sections.append({
                        "section": key,
                        "warnings": remaining,
                        "repair_attempted": True,
                    })
                repaired[key] = merged
            else:
                self.incomplete_sections.append({
                    "section": key,
                    "warnings": warnings,
                    "repair_attempted": True,
                })
                self._log_step(f"{key}_continuation", continuation or "", "markdown", ok=False,
                               note="section_continuation_failed")
                repaired[key] = content
        return repaired["business_plan"], repaired["implementation_plan"], repaired["ai_context"]

    def _continuation_prompt(self, section_title: str, content: str, warnings: list[str]) -> str:
        return f"""Continue and complete only this section. Do not repeat previous sections.
Section: {section_title}
Detected issues: {', '.join(warnings)}

Rules:
- Continue from the incomplete point.
- Complete every unfinished sentence, bullet, task block, and Definition of Done.
- If this is an implementation plan task, include Objective, Priority, Recommended phase/sprint,
  Dependencies, Concrete steps with at least 2 meaningful bullets, and Definition of Done with
  at least 2 meaningful bullets.
- Do not add a new top-level document title.
- Return markdown only.

Current section text:
{_trunc(content, 3200)}"""

    def _merge_continuation(self, content: str, continuation: str) -> str:
        if not _valid_response(continuation):
            return content
        cont = self._strip_top_heading(continuation).strip()
        if not cont:
            return content
        return content.rstrip() + "\n\n" + cont + "\n"

    def _process_summary(self, failed_steps: list[str], review: dict) -> str:
        ok_steps = [s["step"] for s in self.step_results if s.get("ok")]
        review_state = "passed" if review.get("passed") else "completed with warnings"
        if review.get("_error"):
            review_state = "review failed"
        return (
            f"Synthesis ran as {len(self.step_results)} internal steps. "
            f"Successful steps: {', '.join(ok_steps) or 'none'}. "
            f"Failed steps: {', '.join(failed_steps) if failed_steps else 'none'}. "
            f"Consistency review: {review_state}."
        )

    def _conclusion_from_review(self, review: dict, failed_steps: list[str]) -> str:
        warnings = review.get("final_warnings", []) if isinstance(review, dict) else []
        if failed_steps:
            return "A szintĂ©zis elkĂ©szĂĽlt, de rĂ©szleges: " + ", ".join(failed_steps)
        if warnings:
            return "A szintĂ©zis elkĂ©szĂĽlt. Fontos figyelmeztetĂ©sek: " + "; ".join(warnings)
        return "A szintĂ©zis tĂ¶bb lĂ©pĂ©sben elkĂ©szĂĽlt, a konzisztencia-ellenĹ‘rzĂ©s nem jelzett kritikus hibĂˇt."


def phase_judge(judge: AIEngine, evidence: str, issue_matrix: str,
                rebuttals: dict, revisions: dict, engines: dict,
                roles_info: dict, user_prompt: str, output_types: list,
                lang: str, session_log: list,
                task_profile: TaskProfile | dict | None = None,
                output_contract: OutputContract | dict | None = None) -> dict:
    """
    Phase 5: Final Judge
    KĂĽlĂ¶n modell Ă­rja a vĂ©gsĹ‘ szintĂ©zist â€” nem a moderĂˇtor.
    """
    print("\n" + "â•"*64)
    print(_bold(f"âš–ď¸Ź  PHASE 5 â€” Final Judge ({judge.label})"))
    print("â•"*64)

    return SynthesisEngine(
        judge, evidence, issue_matrix, rebuttals, revisions, engines,
        roles_info, user_prompt, output_types, lang, session_log,
        task_profile, output_contract
    ).run()

    lang_i = "MAGYARUL." if lang=="hu" else "In ENGLISH."
    outputs = "\n".join(f"- {o}" for o in output_types)
    rebuttals = {
        rk: text for rk, text in rebuttals.items()
        if _valid_response(text)
    }
    revisions = {
        rk: rv for rk, rv in revisions.items()
        if rk in rebuttals and isinstance(rv, dict)
    }
    participant_roles = sorted(set(rebuttals.keys()) | set(revisions.keys()))
    _require_min_participants(participant_roles, "Final judge input")

    # Revision Ă¶sszefoglalĂł
    rev_summary = "\n".join(
        f"[{r.get('model','?')} | {r.get('role','?')}] "
        f"VĂˇltoztatott: {r.get('changed_my_mind')} | "
        f"Bizalom: {r.get('confidence',0):.0%} | "
        f"Javaslat: {r.get('final_recommendation','')[:100]}"
        for r in revisions.values()
    )
    rebuttals_text = "\n\n".join(
        f"[{engines[rk].label}]\n{text}"
        for rk, text in rebuttals.items()
    )

    today = datetime.now().strftime('%Y-%m-%d')
    participants_list = json.dumps(
        [engines[k].label for k in participant_roles if k in engines],
        ensure_ascii=False
    )

    system = f"""Te egy fĂĽggetlen Final Judge AI vagy. A vitĂˇban NEM vettĂ©l rĂ©szt.
{f'EREDETI CĂ‰L:\n"""{user_prompt}"""\n' if user_prompt else ""}
{lang_i}

FELADATOD: A vita teljes anyaga alapjĂˇn kĂ©szĂ­tsd el az alĂˇbbi HĂROM dokumentumot
RĂ‰SZLETESEN Ă©s KONKRĂ‰TAN â€” ne ĂˇltalĂˇnosĂ­ts, hivatkozz a vitĂˇban elhangzott Ă©rvekre.

A kimenet KIZĂRĂ“LAG vĂˇlid JSON legyen."""

    prompt = f"""EVIDENCE PACKAGE:
{_trunc(evidence, 2500)}

ISSUE MATRIX (moderĂˇtor â€” konszenzusok Ă©s vitapontok):
{_trunc(issue_matrix, 2000)}

REBUTTAL KĂ–RĂ–K (ki mit mondott):
{_trunc(rebuttals_text, 2500)}

ĂLLĂSPONT-FRISSĂŤTĂ‰SEK (ki vĂˇltoztatott vĂ©lemĂ©nyt):
{rev_summary}

---
KĂ©szĂ­tsd el a kĂ¶vetkezĹ‘ hĂˇrom dokumentumot JSON formĂˇtumban.
Minden mezĹ‘ legyen RĂ‰SZLETES Ă©s KONKRĂ‰T â€” legalĂˇbb 150-300 szĂł fejeztenkĂ©nt.

CSAK JSON, semmi mĂˇs:
{{
  "title": "Projekt neve â€” VĂ©gsĹ‘ DokumentĂˇciĂł",
  "date": "{today}",
  "goal_summary": "A felhasznĂˇlĂł cĂ©ljĂˇnak 2-3 mondatos Ă¶sszefoglalĂˇsa",
  "participants": {participants_list},
  "judge": "{judge.label}",
  "verdict": "A vĂ©gsĹ‘ Ă­tĂ©let 3-4 mondatban â€” konkrĂ©t dĂ¶ntĂ©sek, nem ĂˇltalĂˇnossĂˇgok",

  "uzleti_terv": {{
    "piac_es_lehetoseg": "RĂ©szletes piaci helyzet: mĂ©ret, timing, miĂ©rt most, ki a cĂ©lkĂ¶zĂ¶nsĂ©g â€” min. 200 szĂł",
    "termek_es_ertekajanlat": "Mit csinĂˇl a termĂ©k pontosan, mit old meg, kinek, mi az egyedi Ă©rtĂ©ke a KĂ“DBĂZIS valĂłsĂˇgĂˇra alapozva â€” min. 200 szĂł",
    "jelenlegi_allapot_korlatok": "Ami ma tĂ©nyleg mĹ±kĂ¶dik Ă©s ami hiĂˇnyzik â€” Ĺ‘szintĂ©n, a kĂłd alapjĂˇn â€” min. 150 szĂł",
    "bevĂ©teli_modell": "Freemium/Pro/Team/Enterprise rĂ©szletesen: mit tartalmaz minden tier, Ăˇrak, miĂ©rt ez a struktĂşra â€” min. 150 szĂł",
    "gtm_strategia": "Go-to-market: elsĹ‘ 30/60/90 nap konkrĂ©t lĂ©pĂ©sei, cĂ©lszegmensek, csatornĂˇk, elsĹ‘ bevĂ©tel hogyan â€” min. 200 szĂł",
    "versenyelonyok": "Mi kĂĽlĂ¶nbĂ¶ztet meg, miĂ©rt nehĂ©z mĂˇsolni, jĂ¶vĹ‘beli moat â€” min. 150 szĂł",
    "versenytarsak": "Ki a versenytĂˇrs, hol van a rĂ©s, integrĂˇciĂł vs verseny a nagy platformokkal â€” min. 150 szĂł",
    "kockazatok": ["kockĂˇzat: mitigĂˇciĂł", "kockĂˇzat: mitigĂˇciĂł"],
    "sikerkritĂ©riumok": ["KPI 1 â€” mĂ©rĂ©si mĂłd", "KPI 2 â€” mĂ©rĂ©si mĂłd"]
  }},

  "megvalositasi_terv": {{
    "jelenlegi_architektura": "Ami van Ă©s Ă©rtĂ©kelĂ©se â€” konkrĂ©t fĂˇjlokra hivatkozva â€” min. 150 szĂł",
    "cel_architektura": "Hogyan nĂ©zzen ki a vĂ©gsĹ‘ rendszer: frontend/backend/AI/DB/auth rĂ©tegek â€” min. 200 szĂł",
    "azonnali_teendok": ["[0-2 hĂ©t] konkrĂ©t feladat", "[0-2 hĂ©t] konkrĂ©t feladat"],
    "sprint_1": {{
      "fĂłkusz": "Backend alapok",
      "idotartam": "2-4 hĂ©t",
      "feladatok": ["feladat 1", "feladat 2", "feladat 3"],
      "definition_of_done": "Mikor kĂ©sz"
    }},
    "sprint_2": {{
      "fĂłkusz": "AI integrĂˇciĂł",
      "idotartam": "4-6 hĂ©t",
      "feladatok": ["feladat 1", "feladat 2"],
      "definition_of_done": "Mikor kĂ©sz"
    }},
    "sprint_3": {{
      "fĂłkusz": "GitHub App Ă©s CLI",
      "idotartam": "6-10 hĂ©t",
      "feladatok": ["feladat 1", "feladat 2"],
      "definition_of_done": "Mikor kĂ©sz"
    }},
    "sprint_4_plus": {{
      "fĂłkusz": "Team Ă©s Enterprise",
      "idotartam": "10+ hĂ©t",
      "feladatok": ["feladat 1", "feladat 2"],
      "definition_of_done": "Mikor kĂ©sz"
    }},
    "technikai_kockazatok": [{{"kockĂˇzat": "...", "valĂłszĂ­nĹ±sĂ©g": "magas/kĂ¶zepes/alacsony", "mitigĂˇciĂł": "..."}}]
  }},

  "ai_context_block": "# Projekt neve â€” AI Kontextus\n\n## Mi ez\n[1-2 mondat]\n\n## Jelenlegi Ăˇllapot\n- [bullet]\n\n## Ăśzleti cĂ©lok\n- [bullet]\n\n## Kritikus blokkolĂłk\n1. [prioritĂˇs szerint]\n\n## ArchitektĂşra\n[rĂ¶vid leĂ­rĂˇs]\n\n## Ha ezt a kontextust kapod\n[mit vĂˇrunk az AI-tĂłl ha ezt a promptot csatoljĂˇk]",

  "consensus_points": ["Konszenzus pont 1 â€” rĂ©szletesen", "Konszenzus pont 2"],
  "open_issues": ["Nyitott kĂ©rdĂ©s 1 â€” miĂ©rt nem dĹ‘lt el", "Nyitott kĂ©rdĂ©s 2"],
  "risk_register": [{{"risk": "...", "severity": "high/medium/low", "mitigation": "..."}}],
  "decision_log": [{{"decision": "...", "rationale": "...", "confidence": 0.0}}],
  "recommendations": ["AjĂˇnlĂˇs 1 â€” rĂ©szletesen miĂ©rt", "AjĂˇnlĂˇs 2"],
  "implementation_plan": [{{"step": 1, "action": "...", "owner": "...", "priority": "high/medium/low"}}],
  "introduction": "BevezetĹ‘ â€” a vita Ă¶sszefoglalĂˇsa Ă©s cĂ©lja â€” min. 150 szĂł",
  "sections": [{{"heading": "Fejezet", "content": "Tartalom â€” min. 150 szĂł", "bullets": ["pont"]}}],
  "conclusion": "Ă–sszefoglalĂˇs â€” min. 150 szĂł",
  "md_context_block": "RĂ¶vid AI-ready Ă¶sszefoglalĂł (ezt hagyhatod ĂĽresen, az ai_context_block mezĹ‘t hasznĂˇljuk)"
}}"""

    raw = judge.call(system, [{"role":"user","content":prompt}],
                     stream=True, label=f"Final Judge â†’ {judge.label}",
                     max_output_tokens=SYNTHESIS_MAX_OUTPUT_TOKENS)
    if _is_failed_response(raw):
        session_log.append({
            "phase": "judge_error", "role": "judge",
            "speaker": judge.label, "text": raw, "error": True
        })
        _abort(f"Final judge call failed for {judge.label} ({judge.mid}): {raw}")
    session_log.append({
        "phase": "judge", "role": "judge",
        "speaker": judge.label, "text": raw
    })

    clean = raw.strip()
    if "```" in clean:
        for part in clean.split("```"):
            p = part.strip().lstrip("json").strip()
            if p.startswith("{"): clean = p; break
    try:
        return json.loads(clean)
    except:
        return {
            "title": f"SzintĂ©zis â€” {datetime.now().strftime('%Y-%m-%d')}",
            "date": datetime.now().strftime("%Y-%m-%d"),
            "goal_summary": user_prompt[:200] if user_prompt else "",
            "participants": [engines[k].label for k in participant_roles if k in engines],
            "judge": judge.label, "verdict": "LĂˇsd vita log.",
            "consensus_points": [], "open_issues": [],
            "risk_register": [], "decision_log": [],
            "recommendations": [],
            "implementation_plan": [],
            "introduction": "Vita eredmĂ©nye.",
            "sections": [{"heading":"Ă–sszefoglalĂł","content":raw[:2000],"bullets":[]}],
            "conclusion": "LĂˇsd vita rĂ©szleteit.",
            "md_context_block": f"# SzintĂ©zis\n\n{raw[:1500]}"
        }


# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Quick scenario (2-modelles vita vĂˇltozatlan logikĂˇval)
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def run_quick(engines: dict, items: list, user_prompt: str,
              lang: str, autosave_dir: str,
              task_profile: TaskProfile | dict | None = None,
              output_contract: OutputContract | dict | None = None) -> tuple[list, dict]:
    """Quick: evidence â†’ independent â†’ judge (1 kĂ¶r, 3 modell)."""
    session_log = []
    moderator = engines.get("moderator")
    judge     = engines.get("judge")
    debaters  = {k:v for k,v in engines.items() if k not in ("moderator","judge")}
    roles_info = {"debater1": "Strategist / ĂĽzleti szempontok",
                  "debater2": "Technical / megvalĂłsĂ­thatĂłsĂˇg"}

    print("\n" + "â•"*64)
    print(_bold("âšˇ QUICK COUNCIL â€” 3 modell, 1 kĂ¶r"))
    print("â•"*64)

    evidence = phase_evidence(items, user_prompt, moderator, task_profile)
    if _is_failed_response(evidence):
        _abort(f"Evidence generation failed for moderator {moderator.label}: {evidence}")
    session_log.append({"phase":"evidence","role":"moderator",
                         "speaker":moderator.label,"text":evidence})
    _autosave(session_log, autosave_dir, user_prompt)

    opinions = phase_independent(engines, evidence, roles_info,
                                  parallel=False, session_log=session_log,
                                  task_profile=task_profile)
    _require_min_participants(list(opinions.keys()), "Independent opinion phase")
    _autosave(session_log, autosave_dir, user_prompt)

    output_types = _detect_outputs(user_prompt)
    synthesis = phase_judge(judge, evidence, "", opinions, {}, engines,
                             roles_info, user_prompt, output_types, lang, session_log,
                             task_profile, output_contract)
    _autosave(session_log, autosave_dir, user_prompt)
    return session_log, synthesis


def run_expert_council(engines: dict, items: list, user_prompt: str,
                        lang: str, parallel: bool, autosave_dir: str,
                        resume_log: list = None,
                        task_profile: TaskProfile | dict | None = None,
                        output_contract: OutputContract | dict | None = None) -> tuple[list, dict]:
    """Full 5-fĂˇzis expert council."""
    session_log = list(resume_log) if resume_log else []
    done_phases = {e.get("phase") for e in session_log}

    moderator = engines.get("moderator")
    judge     = engines.get("judge")
    roles_info = {
        # expert-council
        "strategist":      "Ăśzleti stratĂ©ga â€” GTM, monetizĂˇciĂł, piac",
        "engineer":        "MĂ©rnĂ¶k / Architect â€” technikai megvalĂłsĂ­thatĂłsĂˇg",
        "market_analyst":  "Piac / VersenytĂˇrs elemzĹ‘ â€” trendek, kĂĽlsĹ‘ nĂ©zet",
        "skeptic":         "Szkeptikus / Red Team â€” kockĂˇzatok, gyenge pontok",
        "cost_reasoner":   "KĂ¶ltsĂ©g Ă©s erĹ‘forrĂˇs elemzĹ‘ â€” ROI, prioritizĂˇlĂˇs",
        # red-team
        "skeptic2":        "MĂˇsodik szkeptikus â€” alternatĂ­v kritika",
        # build-plan
        "product_mgr":     "Product Manager â€” felhasznĂˇlĂłi igĂ©nyek, roadmap priorizĂˇlĂˇs",
        "architect":       "Senior Architect â€” technikai architektĂşra, tech debt",
        "security":        "Security Engineer â€” biztonsĂˇgi kockĂˇzatok, compliance",
        "qa_reviewer":     "QA / Test Reviewer â€” tesztelhetĹ‘sĂ©g, release criteria",
        # quick
        "debater1":        "StratĂ©giai elemzĹ‘",
        "debater2":        "Technikai elemzĹ‘",
    }

    print("\n" + "â•"*64)
    print(_bold("đźŽ™ď¸Ź  EXPERT COUNCIL â€” 5 fĂˇzis"))
    print("â•"*64)

    # Phase 0
    if "evidence" not in done_phases:
        evidence = phase_evidence(items, user_prompt, moderator, task_profile)
        if _is_failed_response(evidence):
            _abort(f"Evidence generation failed for moderator {moderator.label}: {evidence}")
        session_log.append({"phase":"evidence","role":"moderator",
                             "speaker":moderator.label,"text":evidence})
        _autosave(session_log, autosave_dir, user_prompt)
    else:
        evidence = next(e["text"] for e in session_log if e["phase"]=="evidence")
        print(_yellow("   â–¶ď¸Ź  Phase 0 kihagyva (resume)"))

    # Phase 1
    if "independent" not in done_phases:
        opinions = phase_independent(engines, evidence, roles_info,
                                      parallel, session_log, task_profile)
        _autosave(session_log, autosave_dir, user_prompt)
    else:
        opinions = {e["role"]: e["text"] for e in session_log if e["phase"]=="independent"}
        print(_yellow("   â–¶ď¸Ź  Phase 1 kihagyva (resume)"))
    _require_min_participants(list(opinions.keys()), "Independent opinion phase")

    # Phase 2
    if "issue_matrix" not in done_phases:
        issue_matrix = phase_issue_matrix(moderator, opinions, evidence,
                                           engines, session_log)
        _autosave(session_log, autosave_dir, user_prompt)
    else:
        issue_matrix = next(e["text"] for e in session_log if e["phase"]=="issue_matrix")
        print(_yellow("   â–¶ď¸Ź  Phase 2 kihagyva (resume)"))

    # Phase 3
    if "rebuttal" not in done_phases:
        rebuttals = phase_rebuttal(engines, issue_matrix, evidence,
                                    roles_info, session_log,
                                    valid_roles=set(opinions.keys()))
        _autosave(session_log, autosave_dir, user_prompt)
    else:
        rebuttals = {e["role"]: e["text"] for e in session_log if e["phase"]=="rebuttal"}
        print(_yellow("   â–¶ď¸Ź  Phase 3 kihagyva (resume)"))
    _require_min_participants(list(rebuttals.keys()), "Rebuttal phase")

    # Phase 4
    if "revision" not in done_phases:
        revisions = phase_revision(engines, issue_matrix, rebuttals,
                                    roles_info, session_log,
                                    valid_roles=set(rebuttals.keys()))
        _autosave(session_log, autosave_dir, user_prompt)
    else:
        revisions = {}
        for e in session_log:
            if e["phase"]=="revision":
                try: revisions[e["role"]] = json.loads(e["text"])
                except: revisions[e["role"]] = {"model": e["speaker"]}
        print(_yellow("   â–¶ď¸Ź  Phase 4 kihagyva (resume)"))

    # Phase 5
    output_types = _detect_outputs(user_prompt)
    synthesis = phase_judge(judge, evidence, issue_matrix, rebuttals,
                             revisions, engines, roles_info, user_prompt,
                             output_types, lang, session_log, task_profile,
                             output_contract)
    _autosave(session_log, autosave_dir, user_prompt)
    return session_log, synthesis


# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# SegĂ©dfĂĽggvĂ©nyek
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _detect_outputs(user_prompt: str) -> list:
    types = ["SzintĂ©zis dokumentum"]
    if not user_prompt: return types
    pl = user_prompt.lower()
    if "ĂĽzleti terv" in pl:   types.append("VĂ©gsĹ‘ ĂĽzleti terv")
    if "megvalĂłsĂ­tĂˇs" in pl:  types.append("VĂ©gsĹ‘ megvalĂłsĂ­tĂˇsi terv")
    if "markdown" in pl or "ai model" in pl or ".md" in pl:
        types.append("AI-ready markdown kontextus blokk")
    return types

def _autosave(session_log, out_dir, user_prompt):
    if not out_dir: return
    os.makedirs(out_dir, exist_ok=True)
    path = os.path.join(out_dir, "debate_log.json")
    payload = {"user_prompt": user_prompt, "debate": session_log}
    if RUN_METADATA:
        payload["metadata"] = RUN_METADATA
    with open(path,"w",encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    print(_dim(f"   đź’ľ Mentve: {path} ({len(session_log)} bejegyzĂ©s)"))

def _extract_md_section(markdown: str, heading_text: str) -> str:
    lines = (markdown or "").splitlines()
    start = None
    level = 0
    in_fence = False
    for i, line in enumerate(lines):
        if line.strip().startswith("```"):
            in_fence = not in_fence
            continue
        if not in_fence and line.lstrip("#").strip().lower() == heading_text.lower():
            start = i + 1
            level = len(line) - len(line.lstrip("#"))
            break
    if start is None:
        return ""
    end = len(lines)
    in_fence = False
    for j in range(start, len(lines)):
        if lines[j].strip().startswith("```"):
            in_fence = not in_fence
            continue
        if not in_fence and lines[j].startswith("#"):
            next_level = len(lines[j]) - len(lines[j].lstrip("#"))
            if next_level <= level:
                end = j
                break
    return "\n".join(lines[start:end]).strip()

def _lines_outside_code_fences(lines: list[str]) -> list[str]:
    outside = []
    in_fence = False
    for line in lines:
        if line.strip().startswith("```"):
            in_fence = not in_fence
            continue
        if not in_fence:
            outside.append(line)
    return outside

def _fenced_code_blocks_balanced(text: str) -> bool:
    in_fence = False
    for line in (text or "").splitlines():
        if line.strip().startswith("```"):
            in_fence = not in_fence
    return not in_fence

def _ends_with_closed_fenced_code_block(text: str) -> bool:
    lines = [ln.strip() for ln in (text or "").rstrip().splitlines() if ln.strip()]
    return bool(lines and lines[-1].startswith("```") and _fenced_code_blocks_balanced(text))

def _markdown_truncation_warnings(text: str) -> list[str]:
    if not text or _is_failed_response(text):
        return ["empty_or_failed"]
    raw_lines = text.rstrip().splitlines()
    lines = [ln.rstrip() for ln in raw_lines if ln.strip()]
    if not lines:
        return ["empty"]
    last = lines[-1].strip()
    warnings = []
    if not _ends_with_closed_fenced_code_block(text.rstrip()):
        low = last.lower()
        if low in ("- minden", "* minden", "minden", "- all", "* all"):
            warnings.append("suspicious_fragment_minden")
        if "tbd" in low:
            warnings.append("tbd_marker")
        if last.endswith(":"):
            warnings.append("ends_after_colon")
        if last.startswith("#"):
            warnings.append("ends_with_heading")
        if last in ("-", "*") or last.startswith(("- [", "* [")):
            warnings.append("unfinished_bullet")
        if len(last.split()) <= 2 and last.startswith(("- ", "* ")):
            warnings.append("too_short_final_bullet")
        if _looks_mid_word_ending(last):
            warnings.append("possible_mid_word_ending")
    warnings.extend(_abrupt_heading_transitions(_lines_outside_code_fences(raw_lines)))
    warnings.extend(_short_heading_content_warnings(text))
    warnings.extend(_task_block_warnings_md(text))
    return warnings

def _short_heading_content_warnings(markdown: str) -> list[str]:
    lines = _lines_outside_code_fences((markdown or "").splitlines())
    headings = [(i, ln.strip()) for i, ln in enumerate(lines) if ln.strip().startswith("#")]
    warnings = []
    for idx, (start, heading) in enumerate(headings):
        level = len(heading) - len(heading.lstrip("#"))
        if level > 4:
            continue
        end = len(lines)
        for j in range(start + 1, len(lines)):
            stripped = lines[j].strip()
            if stripped.startswith("#"):
                next_level = len(stripped) - len(stripped.lstrip("#"))
                if next_level <= level:
                    end = j
                    break
        content = "\n".join(lines[start + 1:end]).strip()
        if len(content.split()) < 5:
            label = heading.lstrip("#").strip().lower().replace(" ", "_")[:40]
            warnings.append(f"heading_too_little_content:{label}")
    return warnings

def _looks_mid_word_ending(line: str) -> bool:
    stripped = (line or "").strip()
    if not stripped:
        return False
    if stripped.endswith(("...", "â€¦")):
        return True
    if stripped[-1] in ".!?)]}â€ť\"'`":
        return False
    if stripped.startswith(("- ", "* ", "#", "|")):
        return False
    words = stripped.split()
    if not words:
        return False
    last_word = words[-1].strip(".,;:!?)]}â€ť\"'")
    if len(words) <= 3 and len(last_word) >= 4:
        return True
    return any(last_word.lower().endswith(sfx) for sfx in ("eszk", "funk", "valid", "integr", "konfig", "architekt", "implement"))

def _strip_trailing_markdown_emphasis(text: str) -> str:
    cleaned = (text or "").rstrip()
    while cleaned.endswith(("**", "__")):
        cleaned = cleaned[:-2].rstrip()
    while cleaned.endswith(("*", "_", "`")):
        cleaned = cleaned[:-1].rstrip()
    return cleaned

def _abrupt_heading_transitions(lines: list[str]) -> list[str]:
    warnings = []
    prev = ""
    for line in lines:
        cur = line.strip()
        if cur.startswith("#") and prev:
            prev_clean = prev.strip()
            if (
                len(prev_clean) > 24
                and not prev_clean.startswith(("- ", "* ", "|", "**"))
                and not re.match(r"^\d+[\.)]\s+", prev_clean)
                and not _strip_trailing_markdown_emphasis(prev_clean).endswith((".", "!", "?", ":", ";", ")", "]", "â€ť", "\""))
                and not prev_clean.endswith((".", "!", "?", ":", ";", ")", "]", "â€ť", "\""))
            ):
                warnings.append("abrupt_transition_before_heading")
        if cur:
            prev = cur
    return warnings

def _task_blocks_md(markdown: str) -> list[str]:
    lines = _lines_outside_code_fences((markdown or "").splitlines())
    starts = [i for i, ln in enumerate(lines) if ln.strip().lower().startswith("#### task ")]
    blocks = []
    for pos, start in enumerate(starts):
        end = starts[pos + 1] if pos + 1 < len(starts) else len(lines)
        blocks.append("\n".join(lines[start:end]))
    return blocks

def _field_block_md(block: str, marker: str) -> str:
    lines = block.splitlines()
    start = None
    for i, ln in enumerate(lines):
        if ln.strip().lower() == marker.lower():
            start = i + 1
            break
    if start is None:
        return ""
    end = len(lines)
    for j in range(start, len(lines)):
        txt = lines[j].strip()
        if txt.startswith("**") and txt.endswith(":**"):
            end = j
            break
        if txt.startswith("#### Task "):
            end = j
            break
    return "\n".join(lines[start:end])

def _task_block_warnings_md(markdown: str) -> list[str]:
    warnings = []
    for idx, block in enumerate(_task_blocks_md(markdown), 1):
        lower = block.lower()
        required = {
            "objective": "**objective:**",
            "priority": "**priority:**",
            "dependencies": "**dependencies:**",
            "concrete_steps": "**concrete steps:**",
            "definition_of_done": "**definition of done:**",
        }
        for label, marker in required.items():
            if marker not in lower:
                warnings.append(f"task_{idx}_missing_{label}")
        dod = _field_block_md(block, "**Definition of Done:**")
        dod_bullets = [ln for ln in dod.splitlines() if ln.strip().startswith(("- ", "* ")) and len(ln.strip()[2:].split()) >= 3]
        if len(dod_bullets) < 2:
            warnings.append(f"task_{idx}_incomplete_definition_of_done")
        steps = _field_block_md(block, "**Concrete steps:**")
        step_bullets = [ln for ln in steps.splitlines() if ln.strip().startswith(("- ", "* ")) and len(ln.strip()[2:].split()) >= 3]
        if len(step_bullets) < 2:
            warnings.append(f"task_{idx}_incomplete_concrete_steps")
    return warnings

def _validate_synthesis_markdown(markdown: str) -> dict:
    required = {
        "VĂ©gsĹ‘ Ă­tĂ©let": "verdict",
        "VĂ©gsĹ‘ ĂĽzleti terv": "business_plan",
        "VĂ©gsĹ‘ megvalĂłsĂ­tĂˇsi terv": "implementation_plan",
        "AI kontextus blokk": "ai_context",
    }
    missing = []
    trunc_warnings = []
    for heading, key in required.items():
        content = _extract_md_section(markdown, heading)
        if len(content.strip()) < 20:
            missing.append(key)
        for warning in _markdown_truncation_warnings(content):
            trunc_warnings.append({"section": key, "warning": warning})
    top_level_count = sum(1 for line in _lines_outside_code_fences(markdown.splitlines()) if line.startswith("# "))
    if top_level_count != 1:
        trunc_warnings.append({"section": "document", "warning": f"top_level_heading_count={top_level_count}"})
    human_status = "ok"
    if missing:
        human_status = "failed"
    elif trunc_warnings:
        human_status = "ok_with_warnings"
    return {
        "status": human_status,
        "human_artifact_status": human_status,
        "missing_sections": missing,
        "truncation_warnings": trunc_warnings,
    }


# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Markdown + Word output
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def save_markdown(synthesis: dict, items: list, user_prompt: str,
                  session_log: list, revisions: dict, output_path: str):
    """VĂ©gsĹ‘ Ă¶sszefoglalĂł markdown â€” szintĂ©zis, ĂĽzleti terv kivonat, dĂ¶ntĂ©si naplĂł."""
    if synthesis.get("final_markdown"):
        final_md = synthesis["final_markdown"].strip() + "\n"
        validation = dict(synthesis.get("final_validation") or _validate_synthesis_markdown(final_md))
        if (
            validation.get("status") == "failed"
            and not validation.get("missing_sections")
            and not validation.get("weak_sections")
            and not validation.get("incomplete_sections")
            and not validation.get("truncation_warnings")
            and not synthesis.get("synthesis_process", {}).get("failed_steps")
        ):
            validation["status"] = "ok"
            validation["human_artifact_status"] = "ok"
        synthesis["final_validation"] = validation
        proc = synthesis.setdefault("synthesis_process", {})
        proc["failed_steps"] = _normalize_human_failed_steps(
            proc.get("failed_steps", []),
            proc.get("repaired_steps", []),
        )
        metadata_status = proc.get("structured_metadata_status", "unknown")
        if proc.get("failed_steps"):
            validation["status"] = "failed"
            validation["human_artifact_status"] = "failed"
            validation["truncation_warnings"].append({
                "section": "synthesis_steps",
                "warning": "required_step_failed: " + ", ".join(proc.get("failed_steps", [])),
            })
        final_status = validation["status"]
        if final_status != "failed" and metadata_status != "ok":
            final_status = "ok_with_warnings"
        validation["structured_metadata_status"] = metadata_status
        validation["final_validation_status"] = final_status
        proc["human_artifact_status"] = validation.get("human_artifact_status", validation["status"])
        proc["final_validation_status"] = final_status
        proc["missing_sections"] = validation["missing_sections"]
        proc["weak_sections"] = validation.get("weak_sections", [])
        proc["incomplete_sections"] = validation.get("incomplete_sections", proc.get("incomplete_sections", []))
        proc["contract_section_issues"] = validation.get("section_issues", [])
        proc["validation_reasons"] = validation.get("validation_reasons", [])
        proc["truncation_warnings"] = validation["truncation_warnings"]
        RUN_METADATA.setdefault("synthesis", {})
        RUN_METADATA["synthesis"].update({
            "synthesis_steps": proc.get("steps", []),
            "failed_steps": proc.get("failed_steps", []),
            "metadata_failed_steps": proc.get("metadata_failed_steps", []),
            "repaired_steps": proc.get("repaired_steps", []),
            "repaired_sections": proc.get("repaired_sections", []),
            "incomplete_sections": proc.get("incomplete_sections", []),
            "human_artifact_status": proc["human_artifact_status"],
            "structured_metadata_status": metadata_status,
            "final_validation_status": final_status,
            "missing_sections": validation["missing_sections"],
            "weak_sections": validation.get("weak_sections", []),
            "contract_section_issues": validation.get("section_issues", []),
            "validation_reasons": validation.get("validation_reasons", []),
            "truncation_warnings": validation["truncation_warnings"],
            "metadata_available": proc.get("metadata_available", False),
            "structured_metadata_reason": proc.get("structured_metadata_reason", ""),
            "structured_metadata_error_detail": proc.get("structured_metadata_error_detail", ""),
        })
        validation_log = dict(validation)
        validation_log["structured_metadata_status"] = metadata_status
        validation_log["final_validation_status"] = final_status
        session_log.append({
            "phase": "synthesis_validation",
            "role": "judge",
            "speaker": synthesis.get("judge", ""),
            "text": json.dumps(validation_log, ensure_ascii=False),
            "ok": final_status in ("ok", "ok_with_warnings"),
            "error": final_status == "failed",
        })
        if final_status == "failed":
            final_md += (
                "\n## MinĹ‘sĂ©gi megjegyzĂ©s\n\n"
                "A dokumentum rĂ©szben elkĂ©szĂĽlt, de az automatikus ellenĹ‘rzĂ©s hiĂˇnyt vagy befejezetlen szakaszt jelzett. "
                "A rĂ©szletek a meeting reportban Ă©s a debate_log.json fĂˇjlban talĂˇlhatĂłk.\n"
            )
        Path(output_path).write_text(final_md, encoding="utf-8")
        print(f"âś… SzintĂ©zis MD: {output_path}")
        return

    lines = [
        f"# {synthesis.get('title','SzintĂ©zis')}",
        f"> {synthesis.get('date','')} | {synthesis.get('judge','')} Ă­tĂ©lete | "
        f"RĂ©sztvevĹ‘k: {', '.join(synthesis.get('participants',[]))}",
        "", "---", "",
    ]
    if synthesis.get("verdict"):
        lines += ["## âš–ď¸Ź VĂ©gsĹ‘ Ă­tĂ©let", "", f"> {synthesis['verdict']}", ""]
    if synthesis.get("introduction"):
        lines += ["## BevezetĹ‘", "", synthesis["introduction"], ""]

    # â”€â”€ Ăśzleti terv fejezetek â”€â”€
    proc = synthesis.get("synthesis_process", {})
    if proc.get("summary"):
        lines += ["## Synthesis Process", "", proc["summary"], ""]
        failed = proc.get("failed_steps", [])
        if failed:
            lines += ["**Failed synthesis steps:**", ""]
            for step in failed:
                lines.append(f"- {step}")
            lines.append("")

    if synthesis.get("main_deliverable_markdown"):
        lines += ["---", synthesis["main_deliverable_markdown"].strip(), ""]

    if synthesis.get("implementation_plan_markdown"):
        lines += ["---", synthesis["implementation_plan_markdown"].strip(), ""]

    bt = synthesis.get("uzleti_terv", {})
    if bt:
        lines += ["---", "# VĂ‰GSĹ ĂśZLETI TERV", ""]
        field_labels = [
            ("piac_es_lehetoseg",    "## 1. Piaci helyzet Ă©s lehetĹ‘sĂ©g"),
            ("termek_es_ertekajanlat","## 2. TermĂ©k Ă©s Ă©rtĂ©kajĂˇnlat"),
            ("jelenlegi_allapot_korlatok", "## 3. Jelenlegi Ăˇllapot Ă©s korlĂˇtok"),
            ("bevĂ©teli_modell",      "## 4. BevĂ©teli modell"),
            ("gtm_strategia",        "## 5. Go-to-Market stratĂ©gia"),
            ("versenyelonyok",       "## 6. VersenyelĹ‘nyĂ¶k"),
            ("versenytarsak",        "## 7. VersenytĂˇrsak"),
        ]
        for key, heading in field_labels:
            val = bt.get(key)
            if val:
                lines += [heading, "", val, ""]
        if bt.get("kockazatok"):
            lines += ["## 8. KockĂˇzatok", ""]
            for k in bt["kockazatok"]: lines.append(f"- {k}")
            lines.append("")
        if bt.get("sikerkritĂ©riumok"):
            lines += ["## 9. SikerkritĂ©riumok", ""]
            for k in bt["sikerkritĂ©riumok"]: lines.append(f"- {k}")
            lines.append("")

    # â”€â”€ MegvalĂłsĂ­tĂˇsi terv fejezetek â”€â”€
    mt = synthesis.get("megvalositasi_terv", {})
    if mt:
        lines += ["---", "# VĂ‰GSĹ MEGVALĂ“SĂŤTĂSI TERV", ""]
        if mt.get("jelenlegi_architektura"):
            lines += ["## 1. Jelenlegi architektĂşra", "", mt["jelenlegi_architektura"], ""]
        if mt.get("cel_architektura"):
            lines += ["## 2. CĂ©larchitektĂşra", "", mt["cel_architektura"], ""]
        if mt.get("azonnali_teendok"):
            lines += ["## 3. Azonnali teendĹ‘k (0-2 hĂ©t)", ""]
            for t in mt["azonnali_teendok"]: lines.append(f"- {t}")
            lines.append("")
        for sprint_key, sprint_num in [("sprint_1","4"),("sprint_2","5"),("sprint_3","6"),("sprint_4_plus","7")]:
            sp = mt.get(sprint_key, {})
            if sp:
                label = f"Sprint {sprint_num}" if sprint_key != "sprint_4_plus" else "Sprint 4+"
                lines += [f"## {sprint_num}. {label} â€” {sp.get('fĂłkusz','')} ({sp.get('idotartam','')})", ""]
                for f in sp.get("feladatok",[]): lines.append(f"- {f}")
                if sp.get("definition_of_done"):
                    dod = sp['definition_of_done']
                    lines.append(f"\n**KĂ©sz amikor:** {dod}")
                lines.append("")
        if mt.get("technikai_kockazatok"):
            lines += ["## 8. Technikai kockĂˇzatok", "",
                      "| KockĂˇzat | ValĂłszĂ­nĹ±sĂ©g | MitigĂˇciĂł |",
                      "|----------|-------------|-----------|"]
            for r in mt["technikai_kockazatok"]:
                lines.append(f"| {r.get('kockĂˇzat','')} | {r.get('valĂłszĂ­nĹ±sĂ©g','')} | {r.get('mitigĂˇciĂł','')} |")
            lines.append("")

    # â”€â”€ AI kontextus blokk â”€â”€
    ai_ctx = synthesis.get("ai_context_block") or synthesis.get("md_context_block","")
    if ai_ctx:
        lines += ["---", "# AI KONTEXTUS BLOKK",
                  "> CsatolhatĂł system promptba vagy context windowba.", "", ai_ctx, ""]

    # â”€â”€ Konszenzus + dĂ¶ntĂ©sek â”€â”€
    lines += ["---", ""]
    if synthesis.get("consensus_points"):
        lines += ["## âś… Konszenzus-pontok", ""]
        for c in synthesis["consensus_points"]: lines.append(f"- ++ {c}")
        lines.append("")
    if synthesis.get("open_issues"):
        lines += ["## âť“ Nyitott kĂ©rdĂ©sek", ""]
        for o in synthesis["open_issues"]: lines.append(f"- ?? {o}")
        lines.append("")
    if synthesis.get("risk_register"):
        lines += ["## âš ď¸Ź KockĂˇzati naplĂł", ""]
        for r in synthesis["risk_register"]:
            lines.append(f"- [{r.get('severity','?').upper()}] **{r.get('risk','')}** â†’ {r.get('mitigation','')}")
        lines.append("")
    if synthesis.get("decision_log"):
        lines += ["## đź“‹ DĂ¶ntĂ©si naplĂł", ""]
        for d in synthesis["decision_log"]:
            conf = d.get('confidence',0)
            lines.append(f"- **{d.get('decision','')}** ({conf:.0%} bizalom) â€” {d.get('rationale','')}")
        lines.append("")
    if synthesis.get("recommendations"):
        lines += ["## đź“Ś AjĂˇnlĂˇsok", ""]
        for i,r in enumerate(synthesis["recommendations"],1): lines.append(f"{i}. {r}")
        lines.append("")
    if synthesis.get("implementation_plan"):
        lines += ["## đź—şď¸Ź ImplementĂˇciĂłs terv", ""]
        for s in synthesis["implementation_plan"]:
            pri = s.get('priority','?')
            lines.append(f"{s.get('step','')}.  [{pri.upper()}] **{s.get('action','')}** â€” {s.get('owner','')}")
        lines.append("")
    if synthesis.get("conclusion"):
        lines += ["## Ă–sszefoglalĂˇs", "", synthesis["conclusion"], ""]
    if revisions:
        lines += ["---", "## đź”„ ĂllĂˇspont-frissĂ­tĂ©sek", "",
                  "| Modell | Szerepe | VĂˇltoztatott? | Bizalom | VĂ©gsĹ‘ javaslat |",
                  "|--------|---------|--------------|---------|----------------|"]
        for rv in revisions.values():
            changed = "âś… Igen" if rv.get("changed_my_mind") else "âžˇď¸Ź Nem"
            conf = f"{rv.get('confidence',0):.0%}"
            lines.append(
                f"| {rv.get('model','')} | {rv.get('role','')} | {changed} | {conf} | {rv.get('final_recommendation','')[:60]} |"
            )
        lines.append("")

    Path(output_path).write_text("\n".join(lines), encoding="utf-8")
    print(f"âś… SzintĂ©zis MD: {output_path}")


def save_meeting_report(session_log: list, revisions: dict, synthesis: dict,
                         user_prompt: str, output_path: str):
    """
    Meeting Report â€” emberi olvasĂˇsra szĂˇnt vita-Ă¶sszefoglalĂł.
    Ki mit mondott, mibĹ‘l lett a konklĂşziĂł, dĂ¶ntĂ©si nyomvonal.
    """
    date_str = datetime.now().strftime('%Y-%m-%d %H:%M')
    participants = synthesis.get("participants", [])
    judge        = synthesis.get("judge", "")

    lines = [
        f"# Expert Council â€” Meeting Report",
        f"> DĂˇtum: {date_str} | Judge: {judge} | RĂ©sztvevĹ‘k: {', '.join(participants)}",
        "", "---", "",
        "## đźŽŻ MegbeszĂ©lĂ©s cĂ©lja", "", user_prompt, "", "---", "",
    ]
    task_profile = synthesis.get("task_profile") or RUN_METADATA.get("task_profile", {})
    if task_profile:
        lines += [
            "## TaskProfile",
            "",
            f"- **task_type:** {task_profile.get('task_type', 'unknown')}",
            f"- **target_audience:** {task_profile.get('target_audience', 'unknown')}",
            f"- **language:** {task_profile.get('language', 'unknown')}",
            f"- **recommended_scenario:** {task_profile.get('recommended_scenario', 'unknown')}",
            f"- **expected_deliverables:** {', '.join(task_profile.get('expected_deliverables') or [])}",
            f"- **required_perspectives:** {', '.join(task_profile.get('required_perspectives') or [])}",
            "",
            "---",
            "",
        ]
    contract = synthesis.get("output_contract") or RUN_METADATA.get("output_contract", {})
    if contract:
        deliverable_titles = [
            d.get("title", d.get("id", ""))
            for d in contract.get("deliverables", [])
        ]
        lines += [
            "## OutputContract",
            "",
            f"- **contract_id:** {contract.get('contract_id', 'unknown')}",
            f"- **title:** {contract.get('title', 'unknown')}",
            f"- **audience:** {contract.get('audience', 'unknown')}",
            f"- **output_format:** {contract.get('output_format', 'unknown')}",
            f"- **deliverables:** {', '.join(deliverable_titles)}",
            f"- **required_metadata:** {', '.join(contract.get('required_metadata') or [])}",
            "",
            "---",
            "",
        ]

    # RĂ©sztvevĹ‘k tĂˇblĂˇzat
    proc = synthesis.get("synthesis_process", {})
    if proc.get("summary"):
        lines += ["## Synthesis Process", "", proc["summary"], ""]
        if proc.get("steps"):
            lines += ["| Step | Type | Status | Size |",
                      "|------|------|--------|------|"]
            for step in proc["steps"]:
                status = "OK" if step.get("ok") else "FAILED"
                lines.append(
                    f"| {step.get('step','')} | {step.get('output_type','')} | "
                    f"{status} | {step.get('chars',0)} chars |"
                )
            lines.append("")
        lines += [
            f"**Final validation:** {proc.get('final_validation_status', synthesis.get('final_validation', {}).get('status', 'unknown'))}",
            f"**Human artifact:** {proc.get('human_artifact_status', 'unknown')}",
            f"**Structured metadata:** {proc.get('structured_metadata_status', 'unknown')}",
            "",
        ]
        if proc.get("structured_metadata_reason"):
            lines += [f"**Structured metadata note:** {proc.get('structured_metadata_reason')}", ""]
        if proc.get("metadata_failed_steps"):
            lines += [f"**Metadata failed steps:** {', '.join(proc.get('metadata_failed_steps', []))}", ""]
        repaired_sections = proc.get("repaired_sections") or []
        incomplete_sections = proc.get("incomplete_sections") or []
        if repaired_sections:
            lines += ["**Repaired sections:**", ""]
            for r in repaired_sections:
                if isinstance(r, dict):
                    lines.append(f"- {r.get('section')}: {', '.join(r.get('warnings_before', []))}")
                else:
                    lines.append(f"- {r}")
            lines.append("")
        if incomplete_sections:
            lines += ["**Incomplete sections after repair:**", ""]
            for r in incomplete_sections:
                if isinstance(r, dict):
                    lines.append(f"- {r.get('section')}: {', '.join(r.get('warnings', []))}")
                else:
                    lines.append(f"- {r}")
            lines.append("")
        missing = proc.get("missing_sections") or synthesis.get("final_validation", {}).get("missing_sections") or []
        weak_sections = proc.get("weak_sections") or synthesis.get("final_validation", {}).get("weak_sections") or []
        trunc = proc.get("truncation_warnings") or synthesis.get("final_validation", {}).get("truncation_warnings") or []
        if missing:
            lines += ["**Missing/weak sections:**", ""]
            lines += [f"- {m}" for m in missing]
            lines.append("")
        if weak_sections:
            lines += ["**Weak/incomplete sections:**", ""]
            lines += [f"- {m}" for m in weak_sections]
            lines.append("")
        if trunc:
            lines += ["**Truncation warnings:**", ""]
            for w in trunc:
                if isinstance(w, dict):
                    lines.append(f"- {w.get('section')}: {w.get('warning')}")
                else:
                    lines.append(f"- {w}")
            lines.append("")
        lines += ["---", ""]

    phase_speakers = {}
    for e in session_log:
        sp = e.get("speaker","?")
        ph = e.get("phase","?")
        if sp not in phase_speakers:
            phase_speakers[sp] = set()
        phase_speakers[sp].add(ph)

    lines += ["## đź‘Ą RĂ©sztvevĹ‘k Ă©s szerepek", "",
              "| RĂ©sztvevĹ‘ | Szerep | FĂˇzisok |",
              "|-----------|--------|---------|"]
    for e in session_log:
        pass  # collect unique speakers with roles
    seen = {}
    for e in session_log:
        sp = e.get("speaker","?")
        if sp not in seen:
            seen[sp] = e.get("role","?")
    for sp, role in seen.items():
        phases = ", ".join(sorted(phase_speakers.get(sp,set())))
        is_err = any(
            e.get("speaker")==sp and (e.get("error") or _is_failed_response(e.get("text","")))
            for e in session_log
        )
        status = "âš ď¸Ź rĂ©szleges (API hiba)" if is_err else "âś… aktĂ­v"
        lines.append(f"| {sp} | {role} | {phases} | {status} |")
    lines.append("")

    # FĂˇzisok
    def _valid(text):
        return _valid_response(text)

    phase_order = [
        ("evidence",    "âšˇ PHASE 0 â€” Evidence Pack"),
        ("independent", "đź”Ť PHASE 1 â€” FĂĽggetlen ĂˇllĂˇspontok"),
        ("issue_matrix","đź“Š PHASE 2 â€” ModerĂˇtor Issue Matrix"),
        ("rebuttal",    "âš”ď¸Ź  PHASE 3 â€” CĂ©lzott Rebuttal"),
        ("revision",    "đź”„ PHASE 4 â€” ĂllĂˇspont-frissĂ­tĂ©sek"),
        ("judge",       "âš–ď¸Ź  PHASE 5 â€” Final Judge"),
    ]

    for phase_key, phase_title in phase_order:
        entries = [e for e in session_log if e.get("phase")==phase_key]
        if not entries: continue

        lines += ["---", "", f"## {phase_title}", ""]

        if phase_key == "revision":
            # TĂˇblĂˇzat + rĂ©szletek
            lines += ["| RĂ©sztvevĹ‘ | Szerepe | VĂˇltoztatott? | Bizalom | Mit vĂˇltoztatott | VĂ©gsĹ‘ javaslat |",
                      "|-----------|---------|--------------|---------|------------------|----------------|"]
            for e in entries:
                try:
                    text = e.get("text","")
                    if "```" in text:
                        for part in text.split("```"):
                            p = part.strip().lstrip("json").strip()
                            if p.startswith("{"): text = p; break
                    rv = json.loads(text)
                    changed = "âś… Igen" if rv.get("changed_my_mind") else "âžˇď¸Ź Nem"
                    conf    = f"{rv.get('confidence',0.5):.0%}"
                    what    = (rv.get("what_changed") or "â€”")[:70]
                    rec     = (rv.get("final_recommendation") or "â€”")[:70]
                    lines.append(f"| {e['speaker']} | {rv.get('role','?')[:40]} | {changed} | {conf} | {what} | {rec} |")
                except:
                    err = (e.get("text","")[:60]).replace("\n"," ")
                    lines.append(f"| {e['speaker']} | â€” | âš ď¸Ź parse hiba | â€” | {err} | â€” |")
            lines.append("")

        elif phase_key == "judge":
            # Judge: verdict + dĂ¶ntĂ©sek kiemelve
            for e in entries:
                if not _valid(e.get("text","")): continue
                lines += [f"### {e['speaker']} (judge)", ""]
                lines.append(f"**VĂ©gsĹ‘ Ă­tĂ©let:** {synthesis.get('verdict','')}")
                lines.append("")
                if synthesis.get("consensus_points"):
                    lines += ["**Konszenzus-pontok:**"]
                    for c in synthesis["consensus_points"]: lines.append(f"- âś… {c}")
                    lines.append("")
                if synthesis.get("decision_log"):
                    lines += ["**DĂ¶ntĂ©si naplĂł:**"]
                    for d in synthesis["decision_log"]:
                        if isinstance(d, dict):
                            conf = d.get("confidence", 0)
                            try:
                                conf_text = f"{float(conf):.0%}"
                            except (TypeError, ValueError):
                                conf_text = "?"
                            lines.append(f"- [{conf_text}] **{d.get('decision','')}** â€” {d.get('rationale','')}")
                        else:
                            lines.append(f"- {d}")
                    lines.append("")
                if synthesis.get("recommendations"):
                    lines += ["**VĂ©gsĹ‘ ajĂˇnlĂˇsok:**"]
                    for i,r in enumerate(synthesis["recommendations"],1): lines.append(f"{i}. {r}")
                    lines.append("")

        else:
            for e in entries:
                text = e.get("text","")
                is_err = not _valid(text)
                lines += [f"### {e['speaker']} [{e.get('role','')}]", ""]
                if is_err:
                    lines += [f"> âš ď¸Ź **Nem vĂˇlaszolt** â€” {text[:120]}", ""]
                    continue
                # SzĂ¶veg: max 800 szĂł, utĂˇna jelzĂ©s
                words = text.split()
                excerpt = " ".join(words[:800])
                if len(words) > 800:
                    excerpt += f"\n\n*â€” {len(words)-800} szĂł kihagyva. Teljes szĂ¶veg: `debate_transcript.txt` â€”*"
                lines += [excerpt, ""]

    # Ă–sszesĂ­tett konklĂşziĂł
    lines += [
        "---", "",
        "## đź“Ś Ă–sszesĂ­tett konklĂşziĂł Ă©s dĂ¶ntĂ©sek", "",
        f"**Verdict:** {synthesis.get('verdict','')}", "",
    ]
    if synthesis.get("consensus_points"):
        lines += ["**Konszenzusok:**"]
        for c in synthesis["consensus_points"]: lines.append(f"- ++ {c}")
        lines.append("")
    if synthesis.get("open_issues"):
        lines += ["**Nyitott kĂ©rdĂ©sek:**"]
        for o in synthesis["open_issues"]: lines.append(f"- ?? {o}")
        lines.append("")
    if synthesis.get("risk_register"):
        lines += ["**KockĂˇzati naplĂł:**",
                  "| KockĂˇzat | SĂşlyossĂˇg | MitigĂˇciĂł |",
                  "|----------|-----------|-----------|"]
        for r in synthesis["risk_register"]:
            lines.append(f"| {r.get('risk','')} | {r.get('severity','').upper()} | {r.get('mitigation','')} |")
        lines.append("")
    if synthesis.get("implementation_plan"):
        lines += ["**ImplementĂˇciĂłs terv:**"]
        for s in synthesis["implementation_plan"]:
            lines.append(f"{s.get('step','')}. [{s.get('priority','?').upper()}] {s.get('action','')} â€” {s.get('owner','')}")
        lines.append("")

    lines += [
        "---", "",
        "> *Meeting Report â€” AI Expert Council*  ",
        f"> *{date_str} | Modellek: {', '.join(participants)} | Judge: {judge}*",
    ]

    Path(output_path).write_text("\n".join(lines), encoding="utf-8")
    print(f"âś… Meeting Report: {output_path}")


def _js(s): return json.dumps(str(s))

def create_docx(synthesis: dict, revisions: dict, output_path: str,
                session_log: list, items: list, user_prompt: str):
    def H(t,l=1): return f"new Paragraph({{heading:HeadingLevel.HEADING_{l},children:[new TextRun({{text:{_js(t)},bold:true}})]}}), "
    def P(t,c=None):
        col = f",color:{_js(c)}" if c else ""
        return f"new Paragraph({{children:[new TextRun({{text:{_js(t)}{col}}})]}}), "
    def SP(): return "new Paragraph({children:[new TextRun('')]}), "
    def BL(t): return f"new Paragraph({{numbering:{{reference:'b',level:0}},children:[new TextRun({_js(t)})]}}), "
    def NL(t): return f"new Paragraph({{numbering:{{reference:'n',level:0}},children:[new TextRun({_js(t)})]}}), "

    ch = []
    doc_date = synthesis.get("date","")
    participants = ", ".join(synthesis.get("participants",[]))
    judge = synthesis.get("judge","")

    ch.append(f"new Paragraph({{heading:HeadingLevel.TITLE,alignment:AlignmentType.CENTER,children:[new TextRun({{text:{_js(synthesis.get('title','SzintĂ©zis'))},bold:true,size:40}})]}}), ")
    ch.append(f"new Paragraph({{alignment:AlignmentType.CENTER,children:[new TextRun({{text:'Expert Council | Judge: {judge} | {doc_date}',color:'666666'}})]}}), ")
    ch.append(SP())

    if synthesis.get("verdict"):
        ch.append(H("VĂ©gsĹ‘ Ă­tĂ©let"))
        ch.append(P(synthesis["verdict"], "1a5276"))
        ch.append(SP())

    if user_prompt:
        ch.append(H("FelhasznĂˇlĂłi cĂ©l"))
        ch.append(P(synthesis.get("goal_summary", user_prompt[:300])))
        ch.append(SP())

    ch.append(H("RĂ©sztvevĹ‘k"))
    for p in synthesis.get("participants",[]): ch.append(BL(p))
    ch.append(SP())

    ch.append(H("ForrĂˇsanyagok"))
    for i in items: ch.append(BL(("đź–Ľď¸Ź " if i.is_img() else "đź“„ ")+i.name))
    ch.append(SP())

    if synthesis.get("introduction"):
        ch.append(H("BevezetĹ‘")); ch.append(P(synthesis["introduction"])); ch.append(SP())

    for sec in synthesis.get("sections",[]):
        ch.append(H(sec.get("heading",""),2))
        if sec.get("content"): ch.append(P(sec["content"]))
        for b in sec.get("bullets",[]): ch.append(BL(b))
        ch.append(SP())

    if synthesis.get("consensus_points"):
        ch.append(H("Konszenzus-pontok"))
        for c in synthesis["consensus_points"]: ch.append(BL("++ "+c))
        ch.append(SP())

    if synthesis.get("risk_register"):
        ch.append(H("KockĂˇzati naplĂł"))
        for r in synthesis["risk_register"]:
            ch.append(BL(f"[{r.get('severity','?').upper()}] {r.get('risk','')} â†’ {r.get('mitigation','')}"))
        ch.append(SP())

    if synthesis.get("decision_log"):
        ch.append(H("DĂ¶ntĂ©si naplĂł"))
        for d in synthesis["decision_log"]:
            ch.append(NL(f"{d.get('decision','')} ({d.get('confidence',0):.0%}) â€” {d.get('rationale','')}"))
        ch.append(SP())

    if synthesis.get("recommendations"):
        ch.append(H("AjĂˇnlĂˇsok"))
        for r in synthesis["recommendations"]: ch.append(NL(r))
        ch.append(SP())

    if synthesis.get("implementation_plan"):
        ch.append(H("ImplementĂˇciĂłs terv"))
        for s in synthesis["implementation_plan"]:
            ch.append(NL(f"[{s.get('priority','?').upper()}] {s.get('action','')} â€” {s.get('owner','')}"))
        ch.append(SP())

    if synthesis.get("conclusion"):
        ch.append(H("Ă–sszefoglalĂˇs")); ch.append(P(synthesis["conclusion"])); ch.append(SP())

    # Revision tĂˇblĂˇzat
    if revisions:
        ch.append(H("ĂllĂˇspont-frissĂ­tĂ©sek"))
        for rv in revisions.values():
            changed = "âś… VĂˇltoztatott" if rv.get("changed_my_mind") else "âžˇď¸Ź Tartja"
            ch.append(BL(
                f"{rv.get('model','')} [{rv.get('role','')}] | {changed} | "
                f"{rv.get('confidence',0):.0%} | {rv.get('final_recommendation','')[:80]}"
            ))
        ch.append(SP())

    js = f"""const {{Document,Packer,Paragraph,TextRun,HeadingLevel,AlignmentType,LevelFormat}}=require('docx');
const fs=require('fs');
const doc=new Document({{
  numbering:{{config:[
    {{reference:'b',levels:[{{level:0,format:LevelFormat.BULLET,text:'â€˘',alignment:AlignmentType.LEFT,style:{{paragraph:{{indent:{{left:720,hanging:360}}}}}}}}]}},
    {{reference:'n',levels:[{{level:0,format:LevelFormat.DECIMAL,text:'%1.',alignment:AlignmentType.LEFT,style:{{paragraph:{{indent:{{left:720,hanging:360}}}}}}}}]}},
  ]}},
  sections:[{{
    properties:{{page:{{size:{{width:11906,height:16838}},margin:{{top:1440,right:1440,bottom:1440,left:1440}}}}}},
    children:[{chr(10).join(ch)}]
  }}]
}});
Packer.toBuffer(doc).then(buf=>{{
  fs.writeFileSync({_js(output_path)},buf);
  console.log('OK');
}}).catch(e=>{{console.error('HIBA:',e.message);process.exit(1);}});
"""
    jp = os.path.join(tempfile.gettempdir(), "_debate_gen.js")
    Path(jp).write_text(js, encoding="utf-8")
    r = subprocess.run(["node",jp], capture_output=True, text=True)
    if r.returncode != 0: print(f"âťŚ Node: {r.stderr}"); return False
    print(f"âś… Word: {output_path}"); return True


def save_logs(session_log, out_dir, user_prompt):
    os.makedirs(out_dir, exist_ok=True)
    lp = os.path.join(out_dir,"debate_log.json")
    payload = {"user_prompt": user_prompt, "debate": session_log}
    if RUN_METADATA:
        payload["metadata"] = RUN_METADATA
    with open(lp,"w",encoding="utf-8") as f:
        json.dump(payload,f,ensure_ascii=False,indent=2)
    print(f"đź“‹ JSON: {lp}")
    tp = os.path.join(out_dir,"debate_transcript.txt")
    with open(tp,"w",encoding="utf-8") as f:
        f.write(f"EXPERT COUNCIL ĂTIRAT\n{'='*64}\n{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        if user_prompt: f.write(f"\nCĂ‰L:\n{user_prompt}\n")
        cur_phase = ""
        for e in session_log:
            if e.get("phase") != cur_phase:
                cur_phase = e.get("phase","")
                f.write(f"\n{'â•'*64}\n  PHASE: {cur_phase.upper()}\n{'â•'*64}\n")
            f.write(f"\n{'â”€'*40}\n{e.get('speaker','')} [{e.get('role','')}]\n{'â”€'*40}\n{e.get('text','')}\n")
    print(f"đź“„ Ătirat: {tp}")


def estimate_cost(items, scenario_key, quality, role_overrides):
    scenario = SCENARIOS[scenario_key]
    base = dict(scenario["quality_map"][quality])
    base.update(role_overrides)
    total_chars = sum(len(i.text) for i in items if not i.is_img())
    ctx_tok = total_chars // 4
    phases = scenario["phases"]

    print(_bold("\nđź’° KĂ–LTSĂ‰GBECSLĂ‰S"))
    print("â”€"*56)
    print(f"  ForrĂˇs: {total_chars:,} kar â‰ {ctx_tok:,} token")
    print(f"  Scenario: {scenario_key} | Quality: {quality}")
    print(f"  FĂˇzisok: {' â†’ '.join(phases)}")
    print("â”€"*56)

    total = 0.0
    phase_calls = {
        "evidence": [("moderator",1)],
        "independent": [(r,1) for r in base if r not in ("moderator","judge","debaters")],
        "issue_matrix": [("moderator",1)],
        "rebuttal": [(r,1) for r in base if r not in ("moderator","judge","debaters")],
        "revision": [(r,1) for r in base if r not in ("moderator","judge","debaters")],
        "judge": [("judge",1)],
    }
    # quick: debaters
    if "debaters" in base:
        for p in ("independent","rebuttal","revision"):
            if p in phase_calls:
                phase_calls[p] = [(f"debater{i+1}",1) for i in range(len(base["debaters"]))]

    for phase in phases:
        for role, calls in phase_calls.get(phase,[]):
            mk = base.get(role)
            if isinstance(mk, list): mk = mk[0]
            if not mk or mk not in CATALOG: continue
            cfg = CATALOG[mk]
            mid = cfg["id_fast"] if quality=="fast" else cfg["id"]
            in_tok  = (ctx_tok // 3) * calls
            out_tok = 600 * calls
            cost = (in_tok * cfg["pin"] + out_tok * cfg["pout"]) / 1_000_000
            total += cost
            print(f"  {cfg['emoji']} {phase:15s} {cfg['label']:15s} ({mid:30s}): ~${cost:.3f}")

    print("â”€"*56)
    print(_bold(f"  Ă–SSZESEN: ~${total:.3f}"))
    if quality != "fast": print(_yellow("  đź’ˇ --quality fast: ~10-20x olcsĂłbb"))
    print("â”€"*56)

def run_smoke_test() -> int:
    print("Smoke test: local validation only, no provider calls.")
    checks = []
    checks.append(("catalog has quick scenario", "quick" in SCENARIOS))
    checks.append(("gpt-fast uses OpenAI Responses", CATALOG["gpt-fast"]["type"] == "openai_responses"))
    checks.append(("failed marker excluded", _is_failed_response("[ERROR Test: missing API key]")))
    checks.append(("normal response accepted", _valid_response("OK - short useful answer")))
    checks.append(("synthesis token budget configured", SYNTHESIS_MAX_OUTPUT_TOKENS > 0))
    checks.append(("minimum participant count configured", MIN_VALID_PARTICIPANTS >= 1))

    business_profile = build_task_profile(
        "KĂ©szĂ­ts ĂĽzleti tervet, GTM stratĂ©giĂˇt Ă©s megvalĂłsĂ­tĂˇsi tervet egy AI termĂ©khez.",
        [], "hu", "quick",
    )
    checks.append(("business strategy prompt", business_profile.task_type == "business_strategy"))
    checks.append(("business profile preserves current deliverables", "implementation_plan" in business_profile.expected_deliverables and "ai_context_block" in business_profile.expected_deliverables))

    architecture_profile = build_task_profile(
        "Review this technical architecture and decide the best tradeoffs for the backend.",
        [], "en", "expert-council",
    )
    checks.append(("technical architecture prompt", architecture_profile.task_type == "architecture_decision"))

    generic_profile = build_task_profile(
        "Elemezd rĂ¶viden a mellĂ©kelt anyagot Ă©s foglald Ă¶ssze a lĂ©nyeget.",
        [], "hu", "quick",
    )
    checks.append(("generic analysis prompt", generic_profile.task_type == "general_analysis"))

    original_infer = globals()["_infer_task_type"]
    try:
        globals()["_infer_task_type"] = lambda user_prompt, items=None: (_ for _ in ()).throw(RuntimeError("mock profile failure"))
        fallback_profile = build_task_profile("anything", [], "hu", "quick")
        checks.append(("task profile fallback", fallback_profile.task_type == "general_analysis" and fallback_profile.profile_source == "deterministic_fallback"))
    finally:
        globals()["_infer_task_type"] = original_infer

    profile_tmp = tempfile.mkdtemp(prefix="task_profile_smoke_")
    try:
        default_contract = default_output_contract("hu")
        business_contract = load_output_contract(str(Path("contracts") / "business_master_plan.json"), "hu")
        technical_contract = load_output_contract(str(Path("contracts") / "technical_audit.json"), "hu")
        general_contract = load_output_contract(str(Path("contracts") / "general_analysis.json"), "hu")
        checks.append(("default contract works", default_contract.contract_id == "default_business_master_plan" and len(default_contract.deliverables) >= 4))
        checks.append(("business_master_plan contract works", business_contract.contract_id == "business_master_plan" and any(d["id"] == "business_plan" for d in business_contract.deliverables)))
        checks.append(("technical_audit contract works", technical_contract.contract_id == "technical_audit" and any(d["id"] == "technical_findings" for d in technical_contract.deliverables)))
        checks.append(("general_analysis contract works", general_contract.contract_id == "general_analysis" and any(d["id"] == "analysis" for d in general_contract.deliverables)))

        invalid_contract_path = Path(profile_tmp) / "invalid_contract.json"
        invalid_contract_path.write_text('{"contract_id":"broken","title":"Broken"}', encoding="utf-8")
        invalid_stopped = False
        try:
            import contextlib, io
            with contextlib.redirect_stdout(io.StringIO()):
                load_output_contract(str(invalid_contract_path), "hu")
        except SystemExit as e:
            invalid_stopped = e.code == 2
        checks.append(("invalid contract fails with clear error", invalid_stopped))

        save_task_profile_only(business_profile, profile_tmp, business_profile.user_goal, business_contract)
        profile_path = Path(profile_tmp) / "task_profile.json"
        log_path = Path(profile_tmp) / "debate_log.json"
        saved_profile = json.loads(profile_path.read_text(encoding="utf-8"))
        saved_log = json.loads(log_path.read_text(encoding="utf-8"))
        checks.append(("--task-profile-only mode", saved_profile.get("task_type") == "business_strategy" and saved_log.get("metadata", {}).get("task_profile", {}).get("task_type") == "business_strategy"))
        checks.append(("contract summary is stored in debate_log.json", saved_log.get("metadata", {}).get("output_contract", {}).get("contract_id") == "business_master_plan"))

        report_path = Path(profile_tmp) / "meeting_report.md"
        save_meeting_report(
            [],
            {},
            {
                "participants": [],
                "judge": "Smoke Judge",
                "verdict": "ok",
                "task_profile": task_profile_to_dict(business_profile),
                "output_contract": output_contract_summary(technical_contract),
                "synthesis_process": {
                    "summary": "Smoke summary",
                    "final_validation_status": "ok",
                    "human_artifact_status": "ok",
                    "structured_metadata_status": "ok",
                },
            },
            business_profile.user_goal,
            str(report_path),
        )
        report_text = report_path.read_text(encoding="utf-8")
        checks.append(("contract summary appears in meeting_report.md", "## OutputContract" in report_text and "technical_audit" in report_text))
    finally:
        shutil.rmtree(profile_tmp, ignore_errors=True)

    def ok_info(eng):
        return {
            "ok": True, "detail": "mock ok",
            "provider": eng.provider, "model_key": eng.key,
            "model_id": eng.mid, "label": eng.label,
        }

    def fail_info(eng, msg="mock failure"):
        return {
            "ok": False, "detail": msg,
            "provider": eng.provider, "model_key": eng.key,
            "model_id": eng.mid, "label": eng.label,
            "exception_class": "MockHealthError",
            "exception_message": msg,
        }

    def quick_engines():
        return {
            "moderator": AIEngine("claude-sonnet", "fast"),
            "debater1": AIEngine("gpt-fast", "fast"),
            "debater2": AIEngine("gemini-fast", "fast"),
            "judge": AIEngine("gpt-fast", "fast"),
        }

    original_check = globals()["_check_engine_health"]
    try:
        engines = quick_engines()
        health = {role: ok_info(eng) for role, eng in engines.items()}
        health["debater2"] = fail_info(engines["debater2"], "Gemini mock failure")
        globals()["_check_engine_health"] = lambda eng, skip_network=False: (
            ok_info(eng) if eng.key == "claude-sonnet" else fail_info(eng)
        )
        resolve_role_fallbacks(engines, health, "quick", "fast", skip_network=True)
        checks.append(("Gemini fails, Claude fallback succeeds", engines["debater2"].key == "claude-sonnet"))

        engines = quick_engines()
        health = {role: ok_info(eng) for role, eng in engines.items()}
        health["judge"] = fail_info(engines["judge"], "Judge mock failure")
        globals()["_check_engine_health"] = lambda eng, skip_network=False: (
            ok_info(eng) if eng.key in ("gpt-best", "gpt", "claude-sonnet", "gemini-fast") else fail_info(eng)
        )
        resolve_role_fallbacks(engines, health, "quick", "fast", skip_network=True)
        checks.append(("judge fails, judge fallback succeeds", engines["judge"].key == "gpt-best"))

        engines = quick_engines()
        health = {role: ok_info(eng) for role, eng in engines.items()}
        health["debater2"] = fail_info(engines["debater2"], "Gemini mock failure")
        globals()["_check_engine_health"] = lambda eng, skip_network=False: fail_info(eng, "all fallback candidates fail")
        stopped = False
        try:
            resolve_role_fallbacks(engines, health, "quick", "fast", skip_network=True)
        except SystemExit as e:
            stopped = e.code == 2
        checks.append(("all fallback candidates fail, pipeline stops clearly", stopped))
    finally:
        globals()["_check_engine_health"] = original_check

    class FakeJudge:
        label = "Fake Judge"
        mid = "fake-judge"
        captured = ""
        captured_all = []
        def call(self, system, messages, stream=True, label=None, max_output_tokens=None):
            self.captured = messages[0]["content"]
            self.captured_all.append(self.captured)
            return json.dumps({
                "title": "Smoke",
                "date": "2026-05-27",
                "goal_summary": "",
                "participants": [],
                "judge": self.label,
                "verdict": "ok",
                "consensus_points": [],
                "open_issues": [],
                "risk_register": [],
                "decision_log": [],
                "recommendations": [],
                "implementation_plan": [],
                "introduction": "",
                "sections": [],
                "conclusion": "",
                "md_context_block": "",
            })

    class FakeEngine:
        def __init__(self, label):
            self.label = label

    fake_judge = FakeJudge()
    fake_engines = {
        "debater1": FakeEngine("Valid One"),
        "debater2": FakeEngine("Failed One"),
        "debater3": FakeEngine("Valid Two"),
        "judge": fake_judge,
    }
    phase_judge(
        fake_judge, "evidence", "matrix",
        {"debater1": "VALID_A", "debater2": "[ERROR FAILED_TEXT]", "debater3": "VALID_B"},
        {}, fake_engines, {}, "", ["Szintezis"], "en", []
    )
    checks.append((
        "failed participants are not included in final judge input",
        "FAILED_TEXT" not in "\n".join(fake_judge.captured_all)
        and "VALID_A" in "\n".join(fake_judge.captured_all)
        and "VALID_B" in "\n".join(fake_judge.captured_all),
    ))

    truncated_task_md = """# SzintĂ©zis

## VĂ©gsĹ‘ Ă­tĂ©let

Ez egy kellĹ‘en hosszĂş vĂ©gsĹ‘ Ă­tĂ©let konkrĂ©t dĂ¶ntĂ©sekkel.

## VĂ©gsĹ‘ ĂĽzleti terv

### Ăśzleti ajĂˇnlĂˇsok

Az utolsĂł ajĂˇnlĂˇs egy validĂˇciĂłs eszk...

## VĂ©gsĹ‘ megvalĂłsĂ­tĂˇsi terv

#### Task 1: Auth, RBAC Ă©s szervezeti modell
**Objective:** JogosultsĂˇgi modell kialakĂ­tĂˇsa.
**Priority:** high
**Recommended phase/sprint:** Sprint 1
**Dependencies:** szervezeti modell
**Concrete steps:**
- JogosultsĂˇgi mĂˇtrix megtervezĂ©se
**Definition of Done:**
- minden

## AI kontextus blokk

Ez egy elĂ©g hosszĂş AI kontextus blokk a kĂ©sĹ‘bbi agenteknek.

## KockĂˇzatok Ă©s dĂ¶ntĂ©sek

A kockĂˇzatok Ă©s dĂ¶ntĂ©sek rĂ©sz legalĂˇbb rĂ¶viden szerepel.

## KĂ¶vetkezĹ‘ lĂ©pĂ©sek

1. ElsĹ‘ konkrĂ©t kĂ¶vetkezĹ‘ lĂ©pĂ©s vĂ©grehajtĂˇsa.
"""
    truncated_validation = _validate_synthesis_markdown(truncated_task_md)
    warnings_blob = json.dumps(truncated_validation, ensure_ascii=False)
    checks.append(("truncated task detection", "task_1_incomplete_concrete_steps" in warnings_blob))
    checks.append(("incomplete Definition of Done detection", "task_1_incomplete_definition_of_done" in warnings_blob))
    checks.append(("mid-word ending detection", "possible_mid_word_ending" in warnings_blob))
    checks.append(("validation should not be ok when truncation exists", truncated_validation["status"] != "ok"))

    metadata_engine = SynthesisEngine(
        fake_judge, "evidence", "matrix",
        {"debater1": "VALID_A", "debater3": "VALID_B"}, {},
        fake_engines, {}, "", ["Szintezis"], "hu", []
    )
    empty_outline = {
        "date": "2026-05-27",
        "consensus_points": ["stale value"],
        "open_issues": ["stale value"],
        "risk_register": [{"risk": "stale value"}],
        "decision_log": [{"decision": "stale value"}],
        "recommendations": ["stale value"],
    }
    metadata_engine._apply_empty_metadata(empty_outline, "structured metadata extraction failed")
    metadata_keys = ("consensus_points", "open_issues", "risk_register", "decision_log", "recommendations")
    checks.append((
        "structured metadata deterministic empty fallback",
        empty_outline["_structured_metadata_status"] == "failed"
        and empty_outline["_metadata_available"] is False
        and all(empty_outline[k] == [] for k in metadata_keys),
    ))

    clean_final_md = metadata_engine._build_final_markdown(
        empty_outline,
        "A vĂ©gsĹ‘ Ă­tĂ©let kellĹ‘en rĂ©szletes Ă©s hasznĂˇlhatĂł.",
        "### TermĂ©k Ă©s Ă©rtĂ©kajĂˇnlat\n\nA fĹ‘ ĂĽzleti terv teljes, tiszta Ă©s nem tartalmaz heurisztikus metaadat-listĂˇkat.",
        "#### Task 1: StabilizĂˇlĂˇs\n**Objective:** Stabil kimenet.\n**Priority:** high\n**Recommended phase/sprint:** Sprint 1\n**Dependencies:** nincsenek\n**Concrete steps:**\n- ValidĂˇciĂł megtartĂˇsa\n- Riport ellenĹ‘rzĂ©se\n**Definition of Done:**\n- A dokumentum tiszta\n- A log stĂˇtuszai elkĂĽlĂ¶nĂĽlnek",
        "Ez az AI kontextus blokk elĂ©g hosszĂş ahhoz, hogy a validĂˇciĂł elfogadja.",
        {},
    )
    clean_validation = metadata_engine._validate_final_markdown(clean_final_md)
    checks.append(("metadata failure does not create garbage bullets", "## KockĂˇzatok Ă©s dĂ¶ntĂ©sek" not in clean_final_md and "## KĂ¶vetkezĹ‘ lĂ©pĂ©sek" not in clean_final_md))
    checks.append(("metadata failure note is clean", "metaadatok nem" in clean_final_md.lower()))
    checks.append(("random extracted fragments absent", "KĂ¶vetkeztetĂ©s:**" not in clean_final_md and "Startupok Ă©s prototĂ­pus-fĂˇzisĂş csapatok**" not in clean_final_md))
    checks.append(("human artifact status separated from metadata", clean_validation["human_artifact_status"] in ("ok", "ok_with_warnings")))
    checks.append(("structured metadata status remains separate", empty_outline["_structured_metadata_status"] == "failed"))

    metadata_engine.step_results = [
        {"step": "main_deliverable_draft", "output_type": "markdown", "ok": True, "chars": 120},
        {"step": "implementation_plan_draft", "output_type": "markdown", "ok": True, "chars": 120},
        {"step": "ai_context_block", "output_type": "markdown", "ok": True, "chars": 120},
        {"step": "structured_metadata_extraction", "output_type": "json", "ok": False, "chars": 80},
    ]
    assembled = metadata_engine._assemble(
        {
            "title": "Smoke",
            "date": "2026-05-27",
            "goal_summary": "Smoke test",
            "verdict": "A vĂ©gsĹ‘ Ă­tĂ©let kellĹ‘en rĂ©szletes Ă©s hasznĂˇlhatĂł.",
            "_structured_metadata_status": "failed",
            "_metadata_available": False,
            "_structured_metadata_reason": "structured metadata extraction failed",
        },
        "### TermĂ©k Ă©s Ă©rtĂ©kajĂˇnlat\n\nA fĹ‘ ĂĽzleti terv teljes, tiszta Ă©s ellenĹ‘rizhetĹ‘ tartalmat ad a felhasznĂˇlĂłnak.",
        "#### Task 1: StabilizĂˇlĂˇs\n**Objective:** Stabil, tiszta kimenet biztosĂ­tĂˇsa.\n**Priority:** high\n**Recommended phase/sprint:** Sprint 1\n**Dependencies:** meglĂ©vĹ‘ szintĂ©zis motor\n**Concrete steps:**\n- A validĂˇciĂłs stĂˇtuszok elkĂĽlĂ¶nĂ­tĂ©se\n- A metaadat fallback ĂĽresen tartĂˇsa\n**Definition of Done:**\n- A final markdown nem tartalmaz szemĂ©t listĂˇkat\n- A debate log kĂĽlĂ¶n metaadat stĂˇtuszt tĂˇrol",
        "Ez az AI kontextus blokk elĂ©g hosszĂş Ă©s konkrĂ©t ahhoz, hogy egy kĂ©sĹ‘bbi agent hasznos munkakĂ¶rnyezetet kapjon.",
        {"passed": True, "issues": [], "fixes": [], "final_warnings": []},
    )
    checks.append(("metadata failure does not fail clean human artifact", assembled["synthesis_process"]["final_validation_status"] == "ok_with_warnings"))
    checks.append(("metadata failed step tracked separately", assembled["synthesis_process"]["failed_steps"] == [] and assembled["synthesis_process"]["metadata_failed_steps"] == ["structured_metadata_extraction"]))

    technical_contract = load_output_contract(str(Path("contracts") / "technical_audit.json"), "en")
    contract_engine = SynthesisEngine(
        fake_judge, "evidence", "matrix",
        {"debater1": "VALID_A", "debater3": "VALID_B"}, {},
        fake_engines, {}, "", ["Szintezis"], "en", [],
        output_contract=technical_contract,
    )
    contract_results = {
        "final_verdict": {
            "id": "final_verdict", "title": "Audit verdict", "source": "verdict",
            "required": True, "sections": [],
            "markdown": """The AgentReady codebase is usable, but it needs focused reliability hardening before broader expansion.

## Audit verdict

**Overall assessment:** the highest-priority risks are contract validation drift, provider failure handling, and incomplete audit remediation sections.

- The audit verdict identifies the highest-priority technical risks and confirms the remediation direction.
- It ties the recommendation to evidence from the debate and names the expected engineering outcome.""",
        },
        "technical_findings": {
            "id": "technical_findings", "title": "Technical findings", "source": "main",
            "required": True, "sections": technical_contract.deliverables[1]["sections"],
            "markdown": """### Architecture assessment
- The architecture assessment identifies service boundaries, integration risks, and the highest-impact coupling to reduce first.
- It recommends a staged target architecture so the team can improve reliability without blocking current delivery.

### Maintainability and code health
- The maintainability review calls out duplicated logic, unclear ownership, and areas where tests should lock behavior.
- It recommends refactoring only around validated seams so the audit does not become a broad rewrite.

### Security and compliance risks
- The security review prioritizes authentication, authorization, secret handling, and provider error logging.
- It recommends explicit controls and regression checks before any larger platform expansion.

### Operational readiness
- The operational readiness review covers observability, retry behavior, failure modes, and release verification.
- It recommends actionable health checks and smoke tests as the minimum reliable operating baseline.

## VĂ©gsĹ‘ ĂĽzleti terv
This legacy heading must not leak into the assembled contract output.""",
        },
        "remediation_plan": {
            "id": "remediation_plan", "title": "Remediation plan", "source": "implementation",
            "required": True, "sections": technical_contract.deliverables[2]["sections"],
            "markdown": """### Critical fixes
- Fix blocking validation failures and add regression coverage for every required contract section.
- Remove placeholder-producing paths from final assembly and verify failed repairs stay visible in logs.

### Near-term hardening
- Add focused smoke tests for health checks, contract output, metadata status, and final validation.
- Improve provider error visibility and keep failed model outputs excluded from all judge inputs.

### Long-term improvements
- Move contracts toward reusable configuration once the Python contract path is stable and verified.
- Add richer audit-specific contracts after the existing CLI behavior remains stable across real runs.""",
        },
        "ai_handoff": {
            "id": "ai_handoff", "title": "AI handoff context", "source": "ai_context",
            "required": False, "sections": [],
            "markdown": "## AI handoff context\n\n- The next agent should preserve the contract-first synthesis path and avoid default-section leakage.\n- It should verify final validation status, repaired sections, and contract deliverables in debate_log.json.",
        },
    }
    contract_md = contract_engine._build_contract_markdown_from_results(
        {"date": "2026-05-27"},
        contract_results,
    )
    contract_validation = contract_engine._validate_contract_markdown(contract_md)
    checks.append(("technical contract markdown validates ok", contract_validation["status"] == "ok"))
    checks.append(("duplicate deliverable heading is stripped", contract_md.count("## Audit verdict") == 1))
    checks.append(("no-section deliverable with nested headings validates correctly", "final_verdict" not in contract_validation.get("weak_sections", [])))
    checks.append(("no duplicated AI handoff sections", contract_md.count("## AI handoff context") == 1))
    checks.append(("no legacy default headings appended to contract output", "## VĂ©gsĹ‘ ĂĽzleti terv" not in contract_md and "## VĂ©gsĹ‘ megvalĂłsĂ­tĂˇsi terv" not in contract_md))

    fenced_prompt = """## AI context block

```markdown
# Agent prompt

Use the contract-aware synthesis result as input.

#### Task 1: Preserve contract validation
Keep this fenced prompt intact.
```
"""
    fence_warnings = contract_engine._truncation_warnings(fenced_prompt)
    checks.append(("balanced fenced code block does not trigger truncation warning", "possible_mid_word_ending" not in fence_warnings and not any("task_1_" in w for w in fence_warnings)))
    bold_heading_transition = "5. **only then expand into enterprise features or stronger automation claims.**\n\n### Risk judgment\n\n- Concrete risk item with enough words for validation."
    checks.append(("bold sentence before heading is not abrupt truncation", "abrupt_transition_before_heading" not in contract_engine._truncation_warnings(bold_heading_transition)))

    business_contract_validation_engine = SynthesisEngine(
        fake_judge, "evidence", "matrix",
        {"debater1": "VALID_A", "debater3": "VALID_B"}, {},
        fake_engines, {}, "", ["Szintezis"], "hu", [],
        output_contract=load_output_contract(str(Path("contracts") / "business_master_plan.json"), "hu"),
    )
    business_deliverables = business_contract_validation_engine.output_contract["deliverables"]
    business_contract_results = {
        "final_verdict": {
            "id": "final_verdict", "title": business_deliverables[0]["title"], "source": "verdict",
            "required": True, "sections": [],
            "markdown": "- A terv validĂˇlhatĂł piaci irĂˇnyt Ă©s fokozatos megvalĂłsĂ­tĂˇst javasol.\n- A dĂ¶ntĂ©s a gyors tanulĂˇst, a stabil mĹ±kĂ¶dĂ©st Ă©s a kontrollĂˇlt kockĂˇzatkezelĂ©st helyezi elĹ‘tĂ©rbe.",
        },
        "business_plan": {
            "id": "business_plan", "title": business_deliverables[1]["title"], "source": "main",
            "required": True, "sections": business_deliverables[1]["sections"],
            "markdown": "\n\n".join(
                f"### {section['title']}\n"
                f"- A(z) {section['title']} rĂ©sz konkrĂ©t, szerzĹ‘dĂ©s szerinti ĂĽzleti irĂˇnyt Ă©s mĂ©rhetĹ‘ dĂ¶ntĂ©si szempontokat ad.\n"
                f"- A(z) {section['title']} rĂ©sz mĂˇsodik pontja validĂˇlhatĂł kĂ¶vetkezmĂ©nyt, felelĹ‘st Ă©s ellenĹ‘rzĂ©si fĂłkuszt nevez meg."
                for section in business_deliverables[1]["sections"]
            ),
        },
        "implementation_plan": {
            "id": "implementation_plan", "title": business_deliverables[2]["title"], "source": "implementation",
            "required": True, "sections": business_deliverables[2]["sections"],
            "markdown": "\n\n".join(
                f"### {section['title']}\n"
                f"- A(z) {section['title']} rĂ©sz vĂ©grehajthatĂł feladatokat, sorrendet Ă©s ellenĹ‘rzĂ©si pontokat tartalmaz.\n"
                f"- A(z) {section['title']} rĂ©sz mĂˇsodik pontja megadja a stabil CLI-kompatibilitĂˇs Ă©s szerzĹ‘dĂ©ses validĂˇciĂł feltĂ©teleit."
                for section in business_deliverables[2]["sections"]
            ),
        },
        "ai_context_block": {
            "id": "ai_context_block", "title": business_deliverables[3]["title"], "source": "ai_context",
            "required": True, "sections": [],
            "markdown": """```markdown
# KĂ¶vetkezĹ‘ agent kontextus

A kĂ¶vetkezĹ‘ agent tartsa meg a szerzĹ‘dĂ©s szerinti struktĂşrĂˇt, ellenĹ‘rizze a sprint roadmapet, Ă©s ne adjon hozzĂˇ nem kĂ©rt alapĂ©rtelmezett szakaszokat.
```""",
        },
    }
    business_contract_md = business_contract_validation_engine._build_contract_markdown_from_results(
        {"date": "2026-05-27"},
        business_contract_results,
    )
    business_contract_validation = business_contract_validation_engine._validate_contract_markdown(business_contract_md)
    checks.append(("business_master_plan contract validates ok", business_contract_validation["status"] == "ok"))

    placeholder_md = """# Technical Audit Report

## Audit verdict

- The audit verdict is concrete enough for validation.
- It identifies the remediation path and target audience.

## Technical findings

### Architecture assessment
- Architecture finding one has enough concrete detail for validation.
- Architecture finding two has enough concrete detail for validation.

### Maintainability and code health
- Maintainability finding one has enough concrete detail for validation.
- Maintainability finding two has enough concrete detail for validation.

### Security and compliance risks
- Security finding one has enough concrete detail for validation.
- Security finding two has enough concrete detail for validation.

### Operational readiness
- Operations finding one has enough concrete detail for validation.
- Operations finding two has enough concrete detail for validation.

## Remediation plan

### Critical fixes

_Ez a szakasz nem kĂ©szĂĽlt el teljesen._

### Near-term hardening

TBD

### Long-term improvements

This section was not completed
"""
    placeholder_validation = contract_engine._validate_contract_markdown(placeholder_md)
    checks.append(("placeholder sections are detected as incomplete", placeholder_validation["status"] == "failed" and set(["critical_fixes", "near_term", "long_term"]).issubset(set(placeholder_validation.get("weak_sections", [])))))
    checks.append(("weak section issues are not reported as missing headings", not set(["critical_fixes", "near_term", "long_term"]).intersection(set(placeholder_validation.get("missing_sections", [])))))
    checks.append(("debate_log validation fields exist when weak sections exist", bool(placeholder_validation.get("weak_sections")) and "validation_reasons" in placeholder_validation and "section_issues" in placeholder_validation))
    failed_reason_fields = (
        placeholder_validation.get("missing_sections")
        or placeholder_validation.get("truncation_warnings")
        or placeholder_validation.get("weak_sections")
        or placeholder_validation.get("incomplete_sections")
        or placeholder_validation.get("validation_reasons")
    )
    checks.append(("failed validation has a clear reason", placeholder_validation["human_artifact_status"] != "failed" or bool(failed_reason_fields)))

    class RepairJudge:
        label = "Repair Judge"
        mid = "repair-judge"
        def call(self, system, messages, stream=True, label=None, max_output_tokens=None):
            prompt = messages[0]["content"]
            if "- id: critical_fixes" in prompt:
                return "### Critical fixes\n- Patch the contract section extraction path and add regression tests for placeholder sections.\n- Verify health-check failures never enter final judge input and keep provider errors actionable.\n- Re-run technical audit output validation before publishing the final artifact."
            if "- id: near_term" in prompt:
                return "### Near-term hardening\n- Add smoke tests for contract remediation sections, metadata status, and repaired section logs.\n- Improve observability around contract deliverable generation and repair attempts.\n- Document manual checks for synthesis_output.md, meeting_report.md, and debate_log.json."
            if "- id: long_term" in prompt:
                return "### Long-term improvements\n- Move stable contract definitions into reusable configuration after the CLI path is proven.\n- Add richer technical audit contracts for security, reliability, and architecture decisions.\n- Keep default synthesis compatibility covered before expanding scenario configuration."
            return "- Repair response contains concrete fallback action one.\n- Repair response contains concrete fallback action two."

    repair_engine = SynthesisEngine(
        RepairJudge(), "evidence summary with concrete audit facts", "matrix",
        {"debater1": "VALID_A", "debater3": "VALID_B"}, {},
        fake_engines, {}, "technical audit", ["Szintezis"], "en", [],
        output_contract=technical_contract,
    )
    weak_results = {
        "final_verdict": contract_results["final_verdict"],
        "technical_findings": contract_results["technical_findings"],
        "remediation_plan": {
            "id": "remediation_plan", "title": "Remediation plan", "source": "implementation",
            "required": True, "sections": technical_contract.deliverables[2]["sections"],
            "markdown": """### Critical fixes
_Ez a szakasz nem kĂ©szĂĽlt el teljesen._

### Near-term hardening
TBD

### Long-term improvements
This section was not completed""",
        },
        "ai_handoff": contract_results["ai_handoff"],
    }
    repair_engine._repair_contract_deliverables(
        {"date": "2026-05-27", "verdict": "Repair smoke"},
        weak_results,
    )
    repaired_md = repair_engine._build_contract_markdown_from_results({"date": "2026-05-27"}, weak_results)
    repaired_validation = repair_engine._validate_contract_markdown(repaired_md)
    checks.append(("weak sections trigger targeted repair", any(s.startswith("contract_section_repair:remediation_plan") for s in repair_engine.repaired_steps)))
    checks.append(("repaired sections are merged without duplicating headings", repaired_md.count("### Critical fixes") == 1 and repaired_md.count("### Near-term hardening") == 1 and repaired_md.count("### Long-term improvements") == 1))
    checks.append(("technical_audit becomes ok after repair", repaired_validation["status"] == "ok"))

    contract_engine.step_results = [
        {"step": "contract_deliverable:final_verdict", "output_type": "markdown", "ok": True, "chars": 80},
        {"step": "contract_deliverable:technical_findings", "output_type": "markdown", "ok": True, "chars": 300},
        {"step": "contract_deliverable:remediation_plan", "output_type": "markdown", "ok": True, "chars": 240},
        {"step": "contract_deliverable:ai_handoff", "output_type": "markdown", "ok": True, "chars": 80},
    ]
    contract_assembled = contract_engine._assemble_contract(
        {
            "title": "Technical Audit Report",
            "date": "2026-05-27",
            "goal_summary": "Smoke test",
            "verdict": "The audit verdict is sufficiently detailed and contract-specific.",
            "consensus_points": ["valid consensus"],
            "open_issues": [],
            "risk_register": [{"risk": "risk", "severity": "medium", "mitigation": "mitigation"}],
            "decision_log": [{"decision": "decision", "rationale": "rationale", "confidence": 0.8}],
            "recommendations": ["recommendation"],
            "_structured_metadata_status": "ok",
            "_metadata_available": True,
        },
        contract_results,
        {"passed": True, "issues": [], "fixes": [], "final_warnings": []},
    )
    checks.append(("contract full assembly status ok", contract_assembled["synthesis_process"]["final_validation_status"] == "ok"))
    checks.append(("contract steps replace legacy synthesis steps", all(not s.startswith(("main_deliverable_draft", "implementation_plan_draft", "ai_context_block")) for s in contract_assembled["synthesis_process"]["failed_steps"])))
    checks.append(("final status derivation is consistent", contract_assembled["synthesis_process"]["human_artifact_status"] != "failed" and not contract_assembled["synthesis_process"].get("weak_sections")))
    no_failure_reasons = not any((
        contract_assembled["synthesis_process"].get("missing_sections"),
        contract_assembled["synthesis_process"].get("truncation_warnings"),
        contract_assembled["synthesis_process"].get("weak_sections"),
        contract_assembled["synthesis_process"].get("incomplete_sections"),
        contract_assembled["synthesis_process"].get("failed_steps"),
        contract_assembled["synthesis_process"].get("validation_reasons"),
    ))
    checks.append(("empty failure reasons cannot produce failed human status", not (no_failure_reasons and contract_assembled["synthesis_process"]["human_artifact_status"] == "failed")))

    default_status_engine = SynthesisEngine(
        fake_judge, "evidence", "matrix",
        {"debater1": "VALID_A", "debater3": "VALID_B"}, {},
        fake_engines, {}, "default status smoke", ["Szintezis"], "hu", [],
    )
    default_status_engine.step_results = [
        {"step": "final_decision_outline_markdown_fallback", "output_type": "markdown", "ok": False, "chars": 200, "note": "truncation_detected"},
        {"step": "final_decision_outline", "output_type": "json", "ok": True, "chars": 200, "note": "markdown_metadata_fallback"},
        {"step": "main_deliverable_draft", "output_type": "markdown", "ok": True, "chars": 900},
        {"step": "implementation_plan_draft", "output_type": "markdown", "ok": True, "chars": 900},
        {"step": "ai_context_block", "output_type": "markdown", "ok": True, "chars": 300},
        {"step": "structured_metadata_extraction", "output_type": "json", "ok": True, "chars": 300},
        {"step": "consistency_review", "output_type": "json", "ok": True, "chars": 120},
    ]
    default_status_engine.repaired_steps = ["final_decision_outline"]
    default_outline = {
        "title": "Szintezis",
        "date": "2026-05-27",
        "goal_summary": "Default status smoke",
        "verdict": "A vegso itelet reszletes, hasznalhato es a javitott dontesi vazlatra epul.",
        "consensus_points": ["valid consensus"],
        "open_issues": [],
        "risk_register": [{"risk": "risk", "severity": "medium", "mitigation": "mitigation"}],
        "decision_log": [{"decision": "decision", "rationale": "rationale", "confidence": 0.8}],
        "recommendations": ["recommendation"],
        "_structured_metadata_status": "ok",
        "_metadata_available": True,
    }
    default_assembled = default_status_engine._assemble(
        default_outline,
        "### Termek es ertekajanlat\n\nA default uzleti terv teljes es konkret.\n\n### Piac es go-to-market\n\nA piac es az ajanlat ertheto, valid es hasznalhato.",
        "### Roadmap\n\n- Stabilizalni kell a validacios statuszok szamitasat a default es contract utak kozott.\n- Meg kell tartani a javitott dontesi vazlatot ervenyes emberi kimenetkent.\n- A mentesi reteg nem adhat Minosegi megjegyzest, ha a vegso statusz ok.\n\n### Technikai kockazatok\n\n- A nem javitott hibak tovabbra is sikertelen statuszt kapnak.\n- A javitott fallback csak akkor torlodik a failed listabol, ha van sikeres javito lepes.",
        "Ez az AI kontextus blokk eleg hosszu ahhoz, hogy a validacio elfogadja, es lezart mondattal fejezodik be.",
        {"passed": True, "issues": [], "fixes": [], "final_warnings": []},
    )
    checks.append(("default repaired fallback is not a failed human step", default_assembled["synthesis_process"]["failed_steps"] == []))
    checks.append(("default validation uses consistent final status", default_assembled["synthesis_process"]["final_validation_status"] == "ok" and default_assembled["synthesis_process"]["human_artifact_status"] == "ok"))

    save_tmp = tempfile.mkdtemp(prefix="contract_save_smoke_")
    try:
        output_md = Path(save_tmp) / "synthesis_output.md"
        save_markdown(
            contract_assembled,
            [],
            "technical audit smoke",
            [],
            {},
            str(output_md),
        )
        saved_text = output_md.read_text(encoding="utf-8")
        default_output_md = Path(save_tmp) / "default_synthesis_output.md"
        save_markdown(
            default_assembled,
            [],
            "default smoke",
            [],
            {},
            str(default_output_md),
        )
        default_saved_text = default_output_md.read_text(encoding="utf-8")
        checks.append(("no Minosegi megjegyzes for repaired default validation", "Minőségi megjegyzés" not in default_saved_text))
        checks.append(("no MinĹ‘sĂ©gi megjegyzĂ©s when validation is ok", "MinĹ‘sĂ©gi megjegyzĂ©s" not in saved_text))
    finally:
        shutil.rmtree(save_tmp, ignore_errors=True)

    failed = [name for name, ok in checks if not ok]
    for name, ok in checks:
        print(f"   {'OK' if ok else 'FAIL'} - {name}")
    if failed:
        print("Smoke test failed: " + ", ".join(failed))
        return 1
    print("Smoke test passed.")
    return 0


# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Main
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def main():
    global SYNTHESIS_MAX_OUTPUT_TOKENS
    p = argparse.ArgumentParser(
        description="Expert Council â€” Multi-AI strukturĂˇlt vita, 5 fĂˇzis.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Scenariok:
  quick           3 modell, 1 kor + szintezis (default)
  expert-council  5 fazis: evidence->fuggetelen->matrix->rebuttal->revision->judge
  red-team        1 strategist + 1 engineer + 2 skeptic + judge

Peldak:
  python ai_debate.py project.zip --prompt-file feladat.txt
  python ai_debate.py --folder ./docs --prompt "..." --scenario expert-council
  python ai_debate.py ... --quality best --parallel --scenario red-team
  python ai_debate.py ... --roles judge=gpt,skeptic=deepseek --estimate
  python ai_debate.py ... --resume eredmenyek/debate_log.json
"""
    )
    p.add_argument("sources", nargs="*", metavar="FORRAS")
    p.add_argument("--folder",     metavar="MAPPA")
    pp = p.add_mutually_exclusive_group()
    pp.add_argument("--prompt",      metavar="SZOVEG")
    pp.add_argument("--prompt-file", metavar="FAJL")
    p.add_argument("--scenario",
                   default=os.environ.get("DEFAULT_SCENARIO","quick"),
                   choices=list(SCENARIOS.keys()),
                   help="Vita scenario (default: quick / env: DEFAULT_SCENARIO)")
    p.add_argument("--quality",
                   default=os.environ.get("DEFAULT_QUALITY","balanced"),
                   choices=["fast","balanced","best"],
                   help="Modell minĹ‘sĂ©g (default: balanced / env: DEFAULT_QUALITY)")
    p.add_argument("--roles",      metavar="ROLE=MODEL,...",
                   help="Szerepek felĂĽlĂ­rĂˇsa pl: --roles strategist=gpt-best,judge=claude-opus")
    p.add_argument("--parallel",
                   action="store_true",
                   default=os.environ.get("DEFAULT_PARALLEL","false").lower()=="true",
                   help="Phase 1 pĂˇrhuzamos API hĂ­vĂˇsok (env: DEFAULT_PARALLEL)")
    p.add_argument("--output",     default="synthesis_output.docx")
    p.add_argument("--output-dir", default=".", metavar="KONYVTAR")
    p.add_argument("--lang",       default="hu", choices=["hu","en"])
    p.add_argument("--max-chars",  type=int, default=DEFAULT_MAX_CHARS)
    p.add_argument("--estimate",   action="store_true")
    p.add_argument("--no-docx",    action="store_true")
    p.add_argument("--resume",     metavar="LOG.json")
    p.add_argument("--skip-health-check", action="store_true",
                   help="Csak API kulcs jelenletet ellenoriz; nem hiv provider health checket")
    p.add_argument("--health-check-only", action="store_true",
                   help="Provider/model health check utan kilep, vita nelkul")
    p.add_argument("--synthesis-max-output-tokens", type=int,
                   default=SYNTHESIS_MAX_OUTPUT_TOKENS,
                   help="Final Judge max output token (env: SYNTHESIS_MAX_OUTPUT_TOKENS)")
    p.add_argument("--smoke-test", action="store_true",
                   help="Lokalis smoke test API hivas nelkul")
    p.add_argument("--task-profile-only", action="store_true",
                   help="Forras/prompt betoltes es TaskProfile generalas vita nelkul")
    p.add_argument("--contract-file", metavar="CONTRACT.json",
                   help="OutputContract JSON fajl a vegso dokumentum szerkezetehez")
    args = p.parse_args()

    SYNTHESIS_MAX_OUTPUT_TOKENS = args.synthesis_max_output_tokens

    if args.smoke_test:
        raise SystemExit(run_smoke_test())

    # Prompt
    user_prompt = ""
    if args.prompt_file:
        if not os.path.exists(args.prompt_file):
            print(f"âťŚ {args.prompt_file}"); sys.exit(1)
        user_prompt = Path(args.prompt_file).read_text(encoding="utf-8").strip()
        print(f"đź“ť Prompt: {args.prompt_file} ({len(user_prompt)} kar)")
    elif args.prompt:
        user_prompt = args.prompt.strip()

    # Role overrides
    role_overrides = {}
    if args.roles:
        for part in args.roles.split(","):
            if "=" in part:
                k, v = part.strip().split("=", 1)
                role_overrides[k.strip()] = v.strip()
        if role_overrides:
            print(f"đźŽ­ Szerepek felĂĽlĂ­rva: {role_overrides}")

    if args.health_check_only:
        engines = build_roles(args.scenario, args.quality, role_overrides)
        original_mapping = _role_mapping(engines)
        health = run_health_checks(engines, skip_network=args.skip_health_check)
        resolution = resolve_role_fallbacks(
            engines, health, args.scenario, args.quality,
            skip_network=args.skip_health_check,
        )
        print_resolved_role_mapping(engines, original_mapping)
        print("Health check passed.")
        return

    # ForrĂˇs
    all_sources = list(args.sources)
    if args.folder: all_sources.append(args.folder)
    if not all_sources and not args.task_profile_only:
        print("âťŚ Adj meg forrĂˇst!"); p.print_help(); sys.exit(1)

    # BetĂ¶ltĂ©s
    if all_sources:
        print(f"\n{'â•'*64}\nđź“Ą  FORRASANYAGOK\n{'â•'*64}")
        items = load_sources(all_sources, args.max_chars)
    else:
        print(_yellow("\n[--task-profile-only: nincs forras, csak prompt alapjan keszul profil]"))
        items = []

    RUN_METADATA.clear()
    task_profile = build_task_profile(user_prompt, items, args.lang, args.scenario)
    RUN_METADATA["task_profile"] = task_profile_to_dict(task_profile)
    print_task_profile(task_profile)
    output_contract = load_output_contract(args.contract_file, task_profile.language)
    RUN_METADATA["output_contract"] = output_contract_summary(output_contract)
    if args.task_profile_only:
        save_task_profile_only(task_profile, args.output_dir, user_prompt, output_contract)
        return

    # BecslĂ©s
    estimate_cost(items, args.scenario, args.quality, role_overrides)
    if args.estimate:
        print(_yellow("\n[--estimate: csak becslĂ©s, futtatĂˇs kihagyva]")); return

    # Engines
    engines = build_roles(args.scenario, args.quality, role_overrides)
    if "moderator" not in engines:
        print("âťŚ ModerĂˇtor nĂ©lkĂĽl nem futhat."); sys.exit(1)
    original_mapping = _role_mapping(engines)
    health = run_health_checks(engines, skip_network=args.skip_health_check)
    resolution = resolve_role_fallbacks(
        engines, health, args.scenario, args.quality,
        skip_network=args.skip_health_check,
    )
    drop_unhealthy_participants(engines, health)
    resolved_mapping = _role_mapping(engines)
    print_resolved_role_mapping(engines, original_mapping)
    RUN_METADATA.update({
        "scenario": args.scenario,
        "quality": args.quality,
        "required_debate_participants": _required_participant_count(args.scenario),
        "original_role_mapping": original_mapping,
        "resolved_role_mapping": resolved_mapping,
        "health": health,
        "fallback_resolution": resolution,
    })

    # Output dir
    os.makedirs(args.output_dir, exist_ok=True)

    # Resume
    resume_log = None
    if args.resume:
        if not os.path.exists(args.resume):
            print(f"âťŚ Resume log: {args.resume}"); sys.exit(1)
        with open(args.resume, encoding="utf-8") as f:
            rd = json.load(f)
        resume_log = rd.get("debate", rd if isinstance(rd,list) else [])
        print(_yellow(f"\nâ–¶ď¸Ź  Resume: {len(resume_log)} bejegyzĂ©s"))

    # Scenario futtatĂˇs
    scenario_key = args.scenario
    print(f"\nđźš€ Scenario: {scenario_key} | {SCENARIOS[scenario_key]['desc']}")

    if scenario_key == "quick":
        session_log, synthesis = run_quick(
            engines, items, user_prompt, args.lang, args.output_dir,
            task_profile, output_contract)
        revisions = {}
    else:  # expert-council + red-team ugyanaz a logika
        session_log, synthesis = run_expert_council(
            engines, items, user_prompt, args.lang,
            args.parallel, args.output_dir, resume_log, task_profile,
            output_contract)
        revisions = {e["role"]: json.loads(e["text"])
                     for e in session_log
                     if e.get("phase")=="revision"
                     and e["text"].strip().startswith("{")
                     }

    # Kimenetek
    save_logs(session_log, args.output_dir, user_prompt)

    stem = Path(args.output).stem
    md_path = os.path.join(args.output_dir, f"{stem}.md")
    save_markdown(synthesis, items, user_prompt, session_log, revisions, md_path)

    # Meeting Report â€” vita emberi olvasĂˇsra szĂˇnt Ă¶sszefoglalĂłja
    report_path = os.path.join(args.output_dir, f"{stem}_meeting_report.md")
    save_meeting_report(session_log, revisions, synthesis, user_prompt, report_path)

    if not args.no_docx:
        create_docx(synthesis, revisions,
                    os.path.join(args.output_dir, args.output),
                    session_log, items, user_prompt)

    print(f"\n{'â•'*64}")
    print(_bold("đźŽ‰  KĂ‰SZ!"))
    print(f"{'â•'*64}")
    if not args.no_docx: print(f"  đź“„ Word:           {os.path.join(args.output_dir, args.output)}")
    print(f"  đź“ť SzintĂ©zis MD:   {md_path}")
    print(f"  đź“‹ Meeting Report: {report_path}")
    log_path = os.path.join(args.output_dir, "debate_log.json")
    print(f"  đź—„ď¸Ź  Log:            {log_path}")
    transcript_path = os.path.join(args.output_dir, "debate_transcript.txt")
    print(f"  đź“ Ătirat:         {transcript_path}")
    print()
    print(_yellow("  SzintĂ©zis MD: Ăśzleti terv + MegvalĂłsĂ­tĂˇsi terv + AI kontextus blokk"))
    print(_yellow("  Meeting Report: Ki mit mondott, dĂ¶ntĂ©si nyomvonal, konklĂşziĂł"))

if __name__ == "__main__":
    main()
